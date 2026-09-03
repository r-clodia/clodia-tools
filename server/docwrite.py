"""Da Markdown in CriticMarkup a DOCX con REVISIONI TRACCIATE vere.

L'altra metà di `docrev`. Quello legge un DOCX revisionato e lo rende in
CriticMarkup; questo prende il CriticMarkup che l'agente ha prodotto e lo
riscrive come DOCX in cui `{++...++}` e `{--...--}` sono `<w:ins>` e `<w:del>`
di Word — revisioni che la controparte apre e può accettare o rifiutare una per
una, firmate con autore e data.

Perché serve un verbo e non uno script dell'agente: l'avvocato non ha `Bash`
(`native_tools` non dichiarato → solo il pavimento dell'archseed;
`allow_shell_cmds: []` nella sandbox), quindi non può eseguire `python-docx`.
Senza questo verbo l'unica strada era delegare a `sysadmin`, che ha una shell —
cioè aggirare il least-privilege passando da un agente PIÙ privilegiato, e far
manipolare un contratto a chi non ha il contesto per capirlo. Un agente che
tratta documenti `untrusted` di controparte non deve avere una shell; deve avere
un verbo che fa la cosa specifica.

Come, e perché così: si costruisce lo scheletro con `python-docx` — che genera
`[Content_Types].xml`, gli stili e le relazioni corrette, cose che a mano si
sbagliano in silenzio — e poi si INIETTANO gli elementi di revisione nell'albero
`lxml` sottostante. `python-docx` non sa produrre revisioni tracciate (non ha
API per `<w:ins>`/`<w:del>`), ma espone `paragraph._p`, e da lì si lavora
sull'XML vero. Le due metà insieme danno un pacchetto valido con dentro
revisioni autentiche.
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

#: Marcatori CriticMarkup riconosciuti. L'ordine conta: la sostituzione
#: `{~~vecchio~>nuovo~~}` va provata prima delle forme semplici, o il suo
#: contenuto verrebbe letto come testo normale.
_SOSTITUZIONE = re.compile(r"\{~~(.*?)~>(.*?)~~\}", re.S)
_INSERIMENTO = re.compile(r"\{\+\+(.*?)\+\+\}", re.S)
_CANCELLAZIONE = re.compile(r"\{--(.*?)--\}", re.S)
_COMMENTO = re.compile(r"\{>>(.*?)<<\}", re.S)

#: Un commento porta spesso la firma che `docrev` ci ha messo: " [autore — data]".
_FIRMA = re.compile(r"\s*\[([^\]]{0,120}?)\]\s*$")


def _qn(tag: str) -> str:
    return "{%s}%s" % (W, tag)


class Segmento:
    """Un pezzo di paragrafo con il suo stato di revisione.

    `tipo` ∈ `testo` | `ins` | `del` | `commento`.
    """

    __slots__ = ("tipo", "testo")

    def __init__(self, tipo: str, testo: str):
        self.tipo = tipo
        self.testo = testo

    def __repr__(self) -> str:  # pragma: no cover - diagnostica
        return f"Segmento({self.tipo!r}, {self.testo!r})"


def segmenta(riga: str) -> list[Segmento]:
    """Spezza una riga di Markdown nei suoi segmenti di revisione.

    Si scorre la riga UNA volta trovando il marcatore più a sinistra, invece di
    applicare le quattro regex in sequenza: applicandole a turno, il testo
    normale fra due marcatori finirebbe in un ordine diverso da quello del
    documento — e in una revisione l'ordine è il significato (`{--1.000--}`
    seguito da `{++1.500++}` è una sostituzione; invertito è un'altra proposta).
    """
    fuori: list[Segmento] = []
    resto = riga
    while resto:
        candidati = []
        for regex, tipo in ((_SOSTITUZIONE, "sost"), (_INSERIMENTO, "ins"),
                            (_CANCELLAZIONE, "del"), (_COMMENTO, "commento")):
            m = regex.search(resto)
            if m:
                candidati.append((m.start(), m, tipo))
        if not candidati:
            if resto:
                fuori.append(Segmento("testo", resto))
            break
        candidati.sort(key=lambda c: c[0])
        inizio, m, tipo = candidati[0]
        if inizio:
            fuori.append(Segmento("testo", resto[:inizio]))
        if tipo == "sost":
            # `{~~vecchio~>nuovo~~}` è esattamente una cancellazione seguita da
            # un inserimento: si espande così invece di inventare un terzo tipo,
            # perché è ciò che Word rappresenta.
            if m.group(1):
                fuori.append(Segmento("del", m.group(1)))
            if m.group(2):
                fuori.append(Segmento("ins", m.group(2)))
        elif m.group(1):
            fuori.append(Segmento(tipo, m.group(1)))
        resto = resto[m.end():]
    return fuori


#: Separatore di cella: una pipe NON preceduta da backslash. Serve una regex e
#: non `split("|")` perché `docrev` scrive `\|` per una pipe che sta nel testo, e
#: uno split ingenuo spezzerebbe la cella in due lasciando un backslash orfano.
_PIPE = re.compile(r"(?<!\\)\|")


def _riga_tabella(riga: str) -> list[str] | None:
    """Le celle di una riga di tabella Markdown, o None se non lo è."""
    s = riga.strip()
    if not s.startswith("|"):
        return None
    corpo = s[1:]
    if corpo.endswith("|") and not corpo.endswith("\\|"):
        corpo = corpo[:-1]
    return [c.strip().replace("\\|", "|") for c in _PIPE.split(corpo)]


def _e_separatore(riga: str) -> bool:
    celle = _riga_tabella(riga)
    return bool(celle) and all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in celle if c.strip())


class _Blocco:
    """Un blocco di documento: paragrafo, titolo, voce di elenco o tabella."""

    __slots__ = ("genere", "livello", "righe", "celle")

    def __init__(self, genere: str, livello: int = 0):
        self.genere = genere      # paragrafo | titolo | elenco | tabella
        self.livello = livello
        self.righe: list[str] = []
        self.celle: list[list[str]] = []


def blocchi_da_markdown(md: str) -> list[_Blocco]:
    """Markdown → blocchi. Riconosce solo ciò che `docrev` produce.

    Deliberatamente ristretto: titoli ATX, elenchi `-`, tabelle in pipe e
    paragrafi. Un parser Markdown completo qui sarebbe superficie in più senza
    un caso d'uso — questo verbo esiste per richiudere il giro su un documento
    che è passato da `docrev`, e le due grammatiche devono combaciare.
    """
    fuori: list[_Blocco] = []
    righe = (md or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(righe):
        riga = righe[i]
        nuda = riga.strip()
        if not nuda:
            i += 1
            continue
        # Commento HTML (il riepilogo di docrev, i marcatori di pagina): non è
        # contenuto del contratto e non va nel DOCX.
        if nuda.startswith("<!--"):
            while i < len(righe) and "-->" not in righe[i]:
                i += 1
            i += 1
            continue
        if nuda.startswith("---") and set(nuda) <= {"-"}:
            i += 1
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", nuda)
        if m:
            b = _Blocco("titolo", len(m.group(1)))
            b.righe.append(m.group(2))
            fuori.append(b)
            i += 1
            continue
        if _riga_tabella(riga) is not None:
            b = _Blocco("tabella")
            while i < len(righe) and _riga_tabella(righe[i]) is not None:
                if not _e_separatore(righe[i]):
                    b.celle.append(_riga_tabella(righe[i]) or [])
                i += 1
            if b.celle:
                fuori.append(b)
            continue
        m = re.match(r"^[-*+]\s+(.*)$", nuda)
        if m:
            b = _Blocco("elenco")
            b.righe.append(m.group(1))
            fuori.append(b)
            i += 1
            continue
        # Paragrafo: righe consecutive non vuote, riunite con uno spazio.
        b = _Blocco("paragrafo")
        pezzi = []
        while i < len(righe) and righe[i].strip() and not re.match(
                r"^\s*(#{1,6}\s|[-*+]\s|\|)", righe[i]) and not righe[i].strip().startswith("<!--"):
            pezzi.append(righe[i].strip())
            i += 1
        b.righe.append(" ".join(pezzi))
        fuori.append(b)
    return fuori


def _run_xml(testo: str, cancellato: bool):
    """Un `<w:r>` con `<w:t>` o `<w:delText>` secondo lo stato.

    Il testo cancellato va in `<w:delText>` e non in `<w:t>`: è la differenza
    che permette a Word di mostrarlo come rimosso invece che come contenuto, ed
    è anche ciò su cui si basa la rilettura in `docrev`.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if cancellato else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = testo
    r.append(t)
    return r


def _revisione_xml(indice: int, tipo: str, autore: str, data: str, testo: str):
    """Un `<w:ins>` o `<w:del>` completo, con il run dentro."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    el = OxmlElement("w:ins" if tipo == "ins" else "w:del")
    el.set(qn("w:id"), str(indice))
    el.set(qn("w:author"), autore)
    el.set(qn("w:date"), data)
    el.append(_run_xml(testo, cancellato=(tipo == "del")))
    return el


def _scrivi_segmenti(paragrafo, segmenti: list[Segmento], stato: dict) -> None:
    """Appende i segmenti al paragrafo, marcando le revisioni.

    Si lavora su `paragraph._p` — l'elemento lxml — perché `python-docx` non
    ha API per le revisioni: `add_run()` sa creare `<w:r>`, non sa avvolgerlo in
    un `<w:ins>`. Costruire il documento con la libreria e iniettare qui è ciò
    che tiene insieme un pacchetto valido e revisioni autentiche.
    """
    for seg in segmenti:
        if not seg.testo:
            continue
        if seg.tipo == "testo":
            paragrafo.add_run(seg.testo)
        elif seg.tipo in ("ins", "del"):
            stato["n"] += 1
            paragrafo._p.append(_revisione_xml(
                stato["n"], seg.tipo, stato["autore"], stato["data"], seg.testo))
            stato["conteggio"][seg.tipo] += 1
        elif seg.tipo == "commento":
            # I commenti di Word (`comments.xml`) non sono raggiungibili con
            # python-docx senza montare a mano una part e le sue relazioni: una
            # superficie fragile per un contenuto secondario. Si rende come nota
            # in corsivo, marcata in modo da restare riconoscibile a chi legge e
            # da non passare per testo del contratto.
            testo = seg.testo.strip()
            firma = _FIRMA.search(testo)
            if firma:
                testo = testo[:firma.start()].strip()
                nota = f" [Nota — {firma.group(1)}: {testo}]"
            else:
                nota = f" [Nota: {testo}]"
            run = paragrafo.add_run(nota)
            run.italic = True
            stato["conteggio"]["commento"] += 1


def markdown_to_docx(md: str, autore: str = "Studio", data: str | None = None
                     ) -> tuple[bytes, dict]:
    """Markdown in CriticMarkup → (bytes del DOCX, statistiche).

    `autore` e `data` finiscono negli attributi `w:author`/`w:date` di ogni
    revisione: sono ciò che permette alla controparte di vedere CHI ha proposto
    una modifica e QUANDO, che è metà del valore di una revisione tracciata.
    """
    from docx import Document
    from docx.shared import Pt

    quando = data or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    stato = {"n": 0, "autore": autore, "data": quando,
             "conteggio": {"ins": 0, "del": 0, "commento": 0}}

    doc = Document()
    for blocco in blocchi_da_markdown(md):
        if blocco.genere == "tabella":
            larghezza = max(len(r) for r in blocco.celle)
            tabella = doc.add_table(rows=0, cols=larghezza)
            tabella.style = "Table Grid"
            for riga in blocco.celle:
                celle = tabella.add_row().cells
                for k in range(larghezza):
                    testo = riga[k] if k < len(riga) else ""
                    par = celle[k].paragraphs[0]
                    _scrivi_segmenti(par, segmenta(testo), stato)
            continue
        if blocco.genere == "titolo":
            par = doc.add_heading("", level=min(blocco.livello, 6))
        elif blocco.genere == "elenco":
            par = doc.add_paragraph("", style="List Bullet")
        else:
            par = doc.add_paragraph("")
            par.paragraph_format.space_after = Pt(6)
        _scrivi_segmenti(par, segmenta(blocco.righe[0] if blocco.righe else ""), stato)

    buf = io.BytesIO()
    doc.save(buf)
    stats = {
        "inserimenti": stato["conteggio"]["ins"],
        "cancellazioni": stato["conteggio"]["del"],
        "commenti": stato["conteggio"]["commento"],
        "revisioni": stato["conteggio"]["ins"] + stato["conteggio"]["del"],
        "autore": autore,
        "data": quando,
    }
    return buf.getvalue(), stats


def docx_name(path: str) -> str:
    """Nome del `.docx` di destinazione dal path del Markdown sorgente."""
    base = (path or "").rsplit("/", 1)[-1]
    radice = base.rsplit(".", 1)[0] if "." in base else base
    return f"{radice or 'documento'}.docx"
