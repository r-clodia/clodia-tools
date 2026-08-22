"""MCP stdio server entry point — Clodia tools gateway."""
import asyncio
import json
import logging
import sys

from mcp.server.lowlevel import Server
from mcp.server.stdio import stdio_server
from mcp.types import TextContent, Tool

from . import __version__
from . import origin
from . import instance_profile
from . import proxy
from . import taint as _taint
from . import telemetry as _tlm
from . import transfer_channel
from .tools import email, fs, logs, runtime
from .tools import web_fetch, web_post

#: `LOG` era usato in cinque punti di questo modulo e **definito in nessuno**.
#: Ognuno di quei punti sta in un `except` — cioè si scopre solo quando qualcosa
#: è già andato storto, e allora il `NameError` sostituisce l'errore vero con uno
#: che non c'entra. Trovato usando davvero un token MCP umano: due verbi nuovi non
#: erano nella tabella della provenienza, il ramo di ripiego ha provato a
#: scriverlo nel log, e il verbo ha risposto «NameError» a chi chiedeva le proprie
#: menzioni. Nessun test lo vedeva: tutti e cinque i rami sono percorsi solo in
#: presenza di un altro guasto.
LOG = logging.getLogger("clodia-tools.main")
from .tools import eu_corpus
from .whitelist import (agent_config, agent_denies, agent_gates, agent_name,
                        outside_profile,
                        current_chat, current_clearance, current_human_role,
                        current_principal, current_scoped_tools, is_on_behalf,
                        message_kind,
                        current_channel,
                        current_origin,
                        is_unattended)

import os as _os
from .topics.service import TopicService, TopicError
from .topics.local_fs import LocalFsStorage
from .topics.storage import VersionConflict

app = Server("clodia-tools")

# Topic System v2 (P1): storage local-fs in un'area dedicata del datadir del
# gateway (NUOVA e separata dai topic git esistenti). Enforcement tiering OFF in
# P1 (arriva in P3). Reference monitor: gli agenti toccano i topic solo da qui.
# Default in un'area montata SOLO dal gateway (la dir del vault: l'agent-server
# NON la monta) → gli agenti non possono raggiungere i file dei topic by-passando
# i verbi. Reference monitor: l'unica via ai topic è il gateway. Override via
# CLODIA_TOPICS_ROOT.
_TOPICS_ROOT = _os.environ.get("CLODIA_TOPICS_ROOT", "/datadir/clodia-vault/topics-store")
_topic_svc: TopicService | None = None


def _topics() -> TopicService:
    global _topic_svc
    if _topic_svc is None:
        _topic_svc = TopicService(LocalFsStorage(_TOPICS_ROOT))
    return _topic_svc




_EMAIL_TOOLS: list[Tool] = [
    Tool(
        name="email.send",
        description=(
            "Send an email. Specify the sender mailbox via `account` (the account "
            "name, e.g. as shown by email.folders or your instructions). "
            "Plain-text body, optional CC and local file attachments."
        ),
        inputSchema={
            "type": "object",
            "properties": {
                "to": {"type": "string", "description": "recipient email address"},
                "subject": {"type": "string"},
                "body": {"type": "string"},
                "account": {
                    "type": "string",
                    "description": "sender mailbox account name (required if you have more than one)",
                },
                "cc": {"type": "string", "description": "optional CC address"},
                "attachments": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "optional local file paths to attach",
                },
                "topic_files": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": ("allegati PER RIFERIMENTO: path di file del topic "
                                    "(es. 'files/contratto.pdf'). Il gateway li legge e "
                                    "li allega senza che il contenuto entri nel tuo "
                                    "contesto — non serve scaricarli prima, e non "
                                    "consumano token."),
                },
                "tier": {"type": "string", "description": ("tier del topic per topic_files. "
                          "OPZIONALE: senza, si usa il topic del canale in cui stai "
                          "lavorando")},
                "name": {"type": "string", "description": "nome del topic (opzionale, vedi tier)"},
            },
            "required": ["to", "subject", "body"],
        },
    ),
    Tool(
        name="email.folders",
        description="List the IMAP folders of an account (pass the account name in `account`).",
        inputSchema={
            "type": "object",
            "properties": {
                "account": {"type": "string",
                            "description": "account name to inspect"},
            },
        },
    ),
    Tool(
        name="email.list",
        description="List messages of a folder (default INBOX) for a configured account.",
        inputSchema={
            "type": "object",
            "properties": {
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
                "limit": {"type": "integer", "description": "max messages, default 10"},
            },
        },
    ),
    Tool(
        name="email.read",
        description="Read a single message by its IMAP id from a folder (default INBOX).",
        inputSchema={
            "type": "object",
            "properties": {
                "email_id": {"type": "string", "description": "IMAP message id"},
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
            },
            "required": ["email_id"],
        },
    ),
    Tool(
        name="email.get_attachment",
        description=("Contenuto di un allegato di un messaggio, in base64 (per nome file). "
                     "SOLO per allegati piccoli/testuali: per PDF, immagini e binari usa "
                     "email.save_attachment (il base64 non passa dal modello). "
                     "Usa email.read per scoprire i nomi degli allegati."),
        inputSchema={
            "type": "object",
            "properties": {
                "email_id": {"type": "string", "description": "IMAP message id"},
                "filename": {"type": "string", "description": "nome esatto dell'allegato (da email.read)"},
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
            },
            "required": ["email_id", "filename"],
        },
    ),
    Tool(
        name="email.save_attachment",
        description=("Scarica un allegato senza che i byte passino dal contesto del "
                     "modello — usa QUESTO per PDF, immagini e binari. Due "
                     "destinazioni, ne serve una: `tier`+`name` per archiviarlo "
                     "DIRETTAMENTE nei file di un topic (il gateway lo scrive, "
                     "provenienza `untrusted` e il canale risulta contaminato: non "
                     "serve topic.put), oppure `dest` per un path nel tuo scratch se "
                     "devi lavorarlo. Usa email.read per scoprire i nomi degli "
                     "allegati."),
        inputSchema={
            "type": "object",
            "properties": {
                "email_id": {"type": "string", "description": "IMAP message id"},
                "filename": {"type": "string", "description": "nome esatto dell'allegato (da email.read)"},
                "dest": {"type": "string", "description": ("path assoluto nel tuo "
                          "scratch. Alternativa a tier+name: se devi solo archiviare "
                          "l'allegato, usa tier+name e i byte non passano da te")},
                "tier": {"type": "string", "description": ("tier del topic in cui archiviare. "
                          "OPZIONALE: senza `dest` e senza tier/name l'allegato va nel "
                          "topic del canale in cui stai lavorando")},
                "name": {"type": "string", "description": "nome del topic (opzionale, vedi tier)"},
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
            },
            # `dest` non è più obbligatorio: serve UNA delle due destinazioni, e il
            # dispatch lo verifica con un errore che dice quali sono. Uno schema che
            # imponesse entrambe costringerebbe a passare un path finto per
            # archiviare in un topic.
            "required": ["email_id", "filename"],
        },
    ),
    Tool(
        name="email.search",
        description="Search messages with an IMAP query (e.g. FROM \"x@y.it\") in a folder.",
        inputSchema={
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "IMAP search query"},
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
                "limit": {"type": "integer", "description": "max results, default 20"},
            },
            "required": ["query"],
        },
    ),
    Tool(
        name="email.reply",
        description="Reply to a message keeping the thread (plain-text body, optional CC and local attachments).",
        inputSchema={
            "type": "object",
            "properties": {
                "email_id": {"type": "string", "description": "id of the message to reply to"},
                "body": {"type": "string"},
                "account": {"type": "string"},
                "folder": {"type": "string", "description": "IMAP folder, default INBOX"},
                "cc": {"type": "string", "description": "optional CC address"},
                "attachments": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "optional local file paths to attach",
                },
            },
            "required": ["email_id", "body"],
        },
    ),
]


# agent.spawn RIMOSSO (29 giu 2026): creava chat via POST /clodia/chats, endpoint
# eliminato nel passaggio al modello a canali/topic → il tool falliva a runtime
# (tool "fantasma" che illudeva l'agent di poter delegare in background). I
# subagent reali sono in-process (Task tool del Claude SDK), che girano dentro il
# turno (osservabili) e il cui esito rientra nel turno.
#
# agents.* (30 giu 2026): amministrazione delle capability degli ALTRI agent —
# dotare un agent editabile di skill/tool/rules dalla chat. Immutabili (super +
# flag immutable, es. Wainston) non toccabili. Scritture verificate anche dal
# backend (token inoltrato). Vedi tools/agents_admin.py.
_AGENT_TOOLS: list[Tool] = [
    Tool(name="agents.list",
         description=("Elenca gli agent dell'istanza con tipo e flag `immutable`. "
                      "Gli immutabili (super + protetti come Wainston) non sono "
                      "modificabili: si cambiano solo via codice/rebuild."),
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="agents.show",
         description="Capability correnti di un agent: skill (capabilities), rules, tool_permissions, immutabilità.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string", "description": "nome dell'agent"}},
             "required": ["agent"]}),
    Tool(name="agents.list_skills",
         description="Nomi delle skill disponibili nel catalogo (assegnabili come capabilities).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="agents.list_rules",
         description="Nomi delle rule disponibili nel catalogo (assegnabili).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="agents.list_tools",
         description="Namespace dei tool nativi del gateway concedibili a un agent (es. fs, email, topic, gdrive).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="agents.grant_skill",
         description="Aggiunge una skill (capability) a un agent editabile.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "skill": {"type": "string"}},
             "required": ["agent", "skill"]}),
    Tool(name="agents.revoke_skill",
         description="Rimuove una skill da un agent editabile.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "skill": {"type": "string"}},
             "required": ["agent", "skill"]}),
    Tool(name="agents.grant_tool",
         description=("Concede un permesso tool a un agent editabile. Può essere un "
                      "tool puntuale (es. `email.send`) o un namespace (`fs.*`)."),
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "tool": {"type": "string"}},
             "required": ["agent", "tool"]}),
    Tool(name="agents.revoke_tool",
         description="Revoca un permesso tool a un agent editabile.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "tool": {"type": "string"}},
             "required": ["agent", "tool"]}),
    Tool(name="agents.grant_rule",
         description="Aggiunge una rule (regola di stile/comportamento) a un agent editabile.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "rule": {"type": "string"}},
             "required": ["agent", "rule"]}),
    Tool(name="agents.revoke_rule",
         description="Rimuove una rule da un agent editabile.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "rule": {"type": "string"}},
             "required": ["agent", "rule"]}),
    Tool(name="agents.grant_scoped",
         description="Concede temporaneamente skill, regole, tool o runtime a un agent nello scope indicato.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"},
             "scope_kind": {"type": "string", "enum": ["topic", "chat", "run"], "default": "topic"},
             "scope_id": {"type": "string"},
             "ttl_minutes": {"type": "integer", "minimum": 1, "maximum": 120, "default": 15},
             "capabilities": {"type": "array", "items": {"type": "string"}},
             "rules": {"type": "array", "items": {"type": "string"}},
             "tools": {"type": "array", "items": {"type": "string"}},
             "model": {"type": "string"}, "provider": {"type": "string"},
             "reason": {"type": "string"}},
             "required": ["agent"]}),
    Tool(name="agents.list_scoped",
         description="Elenca gli override runtime attivi di un agent.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}}, "required": ["agent"]}),
    Tool(name="agents.revoke_scoped",
         description="Revoca immediatamente un override runtime scoped.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "override_id": {"type": "string"}},
             "required": ["agent", "override_id"]}),
]


_FS_TOOLS: list[Tool] = [
    Tool(
        name="fs.list_dir",
        description=(
            "List the entries of a directory inside the agent's workspace whitelist. "
            "Returns name, kind (file/dir), and size for each child."
        ),
        inputSchema={
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path relative to workspace root or absolute, must be in allowed_paths.",
                }
            },
            "required": ["path"],
        },
    ),
]

_WEB_TOOLS: list[Tool] = [
    Tool(
        name="web.fetch",
        description=(
            "Legge una pagina, un feed o una risposta JSON dal web via HTTP GET. "
            "Non richiede approvazione (leggere non è un'uscita), ma la FONTE conta: "
            "un URL dichiarato fidato in `ingress` non contamina il canale, uno "
            "sconosciuto sì — e da un canale contaminato ogni uscita (email, "
            "Telegram, POST) chiederà conferma umana. Non segue i redirect: li "
            "riporta, perché la destinazione va vagliata per conto suo. Solo "
            "destinazioni pubbliche (niente rete interna), solo contenuto testuale, "
            "corpo tenuto fino a 64 KB (alzabile con `max_bytes`, tetto 512 KB): un risultato grosso resta nel contesto e viene riletto a ogni azione successiva del turno."
        ),
        inputSchema={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL http/https da leggere"},
                "headers": {
                    "type": "object",
                    "additionalProperties": {"type": "string"},
                    "description": ("header opzionali; Host, Cookie, Authorization e "
                                    "hop-by-hop sono vietati"),
                },
                "timeout_seconds": {
                    "type": "number", "minimum": 0.1, "maximum": 30,
                    "description": "timeout, massimo 30 secondi",
                },
                "max_bytes": {
                    "type": "integer", "minimum": 1, "maximum": 524288,
                    "description": ("byte di corpo da tenere; default 65536. Il "
                                    "risultato resta nel contesto e viene riletto a "
                                    "ogni azione successiva del turno, quindi chiedi "
                                    "di più solo se il default ha tagliato ciò che "
                                    "serve — la risposta lo dice."),
                },
            },
            "required": ["url"],
        },
    ),
    Tool(
        name="web.post",
        description=(
            "Invia una richiesta HTTP POST dopo approvazione umana per questa "
            "singola invocazione. Non segue redirect; payload e risposta sono limitati."
        ),
        inputSchema={
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL http/https di destinazione"},
                "json": {"description": "payload JSON (alternativo a body)"},
                "body": {"type": "string", "description": "payload testuale (alternativo a json)"},
                "headers": {
                    "type": "object",
                    "additionalProperties": {"type": "string"},
                    "description": "header opzionali; Host e hop-by-hop sono vietati",
                },
                "timeout_seconds": {
                    "type": "number", "minimum": 0.1, "maximum": 10,
                    "description": "timeout, massimo 10 secondi",
                },
            },
            "required": ["url"],
        },
    ),
]

_LOGS_TOOLS: list[Tool] = [
    Tool(
        name="logs.tail",
        description=(
            "Read-only: le ultime righe del log del server (agent-server) per la "
            "diagnosi. Segreti redatti. Solo log di piattaforma, MAI contenuti dei topic."
        ),
        inputSchema={
            "type": "object",
            "properties": {
                "lines": {"type": "integer", "description": "Numero di righe (default 100, max 500)."},
                "level": {"type": "string", "description": "Filtro livello opzionale: INFO|WARNING|ERROR."},
            },
        },
    ),
]

_EU_CORPUS_TOOLS: list[Tool] = [
    Tool(
        name="eu_corpus.search",
        description=(
            "Retrieval semantico sul corpus normativo UE stabile (Horizon Europe): "
            "AGA (Annotated Grant Agreement), Programme Guide, General Annexes. "
            "Query in linguaggio naturale, IT o EN (l'embedding è multilingue). "
            "Ritorna i passaggi più pertinenti con CITAZIONE (documento, versione, "
            "sezione, pagina, score). Usalo per domande su eleggibilità costi, "
            "categorie di budget, funding rate, regole TRL, FSTP/cascade. "
            "IMPORTANTE: il retrieval trova i candidati, non è la verità — leggi il "
            "testo del passaggio per intero e cita sempre documento+versione+pagina."
        ),
        inputSchema={"type": "object", "properties": {
            "query": {"type": "string", "description": "domanda in linguaggio naturale (IT/EN)"},
            "k": {"type": "integer", "description": "n. passaggi (1-20, default 5)"},
            "doc": {"type": "string", "description": "filtro opzionale per documento: "
                    "AGA | HE-Programme-Guide | HE-General-Annexes"},
        }, "required": ["query"]},
    ),
    Tool(
        name="eu_corpus.ingest",
        description=(
            "Aggiunge un documento PDF alla knowledge base normativa (corpus RAG). "
            "Il file deve già stare nei files/ di un topic di cui sei participant "
            "(es. un PDF che l'utente ha allegato in chat → path 'files/xxx.pdf'). "
            "Lo estrae, chunka, embedda e lo indicizza su pgvector. "
            "USALO per materiale NORMATIVO/DI RIFERIMENTO stabile (guide, regolamenti, "
            "grant agreement), NON per dossier confidenziali specifici di un cliente "
            "(quelli restano nel topic e si leggono live con topic.read_file). "
            "doc_name+version identificano il documento: se stai caricando una NUOVA "
            "versione di un documento già presente, passa supersede=true (le versioni "
            "precedenti restano ma vengono marcate superseded). Ri-ingerire la stessa "
            "(doc_name, version) è idempotente (rimpiazza i chunk)."
        ),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"],
                     "description": "tier del topic da cui leggere il file"},
            "name": {"type": "string", "description": "nome del topic da cui leggere il file"},
            "path": {"type": "string", "description": "path del PDF nel topic, es. 'files/aga.pdf'"},
            "doc_name": {"type": "string", "description": "nome del documento nel corpus (es. 'AGA', 'HE-Programme-Guide')"},
            "version": {"type": "string", "description": "versione, es. '2.0 (2025-04-01)'"},
            "url": {"type": "string", "description": "URL fonte ufficiale (opzionale ma consigliato)"},
            "supersede": {"type": "boolean", "description": "true se è una nuova versione di un doc esistente"},
        }, "required": ["tier", "name", "path", "doc_name", "version"]},
    ),
    Tool(
        name="eu_corpus.list",
        description=("Elenca i documenti indicizzati nella knowledge base (corpus RAG): "
                     "nome, versione, status (active/superseded), n. chunk, fonte. "
                     "Usalo per mostrare all'utente cosa c'è nel corpus."),
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="eu_corpus.remove",
        description=("Rimuove un documento dalla knowledge base (corpus RAG). "
                     "DISTRUTTIVO: cancella il documento e tutti i suoi chunk. "
                     "Se ometti `version` rimuove TUTTE le versioni di quel documento. "
                     "CONFERMA sempre con l'utente cosa stai per rimuovere prima di chiamarlo "
                     "(nome + versione). Usa eu_corpus.list per i nomi/versioni esatti."),
        inputSchema={"type": "object", "properties": {
            "doc_name": {"type": "string", "description": "nome del documento (come da eu_corpus.list)"},
            "version": {"type": "string", "description": "versione specifica; se omessa, tutte le versioni"},
        }, "required": ["doc_name"]},
    ),
]


_RAG_TOOLS: list[Tool] = [
    Tool(
        name="rag.collections",
        description=("Elenca le knowledge base (collection) su cui hai accesso in "
                     "lettura, con tier e conteggi. Usalo per sapere quali corpora "
                     "puoi interrogare."),
        inputSchema={"type": "object", "properties": {}},
    ),
    Tool(
        name="rag.create_collection",
        description=("Crea/provisiona una collection della knowledge base. Riservato "
                     "al provisioning dei pack: usa solo nomi e tier dichiarati nei "
                     "manifest curated. Dopo la creazione usa rag.ingest per le "
                     "risorse iniziali."),
        inputSchema={"type": "object", "properties": {
            "collection": {"type": "string"},
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "description": {"type": "string"},
        }, "required": ["collection"]},
    ),
    Tool(
        name="rag.search",
        description=("Retrieval semantico su una collection della knowledge base. "
                     "Query IT/EN. Ritorna passaggi con citazione (documento, "
                     "versione, sezione, pagina, score). Il retrieval trova i "
                     "candidati, non è la verità: leggi il passaggio per intero e "
                     "cita sempre documento+versione+pagina."),
        inputSchema={"type": "object", "properties": {
            "collection": {"type": "string", "description": "collection su cui cercare"},
            "query": {"type": "string", "description": "domanda in linguaggio naturale (IT/EN)"},
            "k": {"type": "integer", "description": "n. passaggi (1-20, default 5)"},
            "doc": {"type": "string", "description": "filtro opzionale per nome documento"},
        }, "required": ["collection", "query"]},
    ),
    Tool(
        name="rag.list",
        description="Elenca i documenti di una collection (nome, versione, status, n. chunk, fonte).",
        inputSchema={"type": "object", "properties": {
            "collection": {"type": "string"},
        }, "required": ["collection"]},
    ),
    Tool(
        name="rag.ingest",
        description=("Aggiunge un PDF (già nei files/ di un topic di cui sei "
                     "participant) a una collection della knowledge base. Il gateway "
                     "legge i byte server-side, li estrae/chunka/embedda/indicizza. "
                     "Richiede grant di SCRITTURA sulla collection. Solo materiale "
                     "stabile/di riferimento, non dossier confidenziali per-cliente. "
                     "supersede=true per una nuova versione di un doc già presente."),
        inputSchema={"type": "object", "properties": {
            "collection": {"type": "string"},
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"],
                     "description": "tier del topic da cui leggere il file"},
            "name": {"type": "string", "description": "nome del topic da cui leggere il file"},
            "path": {"type": "string", "description": "path del PDF nel topic, es. 'files/x.pdf'"},
            "doc_name": {"type": "string", "description": "nome del documento nella collection"},
            "version": {"type": "string"},
            "url": {"type": "string"},
            "supersede": {"type": "boolean"},
        }, "required": ["collection", "tier", "name", "path", "doc_name", "version"]},
    ),
    Tool(
        name="rag.remove",
        description=("Rimuove un documento da una collection (DISTRUTTIVO). Richiede "
                     "grant di SCRITTURA. Se ometti version rimuovi tutte le versioni. "
                     "Conferma sempre con l'utente cosa stai per rimuovere."),
        inputSchema={"type": "object", "properties": {
            "collection": {"type": "string"},
            "doc_name": {"type": "string"},
            "version": {"type": "string"},
        }, "required": ["collection", "doc_name"]},
    ),
]


_GITHUB_TOOLS: list[Tool] = [
    # github.*: le azioni git che ESCONO dallo scope (§5.2). `add`, `diff` e
    # `commit` restano nel container dell'agente e non passano di qui.
    Tool(
        name="github.clone",
        description=("Clona un repository APPROVATO per questo topic nella tua "
                     "scratch. La credenziale la fornisce l'owner al mount e non "
                     "entra mai nel tuo processo."),
        inputSchema={"type": "object", "properties": {
            "repo": {"type": "string", "description": "https://github.com/<owner>/<repo>"},
            "dest": {"type": "string", "description": "cartella di destinazione nella tua scratch"},
            "branch": {"type": "string"},
        }, "required": ["repo", "dest"]},
    ),
    Tool(
        name="github.pull",
        description="Aggiorna (fast-forward) un working tree clonato con github.clone.",
        inputSchema={"type": "object", "properties": {
            "dir": {"type": "string"}}, "required": ["dir"]},
    ),
    Tool(
        name="github.push",
        description=("Manda al repository i commit già fatti. NON committa: "
                     "`git add`/`git commit` li fai tu nella scratch."),
        inputSchema={"type": "object", "properties": {
            "dir": {"type": "string"}, "branch": {"type": "string"}},
            "required": ["dir"]},
    ),
    Tool(
        name="github.pull_request",
        description="Apre una pull request sul repository approvato.",
        inputSchema={"type": "object", "properties": {
            "repo": {"type": "string"}, "head": {"type": "string"},
            "base": {"type": "string", "description": "default: main"},
            "title": {"type": "string"}, "body": {"type": "string"}},
            "required": ["repo", "head", "title"]},
    ),
]

# I verbi `github.*` che il gateway implementa DA SÉ. Serve perché il namespace
# `github.` è conteso: il backend MCP ufficiale di GitHub è montato con lo stesso
# nome, quindi i suoi tool si chiamano anch'essi `github.<qualcosa>` (issue_write,
# list_issues, …). Il dispatch sceglieva per PREFISSO e mandava tutto al ramo
# nativo, che rispondeva «unknown github verb» a un verbo che il backend esponeva
# davvero e che la whitelist dell'agente permetteva: il tool compariva nella lista
# e falliva alla chiamata.
# Derivata dai Tool dichiarati e non riscritta a mano: due elenchi dello stesso
# insieme divergono al primo verbo aggiunto, e questo qui deciderebbe in silenzio
# quale dei due namespace vince.
_GITHUB_NATIVE_NAMES: frozenset[str] = frozenset(t.name for t in _GITHUB_TOOLS)


_TOPIC_TOOLS: list[Tool] = [
    Tool(
        name="topic.new",
        description=("Crea (idempotente) un topic. tier = SEAL-0..4 "
                     "(Public/Internal/Confidential/Restricted/Sovereign; default SEAL-0). "
                     "meta opzionale: title, type, tags, people, entity, deadline, contact_agent."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"],
                     "description": "classe/sovranità del topic — scala SEAL (default SEAL-0 Public)"},
            "name": {"type": "string", "description": "slug a-z0-9_-"},
            "meta": {"type": "object"},
            "hook_enabled": {
                "type": "boolean",
                "description": "crea il webhook del topic (default true)"},
        }, "required": ["name"]},
    ),
    Tool(
        name="topic.invoke_hook",
        description=(
            "Invoca localmente l'hook di un topic senza segreto. Il chiamante deve "
            "essere participant; messaggero può invocare qualunque topic. Il payload "
            "viene accodato come @caller e sveglia il caller."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "payload": {"type": "string"},
        }, "required": ["tier", "name", "payload"]},
    ),
    Tool(
        name="topic.open",
        description=("Apre un topic (read-only): ritorna meta, summary (col "
                     "summary_version per l'optimistic lock), tldr, lista minute."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.save_summary",
        description=("Riscrive il summary in optimistic lock. Passa base_version "
                     "ottenuto da topic.open. Su conflitto NON sovrascrive: rilegge ed escala."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "text": {"type": "string", "description": "prima riga = TLDR; sezione '## Prossimi passi'"},
            "base_version": {"type": ["string", "null"]},
        }, "required": ["tier", "name", "text"]},
    ),
    Tool(
        name="topic.save_agents_md",
        description=("Riscrive le ISTRUZIONI di scope del topic (AGENTS.md) in optimistic "
                     "lock, come il summary. Non è un file del topic: vive nel "
                     "control-plane, non è sincronizzato dai remote e non si carica con "
                     "topic.put. Testo vuoto = rimuove le istruzioni."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "text": {"type": "string", "description": "markdown; vuoto = rimuove"},
            "base_version": {"type": ["string", "null"],
                             "description": "agents_md_version letto con topic.open"},
        }, "required": ["tier", "name", "text"]},
    ),
    Tool(
        name="topic.add_minute",
        description="Aggiunge una minuta (file append-only datato). Niente contesa concorrente.",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "text": {"type": "string"},
        }, "required": ["tier", "name", "text"]},
    ),
    Tool(
        name="topic.archive",
        description="Imposta status=archived nel meta (non sposta su storage inferiore).",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.telegram_bind",
        description=("Collega un GRUPPO Telegram a questo topic: le menzioni "
                     "delle persone mappate vengono riportate lì, col link alla "
                     "conversazione. `people` = {uid_telegram: nome_utente_clodia}. "
                     "Atto sui muri dello scope: lo decide l'owner."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "chat_id": {"type": "string", "description": "id del gruppo (negativo per i supergruppi)"},
            "mode": {"type": "string", "enum": ["notify", "excerpt"],
                     "description": "notify = solo il fatto · excerpt = anche la riga della menzione"},
            "people": {"type": "object", "description": "{uid telegram: nome utente su Clodia}"},
            "mount": {"type": "string", "description": "nome del mount (default: telegram)"},
        }, "required": ["tier", "name", "chat_id", "people"]},
    ),
    Tool(
        name="topic.telegram_unbind",
        description="Scollega il gruppo Telegram. La voce di egress resta: toglierla è una decisione a parte.",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "mount": {"type": "string"}}, "required": ["tier", "name"]},
    ),
    Tool(
        name="telegram.roster",
        description=("Amministratori e membri noti di un gruppo, con uid e "
                     "username: serve a mappare le persone senza digitare numeri."),
        inputSchema={"type": "object", "properties": {
            "chat_id": {"type": "string"}}, "required": ["chat_id"]},
    ),
    Tool(
        name="telegram.notify_pending",
        description="Le notifiche di menzione da recapitare (testo già composto).",
        inputSchema={"type": "object", "properties": {
            "limit": {"type": "integer"}}},
    ),
    Tool(
        name="telegram.notify_flush",
        description=("Recapita le notifiche di menzione pendenti sul gruppo "
                     "collegato. Meccanico: il testo è già composto, non c'è "
                     "nulla da decidere. Pensato per un job logico."),
        inputSchema={"type": "object", "properties": {
            "limit": {"type": "integer"}}},
    ),
    Tool(
        name="telegram.notify_ack",
        description="Segna una notifica come recapitata, o come fallita con il motivo.",
        inputSchema={"type": "object", "properties": {
            "message_id": {"type": "string"}, "chat_id": {"type": "string"},
            "principal": {"type": "string"}, "ok": {"type": "boolean"},
            "error": {"type": "string"}},
            "required": ["message_id", "chat_id", "principal"]},
    ),
    Tool(
        name="topic.set_portable",
        description=("Dichiara (o revoca) la PORTABILITÀ di un topic: i suoi "
                     "partecipanti possono leggerne i contenuti anche da altre "
                     "stanze, entro il tier della stanza in cui si trovano. "
                     "Atto sui muri dello scope: lo decide l'owner."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "portable": {"type": "boolean", "description": "true = portabile, false = revoca"},
        }, "required": ["tier", "name", "portable"]},
    ),
    Tool(
        name="topic.list",
        description="Elenca i topic (riga sintetica). Gli archived sono nascosti salvo include_archived.",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "include_archived": {"type": "boolean"},
        }},
    ),
    Tool(
        name="topic.search",
        description="Ricerca nei topic (P1: lessicale su meta/summary/minute).",
        inputSchema={"type": "object", "properties": {
            "query": {"type": "string"},
            "mode": {"type": "string", "enum": ["lexical", "semantic"]},
        }, "required": ["query"]},
    ),
    Tool(
        name="topic.files",
        description=("Elenca file e cartelle del topic/canale a partire da `subpath` "
                     "(relativo alla ROOT del topic; vuoto = root). Ritorna name, path, "
                     "kind (file|dir), size, mtime. I file caricati stanno di norma in "
                     "'files/'. Per vedere il CONTENUTO di una sottocartella passa il suo "
                     "path come subpath (es. subpath='files/expenses'). Usa il `path` "
                     "ritornato con topic.read_file."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "subpath": {"type": "string", "description": "cartella da elencare, relativa alla root del topic (es. 'files' o 'files/expenses'); vuoto = root"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.read_file",
        description=("Legge il contenuto di un file del topic/canale. path relativo "
                     "(es. 'files/report.md'). I file di testo tornano come testo; "
                     "i binari (PDF/immagini) tornano come base64 con encoding='base64'."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "path": {"type": "string", "description": "path relativo al topic, es. files/foo.md"},
        }, "required": ["tier", "name", "path"]},
    ),
    Tool(
        name="topic.read_document",
        description=("Estrae il TESTO di un documento del topic (PDF, DOCX, XLSX) — "
                     "l'estrazione avviene server-side nel gateway, quindi ricevi TESTO "
                     "leggibile, non base64. USA QUESTO per leggere un PDF/DOCX/XLSX "
                     "invece di topic.read_file (che restituisce base64 binario). "
                     "Ritorna {text, chars, pages, truncated}. Per PDF lunghi usa "
                     "max_chars per limitare il testo."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "path": {"type": "string", "description": "path relativo al topic, es. files/report.pdf"},
            "max_chars": {"type": "integer", "description": "max caratteri restituiti (default 60000)"},
        }, "required": ["tier", "name", "path"]},
    ),
    Tool(
        name="topic.write_file",
        description=("Carica/sovrascrive un file nella cartella files/ del topic/canale "
                     "(es. un deliverable, o i file estratti da uno zip). filename può "
                     "includere sottocartelle (es. 'archivio/foto/1.jpg'); le dir padre "
                     "vengono create. content = testo; per i binari (xlsx/pdf/docx/"
                     "zip/immagini) passa il base64 COMPLETO del file e encoding='base64' "
                     "(i file con estensione binaria sono comunque decodificati da base64, "
                     "mai scritti come testo). Usa QUESTO, non hosting esterni."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "filename": {"type": "string", "description": "nome file semplice, finisce in files/"},
            "content": {"type": "string", "description": "contenuto (testo o base64)"},
            "encoding": {"type": "string", "enum": ["text", "base64"], "description": "default text"},
        }, "required": ["tier", "name", "filename", "content"]},
    ),
    Tool(
        name="artifact.render",
        description=("Aggiorna il CANVAS LIVE del topic con un artefatto HTML. Scrive lo "
                     "snapshot in files/artifact.html (persistente e riapribile) e la "
                     "finestra di anteprima del topic lo mostra aggiornato. Passa l'INTERO "
                     "documento HTML in `html` a OGNI chiamata (snapshot completo, non un "
                     "frammento/patch). Usalo per mostrare all'utente un artefatto vivo "
                     "(cover, mockup, dashboard) che evolvi durante la conversazione."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "html": {"type": "string", "description": "documento HTML COMPLETO (snapshot del canvas)"},
        }, "required": ["tier", "name", "html"]},
    ),
    Tool(
        name="topic.fetch",
        description=("Scarica una COPIA di un file del topic nel TUO scratch (path "
                     "locale), per trattarlo con le skill standard (xlsx/pdf/docx/…). "
                     "USA QUESTO per i BINARI invece di topic.read_file (che passa "
                     "base64 nel contesto e si tronca sui file grandi). Il trasporto "
                     "usa envelope cifrati effimeri sul volume shared. `dest` è "
                     "OPZIONALE: senza, il file prende il proprio nome nel tuo "
                     "scratch, e il path locale te lo ritorno io in `local_path` — "
                     "non provare a comporlo da `pwd`, che è la radice dello spawn "
                     "e NON lo scratch. Flusso: topic.fetch → skill "
                     "standard su `local_path` → topic.put."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "path": {"type": "string", "description": "path nel topic, es. files/expenses/x.xlsx"},
            "dest": {"type": "string", "description": "opzionale: nome del file nel tuo scratch (default: lo stesso nome che ha nel topic)"},
        }, "required": ["tier", "name", "path"]},
    ),
    Tool(
        name="topic.put",
        description=("Carica nel topic (files/) un file preparato nel TUO scratch. USA "
                     "QUESTO per i BINARI invece di topic.write_file: il gateway legge i "
                     "byte tramite un envelope cifrato effimero, niente base64 nel modello. `src` = path "
                     "assoluto nel tuo scratch; `filename` può includere sottocartelle."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "filename": {"type": "string", "description": "nome file (può includere sottocartelle)"},
            "src": {"type": "string", "description": "path assoluto del file nel tuo scratch"},
        }, "required": ["tier", "name", "filename", "src"]},
    ),
    Tool(
        name="topic.post_message",
        description=("Posta un MESSAGGIO nella chat del topic (una bolla nella "
                     "conversazione), come te. Serve a far comparire nel topic ciò che "
                     "arriva da fuori (es. una mail in arrivo) o un handoff a fine job. "
                     "Se includi una "
                     "@menzione (es. `@messaggero …`), innesca l'agente taggato che "
                     "prende in carico il messaggio; senza menzione è solo una bolla. "
                     "Puoi postare solo in un topic di cui sei participant (cross-topic → gate)."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "text": {"type": "string", "description": "testo del messaggio"},
        }, "required": ["tier", "name", "text"]},
    ),
    Tool(
        name="topic.messages",
        description=("Legge la CONVERSAZIONE del topic (le bolle della chat), dalla più "
                     "vecchia alla più recente. Un agente in turno la riceve già nel "
                     "contesto: questo verbo serve a chi guarda la stanza da fuori — un "
                     "client MCP di una persona, o un agente che riprende il filo."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "limit": {"type": "integer", "description": "quanti messaggi finali (default 200)"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.my_mentions",
        description=("Le menzioni rivolte A TE in questo topic che non hai ancora "
                     "marcato come viste. La domanda è sempre su chi chiama: non si "
                     "può leggere la casella di un altro. Da MCP non arriva nessun "
                     "push — questo è il verbo con cui si CHIEDE se qualcuno ti ha "
                     "chiamato."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "limit": {"type": "integer"},
            "only_unseen": {"type": "boolean", "description": "default true"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.mark_seen",
        description=("Sposta il tuo segnaposto delle menzioni. Passa il `seen_through` "
                     "ritornato da topic.my_mentions, non l'istante corrente: fra le due "
                     "chiamate può essere arrivata una menzione che non hai visto."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "seen_through": {"type": "string", "description": "ts da topic.my_mentions"},
        }, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.delete_file",
        description=("Sposta nel CESTINO (.trash/) un file o una cartella DENTRO files/ del "
                     "topic — NON cancella mai davvero: è sempre recuperabile. Solo sotto "
                     "files/ (meta, summary, minutes sono protetti). path = path relativo alla "
                     "root del topic, come da topic.files (es. 'files/old/x.pdf' o 'files/files')."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "path": {"type": "string", "description": "path da eliminare, dentro files/"},
        }, "required": ["tier", "name", "path"]},
    ),
    Tool(
        name="topic.migrate_storage",
        description=("Migra i FILE del topic da locale a Google Drive o viceversa. "
                     "Su Drive la cartella remota diventa il filesystem live autoritativo; "
                     "tornando a locale, i file remoti vengono materializzati nel topic. "
                     "Guard SEAL: vietato migrare su uno storage con livello inferiore al tier "
                     "(es. SEAL-3 non va su Drive). target.type=local|drive; per drive folder "
                     "(link/id) opzionale (vuoto = crea cartella)."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "target": {"type": "object", "properties": {
                "type": {"type": "string", "enum": ["local", "drive"]},
                "folder": {"type": "string"}, "account": {"type": "string"}},
                "required": ["type"]},
        }, "required": ["tier", "name", "target"]},
    ),
    # ── Remote pluggable: git usa il ciclo di sync; Drive è un filesystem live. ──
    Tool(
        name="topic.remote_enable",
        description=("Attiva un remote per i FILE del topic. Git conserva il ciclo "
                     "add/commit/push/pull. Drive diventa il filesystem autoritativo live: "
                     "config.folder link/id opzionale (vuoto = crea cartella), config.account."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "type": {"type": "string", "enum": ["git", "drive"]},
            "config": {"type": "object", "description": "git: {url,branch} · drive: {folder,account}"},
            "mount": {"type": "string", "description": "nome del mount (default: il tipo). Uno scope può averne più d'uno."},
        }, "required": ["tier", "name", "type"]},
    ),
    Tool(
        name="topic.remote_disable",
        description=("Disattiva il remote preservando i file. Per Drive materializza "
                     "prima la cartella remota nel filesystem locale."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "mount": {"type": "string", "description": "quale mount staccare (default: l'unico)"}},
            "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.remote_add",
        description=("Marca un file per il sync Git. Su Drive è un no-op deprecato "
                     "perché topic.write_file/topic.put caricano immediatamente."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"}, "path": {"type": "string"}}, "required": ["tier", "name", "path"]},
    ),
    Tool(
        name="topic.remote_commit",
        description="Snapshot delle modifiche Git. Su Drive live è un no-op deprecato.",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"}, "message": {"type": "string"}}, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.remote_push",
        description="Invia le modifiche Git. Su Drive live è un no-op deprecato.",
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"}}, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.remote_pull",
        description=("Riceve dal remote Git (conflitto→escala). Su Drive live è un "
                     "no-op deprecato perché le letture vedono già il remoto."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"}}, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.remote_status",
        description=("Stato del remote del topic. Per Drive include mode=live e "
                     "last_write_wins=true."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"}}, "required": ["tier", "name"]},
    ),
    Tool(
        name="topic.suggest_team",
        description=("Proponi la SQUADRA di agenti per un topic, data una breve "
                     "descrizione di cosa tratta. Ritorna gli agenti più "
                     "specializzati (rilevanza) e meno costosi fra quelli idonei "
                     "al tier (SEAL/clearance/provider): `candidates` ordinati con "
                     "score+costo+expertise, `suggested` (specialisti proposti) e "
                     "`coordinator` (super-agent, opzionale). Read-only: NON invita "
                     "nessuno (l'invito lo conferma l'owner). Usalo quando l'owner "
                     "descrive un nuovo topic per proporgli chi coinvolgere."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"],
                     "description": "tier del topic (default SEAL-0)"},
            "description": {"type": "string",
                            "description": "di cosa tratta il topic, in linguaggio naturale"},
        }, "required": ["description"]},
    ),
    Tool(
        name="topic.add_participant",
        description=("Aggiunge un agente ai partecipanti di un topic/chat esistente "
                     "(lo 'invita nella stanza'). Puoi usarlo se sei owner, "
                     "partecipante o super-agent del topic. Decidi TU chi coinvolgere "
                     "leggendo runtime.agents (expertise/skill/clearance/costo); "
                     "l'idoneità SEAL è comunque applicata alla risposta."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string", "description": "slug del topic"},
            "agent": {"type": "string", "description": "nome dell'agent/utente da aggiungere"},
        }, "required": ["tier", "name", "agent"]},
    ),
    Tool(
        name="topic.remove_participant",
        description=("Rimuove un agente dai partecipanti di un topic/chat. Come "
                     "add_participant: owner|partecipante|super."),
        inputSchema={"type": "object", "properties": {
            "tier": {"type": "string", "enum": ["SEAL-0", "SEAL-1", "SEAL-2", "SEAL-3", "SEAL-4"]},
            "name": {"type": "string"},
            "agent": {"type": "string"},
        }, "required": ["tier", "name", "agent"]},
    ),
]


_PROFILE_TOOLS: list[Tool] = [
    Tool(name="profile.get",
         description=("Dati personali (PII) di un agent/umano: email, iban, domicilio, ecc. "
                      "Ritorna i campi SOLO se sei il titolare, un admin, o hai ricevuto il grant "
                      "(ACL per-profilo). Usalo quando ti serve un dato personale di qualcuno."),
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string", "description": "nome dell'agent/umano di cui leggere il profilo"}},
             "required": ["agent"]}),
    Tool(name="profile.set",
         description="Crea/aggiorna i campi del TUO profilo (o, se admin, di un altro). fields = oggetto chiave→valore; valore null rimuove la chiave.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "fields": {"type": "object"}},
             "required": ["fields"]}),
    Tool(name="profile.list_files",
         description="Elenca i file allegati al profilo di un agent (se autorizzato): name, size, mtime.",
         inputSchema={"type": "object", "properties": {"agent": {"type": "string"}}, "required": ["agent"]}),
    Tool(name="profile.read_file",
         description="Legge un file allegato al profilo (se autorizzato). Ritorna testo, o base64 per i binari.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "filename": {"type": "string"}}, "required": ["agent", "filename"]}),
    Tool(name="profile.grant",
         description="Concedi/revoca a un altro agent la lettura del TUO profilo (o, se admin, di un altro). granted=false per revocare.",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string"}, "grantee": {"type": "string"},
             "granted": {"type": "boolean"}},
             "required": ["grantee"]}),
]


_RUNTIME_TOOLS: list[Tool] = [
    Tool(name="runtime.agents",
         description="Introspezione runtime: gli agent dell'istanza col quadro COMPLETO per decidere chi coinvolgere (dominio/expertise, skill, knowledge RAG, clearance SEAL, provider effettivo + suo SEAL, modello, ruolo, stato). Solo metadati, mai segreti — la decisione è tua, non del tool.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.jobs",
         description="Introspezione runtime: i job schedulati (cron/intervallo) e il loro stato.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.skills",
         description="Introspezione runtime: le skill nel catalogo, per pack.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.chats",
         description="Introspezione runtime: le chat aperte (id/kind/titolo/stato, non il contenuto).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.topics",
         description="Introspezione runtime: i topic dell'istanza (metadati). I P3 (Restricted) sono esclusi salvo include_restricted=true.",
         inputSchema={"type": "object", "properties": {
             "include_restricted": {"type": "boolean", "description": "includi i topic P3 Restricted (default false)"}}}),
    Tool(name="runtime.mcp_servers",
         description="Introspezione runtime: i server MCP disponibili (backend montati via Add-MCP + namespace nativi).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.providers",
         description="Introspezione runtime: i provider di inferenza (id/nome/meccanismo/stato di connessione). MAI segreti.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.current_user",
         description="Chi è l'utente umano con cui stai parlando: l'owner/superadmin dell'istanza (+ gli altri principal umani). Usalo per sapere a chi ti stai rivolgendo.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="runtime.restart_agent",
         description="OPS: riavvia le sessioni vive di un agente — ferma i suoi subprocess/runtime; alla prossima interazione la chat rimaterializza il seed da zero. Da usare per sbloccare un agente col runtime impuntato (es. sessione opencode persa, loop). I dati/la history persistono. Capacità di ops (sysadmin).",
         inputSchema={"type": "object", "properties": {
             "agent": {"type": "string", "description": "nome del seed da riavviare (es. 'impiegato-tomato')"}},
             "required": ["agent"]}),
    Tool(name="runtime.inspect_topic",
         description=("OPS/STEWARD: ispeziona UN topic specifico (di norma quello da cui "
                      "l'utente ti sta chiamando nel widget). Ritorna metadati (titolo, "
                      "stato, owner), gli AGENTI del topic e gli ULTIMI messaggi. NB: "
                      "vincolato alla tua clearance — se la tua SEAL effettiva è < tier "
                      "del topic ricevi 403 (il topic è invisibile). Usalo per capire il "
                      "contesto in cui stai assistendo l'utente (chi c'è, cosa si sono detti)."),
         inputSchema={"type": "object", "properties": {
             "tier": {"type": "string", "description": "tier del topic (es. SEAL-1)"},
             "name": {"type": "string", "description": "nome del topic"}},
             "required": ["tier", "name"]}),
]

# jobs.* — gestione dei job schedulati. La CREAZIONE non è diretta: si PROPONE
# e l'owner approva via link firmato (un job è esecuzione autonoma ricorrente →
# superficie di privilegio, deve passare dall'umano).
_JOBS_TOOLS: list[Tool] = [
    Tool(name="jobs.list",
         description="Elenca i job schedulati dell'istanza (cron + stato). Sola lettura.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="jobs.propose",
         description=("PROPONE un nuovo job schedulato: NON lo crea. Registra una "
                      "proposta; il job nasce solo se l'owner approva. Il risultato "
                      "include `render_marker`: presenta il job all'utente e includi "
                      "quel marker in fondo al messaggio → comparirà un popup "
                      "Approva/Annulla in chat (conferma sincrona, l'owner è presente). "
                      "Usalo quando l'utente chiede di schedulare un'attività ricorrente "
                      "(report settimanale, promemoria, backup, ...). Fornisci una "
                      "descrizione della cadenza in linguaggio naturale (schedule_text) "
                      "oppure un cron a 5 campi (cron_expr)."),
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string", "description": "nome univoco del job"},
             "prompt": {"type": "string", "description": "cosa deve fare l'agente al fire del job"},
             "schedule_text": {"type": "string", "description": "cadenza in linguaggio naturale (es. 'ogni lunedì alle 9')"},
             "cron_expr": {"type": "string", "description": "in alternativa, cron a 5 campi"},
             "agent": {"type": "string", "description": "agent (kind) che esegue il job al fire. Default: l'agente chiamante (te stesso). Indicane un altro solo se il job deve girare per conto di qualcun altro."},
             "enabled": {"type": "boolean", "description": "attivo alla creazione (default true)"},
         }, "required": ["name", "prompt"]}),
    Tool(name="jobs.report_status",
         description=(
             "DICHIARA com'è andato il job schedulato che stai eseguendo. Chiamalo "
             "come ULTIMA cosa del turno, sempre, anche quando è andato tutto bene.\n\n"
             "  success — hai fatto il lavoro che il job chiede\n"
             "  error   — l'hai consegnato, ma qualcosa è andato storto e la QUALITÀ "
             "può esserne compromessa (una fonte su tre in 403, mezzo risultato, un "
             "allegato non recuperato)\n"
             "  fatal   — il lavoro NON è stato fatto: non hai potuto spedire, la "
             "fonte era irraggiungibile, il verbo che ti serviva è negato\n\n"
             "In `detail` scrivi cosa è andato storto, in una frase: è il testo che "
             "chi apre lo storico legge per capire se deve intervenire. Se non "
             "chiami questo verbo il run viene registrato `error`, perché un esito "
             "che nessuno dichiara non è un esito riuscito. Vale solo dentro un job "
             "schedulato: fuori non c'è un run da dichiarare."),
         inputSchema={"type": "object", "properties": {
             "status": {"type": "string", "enum": ["success", "error", "fatal"],
                        "description": "esito del run"},
             "detail": {"type": "string",
                        "description": ("cosa è andato storto, in una frase. "
                                        "Obbligatorio di fatto su error e fatal: "
                                        "senza, lo storico dice solo che qualcosa "
                                        "non ha funzionato")},
         }, "required": ["status"]}),
]

# packs.* — import/rimozione dei pack e loro dipendenze. Riservati a sysadmin.
_PACKS_TOOLS: list[Tool] = [
    Tool(name="packs.list",
         description="Elenca i pack installati (nome, versione, plugin/seed contenuti).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="packs.show",
         description="Dettaglio di un pack installato per nome.",
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}}, "required": ["name"]}),
    Tool(name="packs.import_url",
         description=("Importa un pack da URL (repo pubblico / zip remoto). L'import da "
                      "file .zip caricato resta un'operazione della UI (upload)."),
         inputSchema={"type": "object", "properties": {
             "url": {"type": "string"}}, "required": ["url"]}),
    Tool(name="packs.remove",
         description="Rimuove un pack installato per nome.",
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}}, "required": ["name"]}),
    Tool(name="packs.setup_done",
         description=("Marca il SETUP di un pack come COMPLETATO: smarca il flag "
                      "'setup_pending' (la UI toglie il bottone «Finish setup»). Chiamalo "
                      "SOLO al termine del task di setup del pack (deps installate, MCP "
                      "montati, RAG provisionato e verificato)."),
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}}, "required": ["name"]}),
    Tool(name="packs.install_pip",
         description=("Installa package pip dichiarati da un pack nel venv persistente "
                      "$CLODIA_DATA/runtime/venv. Non è shell libera: accetta solo spec "
                      "package/versione validate. Usa solo valori provenienti da "
                      "`requires.pip` del manifest curated."),
         inputSchema={"type": "object", "properties": {
             "packages": {"type": "array", "items": {"type": "string"}}},
             "required": ["packages"]}),
    Tool(name="packs.install_npm",
         description=("Installa package npm dichiarati da un pack nel prefix persistente "
                      "$CLODIA_DATA/runtime/npm. Non è shell libera: accetta solo spec "
                      "package/versione validate. Usa solo valori provenienti da "
                      "`requires.npm` del manifest curated."),
         inputSchema={"type": "object", "properties": {
             "packages": {"type": "array", "items": {"type": "string"}}},
             "required": ["packages"]}),
    Tool(name="packs.check_command",
         description=("Verifica che un binario richiesto dal pack sia disponibile nel "
                      "PATH runtime (venv/bin + npm/bin + PATH di sistema). Usalo per "
                      "`requires.bin` e `requires.system`."),
         inputSchema={"type": "object", "properties": {
             "command": {"type": "string"}}, "required": ["command"]}),
]

_PROVIDERS_TOOLS: list[Tool] = [
    Tool(name="providers.list",
         description="Elenca i provider di inferenza e il loro stato (id/nome/meccanismo/connesso/pausa). MAI segreti.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="providers.pause",
         description="Mette in pausa un provider (escluso dalla selezione; gli agent ripiegano sul prossimo). Non tocca la chiave.",
         inputSchema={"type": "object", "properties": {
             "provider_id": {"type": "string"}}, "required": ["provider_id"]}),
    Tool(name="providers.resume",
         description="Riattiva un provider in pausa.",
         inputSchema={"type": "object", "properties": {
             "provider_id": {"type": "string"}}, "required": ["provider_id"]}),
]

# integrations.* — osservazione dei connettori/integration (stato di connessione).
_INTEGRATIONS_TOOLS: list[Tool] = [
    Tool(name="integrations.list",
         description=("Osserva le integration/connettori e il loro stato di connessione "
                      "(id/nome/provider/connected). NON legge i dati che veicolano."),
         inputSchema={"type": "object", "properties": {}}),
]

# mcp.* — registra/rimuove/elenca i server MCP montati (gateway-local). Sysadmin.
_MCP_TOOLS: list[Tool] = [
    Tool(name="mcp.list",
         description="Elenca i server MCP disponibili (backend montati + namespace nativi).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="mcp.add",
         description=("Registra uno o più server MCP da un config in stile mcp.json "
                      "(oggetto con chiave `mcpServers`). I segreti (secrets: {NAME: val}) "
                      "sono depositati nel vault, mai nel config. Stessa diligenza "
                      "supply-chain dei pack."),
         inputSchema={"type": "object", "properties": {
             "config": {"type": "object", "description": "config con mcpServers"},
             "secrets": {"type": "object", "description": "segreti {NAME: valore} da mettere nel vault"}},
             "required": ["config"]}),
    Tool(name="mcp.remove",
         description="Smonta un server MCP montato per nome (slug).",
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}}, "required": ["name"]}),
]

# settings.* — superficie conversazionale per i settings della piattaforma
# (oggi: backup). SOLO super-agent. MAI segreti (passphrase/credenziali si
# impostano dalla pagina Settings via paste-key).
_IMAGE_TOOLS: list[Tool] = [
    Tool(
        name="image.generate",
        description=("Genera un'immagine PNG (OpenAI gpt-image) dal prompt e la salva "
                     "nei file del topic (files/<filename>). Usa la API key OpenAI del "
                     "vault (server-side, mai esposta). Ritorna il path del file salvato; "
                     "scaricabile via /files/download per portarlo nella working copy."),
        inputSchema={
            "type": "object",
            "properties": {
                "tier": {"type": "string", "description": "tier del topic in cui salvare"},
                "name": {"type": "string", "description": "nome del topic in cui salvare"},
                "prompt": {"type": "string", "description": "prompt fully-baked dell'immagine"},
                "filename": {"type": "string",
                             "description": "nome file PNG di destinazione (es. cover.png)"},
                "size": {"type": "string", "enum": ["1024x1024", "1536x1024", "1024x1536"],
                         "description": "default 1024x1024"},
                "quality": {"type": "string", "enum": ["low", "medium", "high", "auto"],
                            "description": "default auto"},
                "background": {"type": "string", "enum": ["opaque", "transparent", "auto"],
                               "description": "default auto"},
            },
            "required": ["tier", "name", "prompt", "filename"],
        }),
]

_SETTINGS_TOOLS: list[Tool] = [
    Tool(name="settings.backup_get",
         description="Backup della piattaforma (ISO 27001 A.8.13): configurazione SENZA segreti, stato e ultimo snapshot.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="settings.backup_set",
         description=("Aggiorna i campi NON-segreti del backup (backend, repository, "
                      "schedule cron, retention {daily,weekly,monthly}). Le credenziali e la "
                      "passphrase NON si impostano qui: vanno inserite dalla pagina Settings."),
         inputSchema={"type": "object", "properties": {
             "backend": {"type": "string"}, "repository": {"type": "string"},
             "schedule": {"type": "string"},
             "retention": {"type": "object"}}}),
    Tool(name="settings.backup_run",
         description="Esegue subito un backup completo (snapshot + retention + verifica integrità).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="settings.backup_restore_test",
         description="Restore-test: ripristina l'ultimo snapshot in area temporanea e verifica (evidenza A.8.13).",
         inputSchema={"type": "object", "properties": {}}),
]

# gdrive.* — export/import file fra i topic e Drive, riusando le credenziali
# Workspace nel vault. Trasferimento via scratch (come topic.fetch/put).
_GDRIVE_TOOLS: list[Tool] = [
    Tool(name="gdrive.list",
         description=("Elenca file/cartelle di Google Drive. folder_id per il contenuto di "
                      "una cartella; query per una query Drive (es. \"name contains 'x'\")."),
         inputSchema={"type": "object", "properties": {
             "folder_id": {"type": "string"}, "query": {"type": "string"},
             "limit": {"type": "integer"}, "account": {"type": "string"}}}),
    Tool(name="gdrive.search",
         description="Cerca file/cartelle Drive per nome (match parziale).",
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}, "limit": {"type": "integer"},
             "account": {"type": "string"}}, "required": ["name"]}),
    Tool(name="gdrive.mkdir",
         description="Crea una cartella Drive (riusa una omonima nello stesso parent se esiste).",
         inputSchema={"type": "object", "properties": {
             "name": {"type": "string"}, "parent_id": {"type": "string"},
             "account": {"type": "string"}}, "required": ["name"]}),
    Tool(name="gdrive.upload",
         description=("Carica un file su Drive. src = path di un file nello scratch dell'agent "
                      "(prepara con topic.fetch). name = nome su Drive; folder_id = cartella."),
         inputSchema={"type": "object", "properties": {
             "src": {"type": "string"}, "name": {"type": "string"},
             "folder_id": {"type": "string"}, "account": {"type": "string"}},
             "required": ["src"]}),
    Tool(name="gdrive.download",
         description=("Scarica un file Drive in dest (path scratch dell'agent; poi usa topic.put "
                      "per metterlo nel topic). I Google-doc nativi sono esportati (PDF/xlsx)."),
         inputSchema={"type": "object", "properties": {
             "file_id": {"type": "string"}, "dest": {"type": "string"},
             "account": {"type": "string"}}, "required": ["file_id", "dest"]}),
    Tool(name="gdrive.share",
         description="Condivide un file/cartella Drive con un'email. role: writer (editor, default)|reader|commenter.",
         inputSchema={"type": "object", "properties": {
             "file_id": {"type": "string"}, "email": {"type": "string"},
             "role": {"type": "string"}, "account": {"type": "string"}},
             "required": ["file_id", "email"]}),
    Tool(name="gdrive.rename",
         description="Rinomina un file/cartella Drive (anche sui Shared Drive).",
         inputSchema={"type": "object", "properties": {
             "file_id": {"type": "string"}, "new_name": {"type": "string"},
             "account": {"type": "string"}},
             "required": ["file_id", "new_name"]}),
    Tool(name="gdrive.move",
         description=("Sposta un file/cartella in un'altra cartella Drive "
                      "(folder_id di destinazione; anche sui Shared Drive)."),
         inputSchema={"type": "object", "properties": {
             "file_id": {"type": "string"}, "folder_id": {"type": "string"},
             "account": {"type": "string"}},
             "required": ["file_id", "folder_id"]}),
]

# gcalendar.* — Google Calendar sulla stessa credenziale Workspace (scope calendar
# già incluso). Orari ISO8601/RFC3339 (es. 2026-07-22T15:00:00+02:00).
_GCALENDAR_TOOLS: list[Tool] = [
    Tool(name="gcalendar.list_calendars",
         description="Elenca i calendari accessibili con l'account Workspace.",
         inputSchema={"type": "object", "properties": {"account": {"type": "string"}}}),
    Tool(name="gcalendar.list_events",
         description=("Elenca eventi di un calendario in una finestra temporale. "
                      "time_min/time_max ISO8601; query = testo libero opzionale."),
         inputSchema={"type": "object", "properties": {
             "calendar_id": {"type": "string", "description": "default 'primary'"},
             "time_min": {"type": "string"}, "time_max": {"type": "string"},
             "query": {"type": "string"}, "limit": {"type": "integer"},
             "account": {"type": "string"}}}),
    Tool(name="gcalendar.create_event",
         description=("Crea un evento. start/end ISO8601 (dateTime) o date (YYYY-MM-DD "
                      "se all_day=true). attendees = lista di email."),
         inputSchema={"type": "object", "properties": {
             "summary": {"type": "string"}, "start": {"type": "string"},
             "end": {"type": "string"}, "calendar_id": {"type": "string"},
             "description": {"type": "string"}, "location": {"type": "string"},
             "attendees": {"type": "array", "items": {"type": "string"}},
             "all_day": {"type": "boolean"}, "account": {"type": "string"}},
             "required": ["summary", "start", "end"]}),
    Tool(name="gcalendar.update_event",
         description="Modifica un evento esistente (solo i campi passati).",
         inputSchema={"type": "object", "properties": {
             "event_id": {"type": "string"}, "calendar_id": {"type": "string"},
             "summary": {"type": "string"}, "start": {"type": "string"},
             "end": {"type": "string"}, "description": {"type": "string"},
             "location": {"type": "string"}, "account": {"type": "string"}},
             "required": ["event_id"]}),
    Tool(name="gcalendar.delete_event",
         description="Elimina un evento dal calendario.",
         inputSchema={"type": "object", "properties": {
             "event_id": {"type": "string"}, "calendar_id": {"type": "string"},
             "account": {"type": "string"}}, "required": ["event_id"]}),
    Tool(name="gcalendar.freebusy",
         description="Ritorna gli intervalli occupati (busy) in una finestra temporale.",
         inputSchema={"type": "object", "properties": {
             "time_min": {"type": "string"}, "time_max": {"type": "string"},
             "calendar_id": {"type": "string"}, "account": {"type": "string"}},
             "required": ["time_min", "time_max"]}),
]

# gdocs.* — Google Docs sulla stessa credenziale Workspace (scope documents).
_GDOCS_TOOLS: list[Tool] = [
    Tool(name="gdocs.create",
         description="Crea un Google Doc (opz. con testo iniziale). Ritorna id + url.",
         inputSchema={"type": "object", "properties": {
             "title": {"type": "string"}, "text": {"type": "string"},
             "account": {"type": "string"}}, "required": ["title"]}),
    Tool(name="gdocs.read",
         description="Legge il testo di un Google Doc (estratto plain-text).",
         inputSchema={"type": "object", "properties": {
             "document_id": {"type": "string"}, "account": {"type": "string"}},
             "required": ["document_id"]}),
    Tool(name="gdocs.append_text",
         description="Aggiunge testo in fondo a un Google Doc.",
         inputSchema={"type": "object", "properties": {
             "document_id": {"type": "string"}, "text": {"type": "string"},
             "account": {"type": "string"}}, "required": ["document_id", "text"]}),
    Tool(name="gdocs.replace_text",
         description="Sostituisce tutte le occorrenze di `find` con `replace` nel Doc.",
         inputSchema={"type": "object", "properties": {
             "document_id": {"type": "string"}, "find": {"type": "string"},
             "replace": {"type": "string"}, "match_case": {"type": "boolean"},
             "account": {"type": "string"}},
             "required": ["document_id", "find", "replace"]}),
]

# gsheets.* — Google Sheets on the same Workspace credential. The Sheets API
# accepts the `auth/drive` scope the connector already requests, so this needed
# no new consent (clodia-platform#118). Every verb is INCREMENTAL: before this,
# acting on a spreadsheet meant gdrive.download + gdrive.upload, which replaces
# the file and destroys the tabs the agent did not author.
_GSHEETS_TOOLS: list[Tool] = [
    Tool(name="gsheets.list_tabs",
         description=("Tab di un Google Sheet: titolo, id, posizione e dimensioni. "
                      "Da chiamare prima di leggere o scrivere, per sapere i nomi."),
         inputSchema={"type": "object", "properties": {
             "spreadsheet_id": {"type": "string"}, "account": {"type": "string"}},
             "required": ["spreadsheet_id"]}),
    Tool(name="gsheets.read",
         description=("Legge i valori di un range A1 (es. 'Foglio1!A1:D20') oppure di "
                      "un'intera tab. Senza range né tab legge la PRIMA tab. "
                      "formulas=true restituisce il TESTO delle formule invece del "
                      "valore calcolato: usalo se devi riprodurre il foglio, "
                      "altrimenti una formula ti torna come numero."),
         inputSchema={"type": "object", "properties": {
             "spreadsheet_id": {"type": "string"}, "range": {"type": "string"},
             "tab": {"type": "string"}, "formulas": {"type": "boolean"},
             "account": {"type": "string"}},
             "required": ["spreadsheet_id"]}),
    Tool(name="gsheets.add_tab",
         description=("Aggiunge una tab a un Google Sheet ESISTENTE lasciando intatte "
                      "le altre (mutazione, non riscrittura del file). Errore "
                      "azionabile se il titolo è già in uso."),
         inputSchema={"type": "object", "properties": {
             "spreadsheet_id": {"type": "string"}, "title": {"type": "string"},
             "index": {"type": "integer"}, "account": {"type": "string"}},
             "required": ["spreadsheet_id", "title"]}),
    Tool(name="gsheets.append_rows",
         description=("Aggiunge righe DOPO l'ultima riga popolata di una tab: non "
                      "sovrascrive nulla. È la scrittura da preferire. I valori "
                      "entrano come se digitati (le formule restano formule)."),
         inputSchema={"type": "object", "properties": {
             "spreadsheet_id": {"type": "string"}, "tab": {"type": "string"},
             "rows": {"type": "array", "items": {"type": "array"}},
             "account": {"type": "string"}},
             "required": ["spreadsheet_id", "tab", "rows"]}),
    Tool(name="gsheets.write_range",
         description=("Scrive in un range A1 esplicito SOVRASCRIVENDO le celle "
                      "esistenti. Nessun default e nessuna scorciatoia su tab intera: "
                      "per aggiungere dati usa append_rows."),
         inputSchema={"type": "object", "properties": {
             "spreadsheet_id": {"type": "string"}, "range": {"type": "string"},
             "values": {"type": "array", "items": {"type": "array"}},
             "account": {"type": "string"}},
             "required": ["spreadsheet_id", "range", "values"]}),
]

# egress.*/ingress.* — amministrazione delle due whitelist (clodia-platform#128).
# `allow` è GATED: allarga un permesso, e il dialog dice cosa costa. `revoke` e
# `list` no — togliere autorità e leggerla non richiedono un consenso, e chiederlo
# insegnerebbe che anche restringere è un'operazione da negoziare.
_EGRESS_ADMIN_TOOLS: list[Tool] = [
    Tool(name="egress.allow",
         description=("Aggiunge una DESTINAZIONE ammessa in uscita, per tutti gli "
                      "agenti. Notazione URI: mailto:x@y.it · mailto:*@dominio · "
                      "tg:<chat> · https://host/ · https://github.com/owner/repo · "
                      "gdrive://<folder-id> (o l'URL del browser) · gsheets:<id>. "
                      "Richiede approvazione umana."),
         inputSchema={"type": "object", "properties": {"uri": {"type": "string"}},
                      "required": ["uri"]}),
    Tool(name="egress.revoke",
         description="Rimuove una destinazione ammessa. Non richiede approvazione.",
         inputSchema={"type": "object", "properties": {"uri": {"type": "string"}},
                      "required": ["uri"]}),
    Tool(name="egress.list",
         description="Destinazioni ammesse in uscita e modo del confinamento.",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="ingress.allow",
         description=("Aggiunge una FONTE FIDATA: leggere da lì non contaminerà "
                      "più il canale. Notazione URI: mailfrom:x@y.it · "
                      "https://host/prefisso/ · gdrive://<folder-id> · "
                      "gsheets:<id>. Richiede approvazione umana, e il dialog "
                      "avverte che da quel momento le istruzioni nascoste in quella "
                      "fonte non produrranno più nessun segnale."),
         inputSchema={"type": "object", "properties": {"uri": {"type": "string"}},
                      "required": ["uri"]}),
    Tool(name="ingress.revoke",
         description="Rimuove una fonte fidata. Non richiede approvazione.",
         inputSchema={"type": "object", "properties": {"uri": {"type": "string"}},
                      "required": ["uri"]}),
    Tool(name="ingress.list",
         description="Fonti fidate dichiarate.",
         inputSchema={"type": "object", "properties": {}}),
]

# telegram.* — invio + inbound con lease per-chat. Un agente scrive solo a chat
# che hanno già scritto e di cui detiene il lease; chat diverse → lease
# indipendenti. Il bot token vive nel vault, mai nel modello.
_TELEGRAM_TOOLS: list[Tool] = [
    Tool(name="telegram.inbox",
         description=("Chat Telegram con messaggi in arrivo (metadati, NON consuma): "
                      "per ognuna chat_id, titolo, n. messaggi pendenti, anteprima e chi "
                      "detiene il lease. Punto di partenza prima di prendere un lease."),
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="telegram.lease_acquire",
         description=("Acquisisce il lease ESCLUSIVO su una chat per N minuti: finché è "
                      "valido sei l'unico a consumarne i messaggi e a poterle scrivere. "
                      "Fallisce se un altro agente la detiene. Solo chat che hanno scritto."),
         inputSchema={"type": "object", "properties": {
             "chat_id": {"type": "string", "description": "ID chat (da telegram.inbox)"},
             "minutes": {"type": "integer", "description": "Durata lease (default 10, max 120)"}},
             "required": ["chat_id"]}),
    Tool(name="telegram.poll",
         description=("Consuma (svuota) i messaggi in coda di una chat. Richiede un lease "
                      "attivo del chiamante su quella chat."),
         inputSchema={"type": "object", "properties": {
             "chat_id": {"type": "string"}}, "required": ["chat_id"]}),
    Tool(name="telegram.send",
         description=("Invia un messaggio a una chat/gruppo (lease-free: sei l'unico "
                      "mittente). Vincolo Telegram: la chat deve aver già contattato il "
                      "bot, o il bot dev'essere membro del gruppo. `chat_id` accetta anche "
                      "il NOME del gruppo."),
         inputSchema={"type": "object", "properties": {
             "chat_id": {"type": "string"}, "text": {"type": "string"}},
             "required": ["chat_id", "text"]}),
    Tool(name="telegram.send_file",
         description=("Invia un FILE del topic a una chat/gruppo Telegram come allegato "
                      "(o come foto se è un'immagine). Solo tu (messaggero) puoi spedire. "
                      "Passa `chat_id` (id o NOME del gruppo) e `path` (il file dentro il "
                      "topic, es. `files/foo.png`): il TOPIC si ricava dal gruppo. `tier`/"
                      "`name` solo se il file è in un topic diverso da quello del gruppo."),
         inputSchema={"type": "object", "properties": {
             "chat_id": {"type": "string", "description": "chat_id o nome del gruppo"},
             "path": {"type": "string", "description": "path del file nel topic, es. files/foo.png"},
             "tier": {"type": "string", "description": "opzionale (override topic)"},
             "name": {"type": "string", "description": "opzionale: nome del topic (override)"},
             "caption": {"type": "string"}},
             "required": ["chat_id", "path"]}),
    Tool(name="telegram.lease_release",
         description="Rilascia anticipatamente il lease su una chat (no-op se non lo detieni).",
         inputSchema={"type": "object", "properties": {
             "chat_id": {"type": "string"}}, "required": ["chat_id"]}),
    Tool(name="telegram.listen",
         description=("Collega una chat Telegram a un topic: da ora il messaggero ne "
                      "riporta VERBATIM i messaggi nella chat del topic, con l'handle "
                      "autenticato del mittente. Il messaggero NON esegue né risponde "
                      "ai messaggi: riportano soltanto, decidono gli agenti del topic. "
                      "Richiede che tu sia partecipante del topic. Binding a livello di "
                      "istanza: puoi ascoltare più chat."),
         inputSchema={"type": "object", "properties": {
             "tier": {"type": "string"}, "name": {"type": "string"},
             "chat_id": {"type": "string"}},
             "required": ["tier", "name", "chat_id"]}),
    Tool(name="telegram.unlisten",
         description=("Scollega una chat Telegram da un topic: il messaggero smette di "
                      "riportarne i messaggi. Simmetrico a telegram.listen."),
         inputSchema={"type": "object", "properties": {
             "tier": {"type": "string"}, "name": {"type": "string"},
             "chat_id": {"type": "string"}},
             "required": ["tier", "name", "chat_id"]}),
]


# memory.* — seed memory scrivibile dell'agente (universale, non richiede grant).
_MEMORY_TOOLS: list[Tool] = [
    Tool(name="memory.read",
         description=("Legge un file della tua seed memory (default `MEMORY.md`, la tua "
                      "memoria di note/esperienza sempre disponibile). La memory è "
                      "condivisa fra le tue istanze."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string", "description": "default MEMORY.md"}}}),
    Tool(name="memory.write",
         description=("Scrive (sovrascrive) un file della tua seed memory. Usa per "
                      "aggiornare note durature o dati strutturati (es. una whitelist "
                      "JSON). Cap 64KB per file."),
         inputSchema={"type": "object", "properties": {
             "content": {"type": "string"},
             "filename": {"type": "string", "description": "default MEMORY.md"}},
             "required": ["content"]}),
    Tool(name="memory.append",
         description="Aggiunge una nota in coda a un file della seed memory (default MEMORY.md).",
         inputSchema={"type": "object", "properties": {
             "content": {"type": "string"},
             "filename": {"type": "string", "description": "default MEMORY.md"}},
             "required": ["content"]}),
    Tool(name="memory.list",
         description="Elenca i file di NOTE (testo) nella tua seed memory.",
         inputSchema={"type": "object", "properties": {}}),
    # Document store per-seed: DOCUMENTI (PDF/docx/dataset…) che sopravvivono agli
    # spawn, in agents/<seed>/files/. NON caricati in automatico: leggili su richiesta.
    Tool(name="memory.put_document",
         description=("Salva un DOCUMENTO (PDF, docx, xlsx, dataset, immagine…) nella tua "
                      "libreria personale del seed (persistente, sopravvive agli spawn). "
                      "content_b64 = contenuto in base64. Max 25MB. ⚠️ SOLO per file "
                      "PICCOLI: per binari o file grandi usa memory.put (byte dallo "
                      "scratch, niente base64 nel modello)."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}, "content_b64": {"type": "string"}},
             "required": ["filename", "content_b64"]}),
    Tool(name="memory.put",
         description=("Salva un file (anche BINARIO/GRANDE) nella tua libreria del seed "
                      "leggendo i BYTE dal tuo scratch — niente base64 nel modello (come "
                      "topic.put). USA QUESTO per PDF/immagini/file grandi. src = path "
                      "assoluto del file nel tuo scratch; filename = nome nella libreria."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}, "src": {"type": "string",
                 "description": "path assoluto del file nel tuo scratch"}},
             "required": ["filename", "src"]}),
    Tool(name="memory.fetch",
         description=("Scarica un DOCUMENTO della tua libreria nel TUO scratch (byte su "
                      "file, niente base64 nel modello — come topic.fetch). USA QUESTO al "
                      "posto di memory.get_document per binari/file grandi. filename = "
                      "nome nella libreria; dest = path assoluto nel tuo scratch."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}, "dest": {"type": "string",
                 "description": "path assoluto di destinazione nel tuo scratch"}},
             "required": ["filename", "dest"]}),
    Tool(name="memory.list_documents",
         description="Elenca i DOCUMENTI nella tua libreria del seed (nome + dimensione).",
         inputSchema={"type": "object", "properties": {}}),
    Tool(name="memory.read_document",
         description=("Legge un DOCUMENTO della tua libreria estraendone il TESTO "
                      "(PDF/docx/xlsx/txt/md → testo per l'uso). `max_chars` opzionale."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}, "max_chars": {"type": "integer"}},
             "required": ["filename"]}),
    Tool(name="memory.get_document",
         description=("Recupera un DOCUMENTO della libreria come base64 grezzo (per "
                      "ri-allegarlo o passarlo a un tool che accetta binari). ⚠️ SOLO "
                      "file PICCOLI: per binari/file grandi usa memory.fetch (byte nello "
                      "scratch)."),
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}}, "required": ["filename"]}),
    Tool(name="memory.delete_document",
         description="Rimuove un documento dalla tua libreria del seed.",
         inputSchema={"type": "object", "properties": {
             "filename": {"type": "string"}}, "required": ["filename"]}),
]


def _dispatch_profile(name: str, a: dict, caller: str | None):
    from . import profile as prof
    sub = name.split(NS_SEP_DOT, 1)[1]
    target = a.get("agent") or caller
    if sub == "get":
        return prof.get(caller, target)
    if sub == "set":
        return prof.set_fields(caller, target, a.get("fields") or {})
    if sub == "list_files":
        return {"files": prof.list_files(caller, target)}
    if sub == "read_file":
        import base64 as _b64
        raw = prof.read_file(caller, target, a["filename"])
        try:
            return {"filename": a["filename"], "text": raw.decode("utf-8")}
        except UnicodeDecodeError:
            return {"filename": a["filename"], "encoding": "base64", "data": _b64.b64encode(raw).decode()}
    if sub == "grant":
        return prof.grant(caller, target, a["grantee"], bool(a.get("granted", True)))
    raise ValueError(f"unknown profile tool: {name}")


def _dispatch_settings(name: str, arguments: dict, agent: str | None):
    # Autorizzazione GIÀ fatta a monte in call_tool: super-agent (bypass) OPPURE
    # verbo nella whitelist `tool_permissions` dell'agente (per-verbo) + M-gate.
    # NIENTE guard super-only qui: era ridondante e SBAGLIATO — ri-bloccava anche
    # chi ha il grant puntuale (es. sysadmin ha `settings.backup_run` per il
    # backup pre-flight delle migrazioni dati dei pack, ma non backup_set/get).
    from . import backup
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "backup_get":
        return backup.config_redacted()
    if sub == "backup_set":
        return backup.set_config(arguments or {})
    if sub == "backup_run":
        return backup.run_backup()
    if sub == "backup_restore_test":
        return backup.restore_test()
    raise ValueError(f"unknown settings tool: {name}")


def _dispatch_github(name: str, a: dict):
    """`github.*` — le azioni git che escono dallo scope (§5.2).

    Tre cose si decidono QUI e non nel modulo: in quale stanza siamo, se il
    repository è nella lista di quella stanza, e con quale credenziale. Sono le
    tre che l'agente non deve poter dire di sé — la stanza arriva dal claim
    firmato, il perimetro dal meta del topic, la credenziale dal vault.
    """
    from .tools import github_repo as gh
    verb = name.split(NS_SEP_DOT, 1)[1]
    tier, tname = _current_topic()
    if not tier:
        # Fuori da una stanza non c'è un perimetro cui appartenere: rifiutare è
        # l'unica risposta che non inventa uno scope.
        raise ValueError("i verbi github.* agiscono dentro un topic: "
                         "questa chiamata non ha un canale")
    svc = _topics()
    repo = a.get("repo") or ""
    if verb in ("clone", "pull_request"):
        canonico = gh.normalize_repo(repo)
        svc._require_approved_repo(canonico, tier, tname)
        token = _repo_credential(svc, tier, tname, canonico)
    if verb == "clone":
        dest = _safe_scratch_path(a["dest"])
        return gh.clone(canonico, dest, token=token, branch=a.get("branch"))
    if verb in ("pull", "push"):
        workdir = _safe_scratch_path(a["dir"])
        # Il repository di questo working tree non lo dice il chiamante: lo dice
        # l'origin che il gateway stesso ha scritto al clone. Prenderlo dal
        # parametro permetterebbe di far passare un repo approvato per
        # autorizzare un push verso un altro.
        canonico = gh.normalize_repo(_origin_of(workdir))
        svc._require_approved_repo(canonico, tier, tname)
        token = _repo_credential(svc, tier, tname, canonico)
        if verb == "pull":
            return gh.pull(workdir, token=token)
        return gh.push(workdir, token=token, branch=a.get("branch"))
    if verb == "pull_request":
        return gh.pull_request(canonico, a["head"], a.get("base") or "main",
                               a["title"], a.get("body") or "", token=token)
    raise ValueError(f"unknown github verb: {name}")


def _origin_of(workdir: str) -> str:
    import subprocess as _sp
    r = _sp.run(["git", "-C", workdir, "remote", "get-url", "origin"],
                capture_output=True, text=True, timeout=30)
    if r.returncode != 0 or not (r.stdout or "").strip():
        raise ValueError(f"nessun origin in {workdir}: clona con github.clone")
    return r.stdout.strip()


def _repo_credential(svc, tier: str, tname: str, repo_canonico: str):
    """La credenziale del MOUNT che porta QUESTO repository nello scope.

    Cercata per repository e non per nome del mount: chi chiama `github.push`
    non sa (e non deve sapere) come l'owner ha battezzato il mount, e chiedergli
    il nome significherebbe lasciargli scegliere quale credenziale usare.
    """
    from .tools import github_repo as gh
    from .topics.service import mounts as _mounts
    try:
        meta, _ = svc._read_meta(tier, tname)
    except Exception:  # noqa: BLE001 — meta illeggibile → nessuna credenziale di mount
        meta = {}
    for m in _mounts(meta):
        if m.get("type") != "git":
            continue
        try:
            if gh.normalize_repo((m.get("config") or {}).get("url") or "") == repo_canonico:
                return svc.git_credential(tier, tname, m.get("name"))[0]
        except Exception:  # noqa: BLE001 — URL anomalo nel meta: non è questo
            continue
    # Nessun mount per questo repo: resta il ripiego dello scope/piattaforma,
    # che `git_credential` rende esplicito. Rifiutare qui romperebbe i repo
    # approvati per lista ma non montati — che la voce 31 prevede.
    return svc.git_credential(tier, tname)[0]


def _dispatch_gdrive(name: str, a: dict):
    from .tools import gdrive as gd
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "list":
        return gd.list_files(folder_id=a.get("folder_id"), query=a.get("query"),
                             limit=a.get("limit", 50), account=a.get("account"))
    if verb == "search":
        return gd.search(a["name"], limit=a.get("limit", 20), account=a.get("account"))
    if verb == "mkdir":
        return gd.mkdir(a["name"], parent_id=a.get("parent_id"), account=a.get("account"))
    if verb == "upload":
        src = _safe_scratch_path(a["src"])  # i byte vengono dallo scratch, mai dal modello
        return gd.upload(src, name=a.get("name"), folder_id=a.get("folder_id"),
                         account=a.get("account"))
    if verb == "download":
        dest = _safe_scratch_path(a["dest"])
        _os.makedirs(_os.path.dirname(dest), exist_ok=True)
        return gd.download(a["file_id"], dest, account=a.get("account"))
    if verb == "share":
        return gd.share(a["file_id"], a["email"], role=a.get("role", "writer"),
                        account=a.get("account"))
    if verb == "rename":
        return gd.rename(a["file_id"], a["new_name"], account=a.get("account"))
    if verb == "move":
        return gd.move(a["file_id"], a["folder_id"], account=a.get("account"))
    raise ValueError(f"unknown gdrive verb: {name}")


def _dispatch_gcalendar(name: str, a: dict):
    from .tools import gcalendar as gc
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "list_calendars":
        return gc.list_calendars(account=a.get("account"))
    if verb == "list_events":
        return gc.list_events(calendar_id=a.get("calendar_id", "primary"),
                              time_min=a.get("time_min"), time_max=a.get("time_max"),
                              query=a.get("query"), limit=a.get("limit", 25),
                              account=a.get("account"))
    if verb == "create_event":
        return gc.create_event(a["summary"], a["start"], a["end"],
                               calendar_id=a.get("calendar_id", "primary"),
                               description=a.get("description"), location=a.get("location"),
                               attendees=a.get("attendees"), all_day=a.get("all_day", False),
                               account=a.get("account"))
    if verb == "update_event":
        return gc.update_event(a["event_id"], calendar_id=a.get("calendar_id", "primary"),
                               summary=a.get("summary"), start=a.get("start"),
                               end=a.get("end"), description=a.get("description"),
                               location=a.get("location"), account=a.get("account"))
    if verb == "delete_event":
        return gc.delete_event(a["event_id"], calendar_id=a.get("calendar_id", "primary"),
                               account=a.get("account"))
    if verb == "freebusy":
        return gc.freebusy(a["time_min"], a["time_max"],
                           calendar_id=a.get("calendar_id", "primary"),
                           account=a.get("account"))
    raise ValueError(f"unknown gcalendar verb: {name}")


def _dispatch_gdocs(name: str, a: dict):
    from .tools import gdocs as gdo
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "create":
        return gdo.create(a["title"], text=a.get("text"), account=a.get("account"))
    if verb == "read":
        return gdo.read(a["document_id"], account=a.get("account"))
    if verb == "append_text":
        return gdo.append_text(a["document_id"], a["text"], account=a.get("account"))
    if verb == "replace_text":
        return gdo.replace_text(a["document_id"], a["find"], a["replace"],
                                match_case=a.get("match_case", True), account=a.get("account"))
    raise ValueError(f"unknown gdocs verb: {name}")


def _dispatch_egress_admin(name: str, a: dict):
    from . import egress as eg
    ns, verb = name.split(NS_SEP_DOT, 1)
    if verb == "allow":
        return eg.allow(ns, a["uri"])
    if verb == "revoke":
        return eg.revoke(ns, a["uri"])
    if verb == "list":
        return eg.listing(ns)
    raise ValueError(f"unknown {ns} verb: {name}")


def _dispatch_gsheets(name: str, a: dict):
    from .tools import gsheets as gsh
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "list_tabs":
        return gsh.list_tabs(a["spreadsheet_id"], account=a.get("account"))
    if verb == "read":
        return gsh.read(a["spreadsheet_id"], range=a.get("range"), tab=a.get("tab"),
                        formulas=bool(a.get("formulas")), account=a.get("account"))
    if verb == "add_tab":
        return gsh.add_tab(a["spreadsheet_id"], a["title"], index=a.get("index"),
                           account=a.get("account"))
    if verb == "append_rows":
        return gsh.append_rows(a["spreadsheet_id"], a["tab"], a["rows"],
                               account=a.get("account"))
    if verb == "write_range":
        return gsh.write_range(a["spreadsheet_id"], a["range"], a["values"],
                               account=a.get("account"))
    raise ValueError(f"unknown gsheets verb: {name}")


def _dispatch_telegram(name: str, a: dict):
    sub0 = name.split(NS_SEP_DOT, 1)[1]
    if sub0 == "roster":
        from .tools import telegram as _tg
        tok = _tg._token_internal()
        # `api_call` ritorna già il `result`: qui è direttamente la lista.
        got = _tg.api_call(tok, "getChatAdministrators", {"chat_id": a["chat_id"]}) or []
        out = []
        for m in got:
            u = m.get("user") or {}
            if u.get("is_bot"):
                continue
            out.append({"uid": str(u.get("id")), "username": u.get("username"),
                        "name": " ".join(x for x in (u.get("first_name"),
                                                     u.get("last_name")) if x),
                        "status": m.get("status")})
        # Telegram non espone l'elenco COMPLETO dei membri di un gruppo a un
        # bot: si ottengono gli amministratori. Lo si dice invece di far
        # credere che la lista sia tutta — chi mappa deve sapere che i membri
        # non amministratori vanno aggiunti a mano dal loro uid.
        return {"members": out, "complete": False,
                "note": ("Telegram espone a un bot i soli amministratori. Gli "
                         "altri membri vanno mappati col loro uid, che compare "
                         "quando scrivono nel gruppo.")}
    if sub0 == "notify_pending":
        from .topics import telegram_notify as _tn
        items = _tn.pending(int(a.get("limit") or 20))
        return {"pending": [{**i, "text": _tn.render(i)} for i in items]}
    if sub0 == "notify_flush":
        from .topics import telegram_notify as _tn
        return _tn.flush(int(a.get("limit") or 20))
    if sub0 == "notify_ack":
        from .topics import telegram_notify as _tn
        return _tn.ack(a["message_id"], a["chat_id"], a["principal"],
                       ok=bool(a.get("ok", True)), error=a.get("error", ""))

    from .tools import telegram as tg
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "inbox":
        return tg.inbox()
    if verb == "lease_acquire":
        return tg.lease_acquire(a["chat_id"], a.get("minutes", 10))
    if verb == "poll":
        return tg.poll(a["chat_id"])
    if verb == "send":
        return tg.send(a["chat_id"], a["text"])
    if verb == "send_file":
        # Legge il file dal topic (compartimento: dev'essere participant) e lo invia.
        # Il TOPIC si ricava dal gruppo (binding chat→topic), così basta chat + path;
        # `tier`/`name` sono override opzionali per topic diversi da quello del gruppo.
        import base64
        import os as _os
        from .tools import telegram_bindings as _tb
        cid = tg._resolve_chat(a["chat_id"])
        tier, tname = a.get("tier"), a.get("name")
        if not (tier and tname):
            b = _tb.get(cid)
            if not b:
                raise ValueError(f"chat {cid} non legata a un topic: passa tier+name del topic")
            tier, tname = b["tier"], b["topic"]
        _require_topic_member(_topics(), tier, tname)
        data = _topics().read_file(tier, tname, a["path"])
        return tg.send_file(cid, _os.path.basename(a["path"]),
                            base64.b64encode(data).decode("ascii"), a.get("caption", ""))
    if verb == "lease_release":
        return tg.lease_release(a["chat_id"])
    if verb in ("listen", "unlisten"):
        # Binding sull'ISTANZA del messaggero (telegram-bindings.json), NON nel
        # meta del topic. Il messaggero dev'essere partecipante del topic in cui
        # ripeterà. Enforcement compartimento come i topic.*.
        from .tools import telegram_bindings as tb
        from .topics.service import _check_channel_cap
        cid = str(a["chat_id"])
        tier, tname = a["tier"], a["name"]
        _require_topic_member(_topics(), tier, tname)
        if verb == "unlisten":
            return {"ok": True, "chat_id": cid, "removed": tb.remove(cid)}
        # listen: SEAL-cap (telegram cappa a SEAL-1) + una chat → un solo binding.
        meta = _topics().open(tier, tname).get("meta", {})
        _check_channel_cap({"type": "telegram"}, meta.get("tier", tier))
        ex = tb.get(cid)
        if ex and (ex.get("tier"), ex.get("topic")) != (tier, tname):
            raise ValueError(
                f"chat {cid} già collegata a {ex.get('tier')}/{ex.get('topic')}: "
                f"fai prima telegram.unlisten lì (una chat → un solo topic)")
        tb.set_binding(cid, agent_name(), tier, tname)
        return {"ok": True, "chat_id": cid, "instance": agent_name(),
                "topic": f"{tier}/{tname}"}
    raise ValueError(f"unknown telegram verb: {name}")


def _dispatch_memory(name: str, a: dict):
    from .tools import memory as mem
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "read":
        return mem.read(a.get("filename"))
    if verb == "write":
        return mem.write(a["content"], a.get("filename"))
    if verb == "append":
        return mem.append(a["content"], a.get("filename"))
    if verb == "list":
        return mem.list_files()
    # ── Document store per-seed ──────────────────────────────────────────────
    if verb == "put_document":
        return mem.put_document(a["filename"], a["content_b64"])
    if verb == "put":
        # Byte dallo scratch dell'agent → libreria del seed (niente base64 nel
        # modello). Il gateway legge il file locale; path validato sotto spawns/.
        src = _safe_scratch_path(a["src"])
        with open(src, "rb") as f:
            data = f.read()
        return mem.put_document_bytes(a["filename"], data)
    if verb == "fetch":
        # Documento del seed → file nello scratch dell'agent (niente base64).
        fn, data = mem.read_document_bytes(a["filename"])
        dest = _safe_scratch_path(a["dest"])
        _os.makedirs(_os.path.dirname(dest) or ".", exist_ok=True)
        with open(dest, "wb") as f:
            f.write(data)
        return {"file": fn, "bytes": len(data), "dest": dest, "ok": True}
    if verb == "list_documents":
        return mem.list_documents()
    if verb == "read_document":
        fn, data = mem.read_document_bytes(a["filename"])
        cap = int(a.get("max_chars") or 60000)
        try:
            text, pages = _extract_document_text(fn, data)
        except Exception as e:  # noqa: BLE001
            return {"ok": False, "error": f"estrazione fallita: {str(e)[:160]}"}
        return {"file": fn, "text": text[:cap], "chars": len(text),
                "pages": pages, "truncated": len(text) > cap}
    if verb == "get_document":
        import base64 as _b64
        fn, data = mem.read_document_bytes(a["filename"])
        return {"file": fn, "bytes": len(data), "encoding": "base64",
                "content_b64": _b64.b64encode(data).decode("ascii")}
    if verb == "delete_document":
        return mem.delete_document(a["filename"])
    raise ValueError(f"unknown memory verb: {name}")


def _native_tool_namespaces() -> list[str]:
    """Namespace dei tool nativi del gateway (per agents.list_tools).

    ⚠️ Questa lista NON è `_all_native_tools()` e la differenza è reale: omette
    `_SETTINGS_TOOLS`, quindi `settings` non compare fra i namespace concedibili.
    Non la si allinea qui perché cambiare cosa è concedibile è una scelta di
    policy, non un refactor — e `settings.*` è gated globalmente per prefisso,
    quindi l'omissione può essere deliberata. Va deciso, non uniformato di
    soppiatto: clodia-platform#140.
    """
    tools = (_FS_TOOLS + _WEB_TOOLS + _LOGS_TOOLS + _EMAIL_TOOLS + _TOPIC_TOOLS + _GITHUB_TOOLS + _IMAGE_TOOLS
             + _RUNTIME_TOOLS + _JOBS_TOOLS + _PROFILE_TOOLS + _TELEGRAM_TOOLS + _MEMORY_TOOLS + _GDRIVE_TOOLS
             + _GCALENDAR_TOOLS + _GDOCS_TOOLS + _GSHEETS_TOOLS
             + _EGRESS_ADMIN_TOOLS + _AGENT_TOOLS
             + _PACKS_TOOLS + _PROVIDERS_TOOLS + _INTEGRATIONS_TOOLS + _MCP_TOOLS)
    if instance_profile.rag_enabled():
        tools = tools + _EU_CORPUS_TOOLS + _RAG_TOOLS
    ns = sorted({t.name.split(NS_SEP_DOT, 1)[0] for t in tools})
    return ns


def runtime_configuration_warnings() -> list[str]:
    """Coerenza fra whitelist, namespace montati e credenziali consumabili."""
    from .whitelist import CONFIG

    mounted = set(_native_tool_namespaces())
    mounted.update(
        str(backend.get("name") or "")
        for backend in (CONFIG.get("mcp_backends") or [])
    )
    warnings = []
    for agent, spec in (CONFIG.get("agents") or {}).items():
        for tool in spec.get("allowed_tools") or []:
            if tool == "*" or NS_SEP_DOT not in tool:
                continue
            namespace = tool.split(NS_SEP_DOT, 1)[0]
            if namespace not in mounted:
                warnings.append(
                    f"agent '{agent}': namespace '{namespace}' in whitelist "
                    "ma non esposto da tool nativi o backend MCP"
                )
    for row in email.credential_diagnostics():
        if row["operational"]:
            continue
        detail = (
            f"campi mancanti: {', '.join(row['missing'])}"
            if row["missing"] else f"errore: {row['error']}"
        )
        warnings.append(
            f"credenziale email '{row['credential']}' non materializzabile ({detail})"
        )
    return warnings


def _dispatch_agents(name: str, a: dict, caller: str | None,
                     gate_approval: dict | None = None):
    from .tools import agents_admin as adm
    verb = name.split(NS_SEP_DOT, 1)[1]
    if verb == "list":
        return adm.list_agents()
    if verb == "show":
        return adm.show(a["agent"])
    if verb == "list_skills":
        return adm.list_skills()
    if verb == "list_rules":
        return adm.list_rules()
    if verb == "list_tools":
        return {"namespaces": _native_tool_namespaces(),
                "note": "concedi un namespace intero con `<ns>.*` o un tool puntuale `<ns>.<verbo>`"}
    if verb == "grant_skill":
        return adm.grant_skill(a["agent"], a["skill"])
    if verb == "revoke_skill":
        return adm.revoke_skill(a["agent"], a["skill"])
    if verb == "grant_tool":
        return adm.grant_tool(a["agent"], a["tool"])
    if verb == "revoke_tool":
        return adm.revoke_tool(a["agent"], a["tool"])
    if verb == "grant_rule":
        return adm.grant_rule(a["agent"], a["rule"])
    if verb == "revoke_rule":
        return adm.revoke_rule(a["agent"], a["rule"])
    if verb == "grant_scoped":
        token = (gate_approval or {}).get("token")
        if not token:
            raise PermissionError("agents.grant_scoped richiede approvazione firmata one-shot")
        return adm.grant_scoped(a["agent"], a, token)
    if verb == "list_scoped":
        return adm.list_scoped(a["agent"])
    if verb == "revoke_scoped":
        token = (gate_approval or {}).get("token")
        if not token:
            raise PermissionError("agents.revoke_scoped richiede approvazione firmata one-shot")
        return adm.revoke_scoped(a["agent"], a["override_id"], token)
    raise ValueError(f"unknown agents verb: {name}")


def _dispatch_runtime(name: str, arguments: dict, caller: str | None = None):
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "agents":
        return runtime.agents()
    if sub == "jobs":
        return runtime.jobs()
    if sub == "skills":
        return runtime.skills()
    if sub == "chats":
        return runtime.chats()
    if sub == "topics":
        return runtime.topics(include_restricted=bool(arguments.get("include_restricted")))
    if sub == "mcp_servers":
        return runtime.mcp_servers()
    if sub == "providers":
        return runtime.providers()
    if sub == "current_user":
        return runtime.current_user()
    if sub == "restart_agent":
        return runtime.restart_agent(arguments.get("agent"))
    if sub == "inspect_topic":
        return runtime.inspect_topic(arguments.get("tier"), arguments.get("name"),
                                     by=caller or "")
    raise ValueError(f"unknown runtime tool: {name}")


def _dispatch_jobs(name: str, a: dict, caller: str | None):
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "list":
        return runtime.jobs()
    if sub == "propose":
        # l'agente PROPONE un job → l'owner approva via gate. `requested_by` è
        # l'identità del chiamante, impostata qui (non fidarsi dell'input).
        # NB: NESSUN create/delete diretto — anche gli agent di piattaforma
        # (sysadmin) passano dal gate owner. La creazione autonoma ricorrente è
        # superficie di privilegio: deve confermarla l'owner (Prima Legge).
        return runtime.propose_job(
            name=a.get("name"), prompt=a.get("prompt"),
            schedule_text=a.get("schedule_text"), cron_expr=a.get("cron_expr"),
            # default = l'agente CHIAMANTE (non clodia): chi propone un job
            # ricorrente di norma lo esegue lui stesso (es. messaggero/check-email).
            # Evita che l'executor "scivoli" a clodia quando l'agent è omesso.
            agent=a.get("agent") or caller or "clodia", enabled=a.get("enabled", True),
            requested_by=caller or "agente")
    if sub == "report_status":
        # L'ESITO del run, dichiarato dall'agente (clodia-platform#206). Prima lo
        # stato era il valore di verità di «il turno ha sollevato?», che misura la
        # fine del turno e non il lavoro: un job ha girato 652 secondi, ha fallito
        # tre invii e si è registrato `success`.
        #
        # `chat_id` NON è fra gli argomenti, di proposito: viene dal claim `chat`
        # del token di sessione, firmato dall'agent-server. Se lo passasse
        # l'agente, dichiarare l'esito del run di qualcun altro sarebbe questione
        # di cambiare un campo — e nel modello quel campo è testo come tutto il
        # resto. Stessa ragione per cui `topic.fetch`/`topic.put` lo leggono qui.
        chat_id = current_chat()
        if not chat_id:
            raise ValueError(
                "jobs.report_status richiede una sessione agent con chat_id: "
                "si dichiara l'esito di un run schedulato, e fuori da un job "
                "non c'è alcun run da dichiarare")
        return runtime.report_run_status(
            chat_id=chat_id, status=a.get("status"), detail=a.get("detail"),
            agent=caller or "")
    raise ValueError(f"unknown jobs tool: {name}")


def _dispatch_packs(name: str, a: dict):
    from .tools import platform_ops as ops
    from .tools import pack_runtime
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "list":
        return ops.packs_list()
    if sub == "show":
        return ops.packs_show(a["name"])
    if sub == "import_url":
        return ops.packs_import_url(a["url"])
    if sub == "remove":
        return ops.packs_remove(a["name"])
    if sub == "setup_done":
        return ops.packs_setup_done(a["name"])
    if sub == "install_pip":
        return pack_runtime.install_pip(a["packages"])
    if sub == "install_npm":
        return pack_runtime.install_npm(a["packages"])
    if sub == "check_command":
        return pack_runtime.check_command(a["command"])
    raise ValueError(f"unknown packs tool: {name}")


def _dispatch_providers(name: str, a: dict):
    from .tools import platform_ops as ops
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "list":
        return ops.providers_list()
    if sub == "pause":
        return ops.providers_pause(a["provider_id"])
    if sub == "resume":
        return ops.providers_resume(a["provider_id"])
    raise ValueError(f"unknown providers tool: {name}")


def _dispatch_integrations(name: str, a: dict):
    from .tools import platform_ops as ops
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "list":
        return ops.integrations_list()
    raise ValueError(f"unknown integrations tool: {name}")


def _dispatch_mcp(name: str, a: dict):
    """mcp.* — montaggio server MCP (gateway-local, via tools_api core)."""
    from . import tools_api
    sub = name.split(NS_SEP_DOT, 1)[1]
    if sub == "list":
        return runtime.mcp_servers()
    if sub == "add":
        try:
            return tools_api.register_mcp_core(a["config"], a.get("secrets") or {})
        except tools_api.McpRegisterError as e:
            raise PermissionError(str(e)) if getattr(e, "status", 400) == 403 else ValueError(str(e))
    if sub == "remove":
        return tools_api.unregister_mcp_core(a["name"])
    raise ValueError(f"unknown mcp tool: {name}")


# Super-agent nativi: hanno accesso a TUTTI i tool (inclusi i connettori/email
# delegati), bypassando la whitelist per-agent.
# `clodia` NON è più qui (6 ago 2026): «clodia è solo un agent con tanti verbi».
# Il concetto di super resta per ora — rimuoverlo del tutto significa riscrivere
# sette punti con tre definizioni indipendenti, di cui due non riguardano
# l'autorità di un agente ma l'identità di SERVIZIO dell'agent-server (i profili
# umani non hanno una chiave server-side per coniare un token a proprio nome).
# Quel lavoro è tracciato e si fa separatamente; togliere l'attributo è
# l'incremento che si può verificare oggi.
#: VUOTO dal 7 ago 2026. Nessun agente bypassa più la whitelist per il proprio
#: nome. `clodia` era uscita il 6, `ophelia` esce ora — e togliere l'ultima è
#: quello che rende il concetto verificabile invece che convenzionale: finché
#: uno solo resta dentro, la matrice non è mai davvero il documento che decide.
#:
#: Cosa NON era autorità dell'agente, e resta: l'identità di SERVIZIO con cui
#: l'agent-server conia un token per conto di un umano (`clodia`, in gate.py) —
#: i profili umani non hanno una chiave lato server per firmare a proprio nome.
#: Sono due cose che portavano lo stesso nome; questa non passa di qui.
#:
#: L'insieme resta, vuoto, ed estendibile via env: rimetterci un nome è ancora
#: possibile, ma deve essere un atto esplicito di chi amministra l'istanza.
_SUPER_AGENTS: set = set()


def _is_super(name: str | None) -> bool:
    return (name or "") in _SUPER_AGENTS


def _origin_chain(verb: str) -> list:
    """Catena d'origine del turno corrente, con un fallback ESPLICITO.

    Quando il claim non c'è — un agent-server non ancora aggiornato, un percorso
    non strumentato — si ricostruisce la catena minima da ciò che si sa: il
    principal umano se la chiamata è on-behalf, e l'agente che esegue. Non è la
    catena vera (mancano gli anelli intermedi) ma non è nemmeno un via libera, e
    soprattutto non è silenziosa: `origin.evaluate` distingue «sconosciuta» da
    «vuota», e la modalità di osservazione mostra la differenza.
    """
    chain = origin.parse(current_origin())
    if chain:
        return chain
    out = []
    p = current_principal()
    if p and is_on_behalf():
        out.append(("human", p))
    ag = agent_name()
    if ag:
        out.append(("agent", ag))
    return out


def _human_tool_allowed(name: str) -> bool:
    """RBAC UMANA (chiamata on-behalf): il gateway è il PDP unico anche per gli
    umani. Un tool `super-only` (packs/providers/mcp/agents/settings/pki/ca…,
    definita da M-gate) richiede ruolo **admin**; tutto il resto è concesso a
    qualunque umano autenticato. Il ruolo è un claim FIRMATO dall'agent-server →
    non forgiabile dal modello. Chiude la Broken Access Control del path REST."""
    from . import gate as _gate
    if _gate.is_gated(name):
        # `in _ADMIN_ROLES` e NON `== "admin"`. Il confronto letterale escludeva
        # `superadmin`, cioè l'unica persona che possiede l'istanza: il 7 ago 2026
        # Davide si è visto rifiutare `packs.import_url` da amministratore.
        #
        # Il difetto stava in piedi perché ovunque altrove — `human.is_admin`,
        # `admin._is_admin_yaml`, `origin.principal_may`, `tools_api` — la
        # verifica passa da `_ADMIN_ROLES = ("superadmin", "admin")`. Un solo
        # punto duplicava la regola invece di usarla, e le duplicazioni di una
        # regola divergono: questa divergeva sul caso più privilegiato, quindi
        # sbagliava verso il rifiuto e nessun test la vedeva.
        from .human import _ADMIN_ROLES
        return (current_human_role() or "user") in _ADMIN_ROLES
    return True


def _scoped_ceiling_ok(name: str) -> bool:
    """Per una chiamata ON-BEHALF, `scoped_tools` è un **tetto**, non un'aggiunta.

    Sul ramo degli agenti il claim SOMMA: un agente riceve una delega e con essa
    dei verbi in più (`allowed_tools | scoped_tools`). Ha senso lì, dove il token
    *concede*.

    Sul ramo umano non sommava e non limitava: `_human_tool_allowed` guarda solo
    il ruolo, quindi un token coniato per far parlare una persona in **una**
    stanza le apriva ogni verbo non-gated del gateway — leggere qualunque topic,
    scrivere file ovunque. Il claim c'era, era firmato, e non lo leggeva nessuno:
    è il difetto «qualcosa di dichiarato che nessuno porta», nella variante
    peggiore, perché la dichiarazione *sembra* già una restrizione.

    Un tetto assente non è un tetto vuoto: senza il claim vale la RBAC del ruolo
    come prima, altrimenti ogni sessione umana della webui si spegnerebbe.
    """
    tetto = set(current_scoped_tools())
    if not tetto:
        return True
    if name in tetto:
        return True
    # `<ns>.*` resta ammesso — è così che si concede un backend MCP montato per
    # intero. Il `*` nudo no, e non serve chiuderlo solo qui: `mint_session_token`
    # lo rifiuta già alla coniazione. Due controlli sulla stessa cosa divergono;
    # questo si fida di quello, e il test lo verifica sul minter.
    return "." in name and f"{name.split('.', 1)[0]}.*" in tetto


def _vault_grants(agent: str | None) -> set:
    if not agent:
        return set()
    try:
        from . import vault
        return set(vault.grants_for(agent).keys())
    except Exception:  # noqa: BLE001
        return set()


def _declared_tools(agent: str | None) -> set:
    """Verbi EFFETTIVI di un principal: propri, antenati e floor dell'archseed.

    Il fallback sul seed serve agli agenti non registrati — un clone per-topic,
    un responder appena materializzato — per i quali `agent_config` solleva
    KeyError. Senza, l'intersezione qui sotto li ridurrebbe a zero verbi di
    connettore, cioè romperebbe proprio il caso che `_connector_allows` esisteva
    per servire.

    ``effective_tools`` è l'unico risolutore della matrice. La lista grezza in
    config resta la dichiarazione propria dell'agente, non una copia derivata
    dell'ereditarietà che diventerebbe stale quando cambia un antenato.

    Il seed è la fonte autorevole della dichiarazione (il repo lo è del seed), e
    vive dove uno spawn non può riscriverlo: `/datadir/agents/` è `drwx------
    root` e gli spawn girano come uid 60000.
    """
    from .whitelist import effective_tools
    return effective_tools(agent)


def _agent_tool_reachable(name: str, agent: str | None,
                          resolved: set | None = None) -> bool:
    """One allow-matrix decision for agent discovery and dispatch.

    ``effective_tools`` is authoritative: it resolves the seed, its ancestry and
    the archseed floor. Scoped grants are request-local additions. Explicit
    denies remain a separate subtraction because dispatch reports their more
    useful denial reason after this reachability check.
    """
    allowed = (resolved if resolved is not None
               else _declared_tools(agent) | set(current_scoped_tools()))
    return _tool_allowed(name, allowed) or _connector_allows(name, agent)


def _connector_intersect_on() -> bool:
    """Interruttore d'emergenza. Acceso per default perché la misura lo sostiene —
    nessuno dei verbi che l'intersezione toglie è mai stato usato secondo la
    telemetria delle due istanze — ma una telemetria è una finestra, non la
    storia completa, e se salta fuori un flusso legittimo va sbloccato in un
    minuto senza un deploy."""
    return (_os.environ.get("CLODIA_CONNECTOR_INTERSECT") or "on").strip().lower() != "off"


def _connector_allows(name: str, agent: str | None) -> bool:
    """Un grant sul vault apre la CREDENZIALE, non i verbi.

    Com'era e perché è cambiato. Questa funzione ritornava True per l'intero
    namespace di una credenziale concessa: chi aveva `google_<account>` otteneva
    `email.*`, `gdrive.*`, `gdocs.*`, `gsheets.*` e `gcalendar.*` — 23 verbi —
    **indipendentemente da ciò che il suo seed dichiara**. Il commento originale
    lo giustificava così: «la delega non dipende da config.yaml (effimero al
    rebuild)».

    Due ragioni per cambiarlo. La prima è che quel presupposto non vale più:
    `config.yaml` sta su un volume del gateway (`/gateway-state`, bind mount) e
    sopravvive alla ricreazione del container — verificato. La seconda è che
    rendeva la dichiarazione **decorativa** su cinque namespace: il refactoring
    per classe di seed, i `profile_tools`, il modello del mestiere non decidevano
    nulla là dove decideva il grant. E il modello di sicurezza afferma che la
    matrice del principal delimita i suoi verbi: su quei namespace era falso.

    Ora servono ENTRAMBI: il grant sulla credenziale **e** la dichiarazione del
    verbo. Concedere l'account a un postino gli dà la posta che dichiara, non il
    Drive che non dichiara. Per dargli il Drive si aggiunge il verbo al suo seed —
    cioè lo si decide, invece di ottenerlo come effetto collaterale.
    """
    grants = _vault_grants(agent)
    if not _grant_covers(name, grants):
        return False
    if not _connector_intersect_on():
        return True          # interruttore spento: comportamento storico
    # L'INTERSEZIONE. Il grant è necessario, non sufficiente.
    return _tool_allowed(name, _declared_tools(agent))


def _grant_covers(name: str, grants: set) -> bool:
    """Il grant copre il namespace del verbo? (metà «credenziale» della regola)"""
    # La credenziale Google UNIFICATA (google_<account>) abilita SIA email.* SIA
    # gdrive.* (ha entrambi gli scope); i legacy gmail_/gworkspace_ restano validi.
    if name.startswith("email.") and any(
            c.startswith("google_") or c.startswith("gmail_") or c.startswith("mailbox_")
            for c in grants):
        return True
    if name.startswith("telegram.") and "telegram_bot_token" in grants:
        return True
    _gws_grant = any(c.startswith("google_") or c.startswith("gworkspace_") for c in grants)
    if name.startswith(("gdrive.", "gcalendar.", "gdocs.", "gsheets.")) and _gws_grant:
        return True
    return False


def _email_account(arguments: dict) -> str:
    """Account per una chiamata email.*: quello richiesto esplicitamente,
    altrimenti l'UNICO account operativo con grant vault. Con zero o più account
    solleva un errore azionabile: nessun fallback verso mailbox inesistenti."""
    agent = agent_name()
    accounts = email.available_accounts(agent)
    acct = (arguments.get("account") or "").strip()
    if acct:
        if acct in accounts:
            return acct
        raise ValueError(
            f"email: account '{acct}' non disponibile per '{agent}'. "
            f"Passa il parametro 'account' con uno di questi valori: {accounts}"
        )
    if len(accounts) == 1:
        return accounts[0]
    if not accounts:
        raise ValueError(
            f"email: nessun account operativo con grant per '{agent}'. "
            "Configura una mailbox e assegna il relativo grant nel vault."
        )
    raise ValueError(
        "email: il parametro 'account' è obbligatorio quando sono disponibili "
        f"più mailbox. Valori accettati: {accounts}"
    )


# Nessun namespace è più universale. `memory` lo era: concesso a ogni agente senza
# comparire nella sua scheda, quindi invisibile leggendo la configurazione e
# impossibile da togliere a uno in particolare.
#
# Non è stato tolto e basta: lo dà ora l'ARCISEED, da cui ogni seed discende
# (specification §1.3). La differenza è che adesso si vede — l'insieme risolto
# dice da dove viene ogni verbo — e si può sottrarre con `denied_tools`, che con
# un namespace universale era impossibile.
#
# L'insieme resta, vuoto, invece di sparire con la sua funzione: se tornasse la
# tentazione di un namespace implicito, il posto è questo, col commento che dice
# perché ne siamo usciti.
_UNIVERSAL_NS: set = set()


def _tool_allowed(name: str, allowed: set) -> bool:
    """True se il tool è in whitelist.

    Supporta il wildcard ``<backend>.*`` (tutti i tool di un backend MCP montato,
    usato dall'Add-MCP UI) e il ``*`` NUDO (tutti i verbi).

    Il `*` nudo non era gestito, e l'asimmetria era una trappola: `allowed_tools:
    ["*"]` sembrava concedere tutto e concedeva **zero**. Funzionava solo per
    `clodia` e `ophelia`, che bypassavano da `_SUPER_AGENTS` — quindi la stessa
    configurazione dava «tutto» o «niente» a seconda del NOME dell'agente. Emerso
    togliendo l'attributo super a clodia: due test sono passati da verdi a rossi
    non perché la logica fosse cambiata, ma perché per la prima volta veniva
    consultata.
    """
    if "*" in allowed:
        return True
    if NS_SEP_DOT in name and name.split(NS_SEP_DOT, 1)[0] in _UNIVERSAL_NS:
        return True
    if name in allowed:
        return True
    if NS_SEP_DOT in name and f"{name.split(NS_SEP_DOT, 1)[0]}.*" in allowed:
        return True
    return False


NS_SEP_DOT = "."


def native_verb_descriptions() -> dict[str, str]:
    """Verbo → prima frase della sua descrizione, per la scheda del seed.

    Solo la prima frase: la descrizione completa di alcuni verbi è un paragrafo
    che spiega quando NON usarli, e in un albero di 159 righe sarebbe illeggibile.
    Chi vuole il resto invoca `--help` del tool.
    """
    out: dict[str, str] = {}
    for t in _all_native_tools():
        d = " ".join((t.description or "").split())
        # Si taglia alla prima frase, ma non su un'abbreviazione tipo `es.`
        for stop in (". ", " — ", ": "):
            i = d.find(stop)
            if 0 < i < 160:
                d = d[:i]
                break
        out[t.name] = d[:160]
    return out


def _all_native_tools() -> list:
    """I Tool nativi. Sorgente unica per `list_tools`, i nomi e le descrizioni.

    Resta fuori `_native_tool_namespaces`, che omette `settings` — una divergenza
    preesistente che non si sana con un refactor (vedi la nota là).
    """
    native = list(_FS_TOOLS + _WEB_TOOLS + _LOGS_TOOLS + _EMAIL_TOOLS
                  + _TOPIC_TOOLS + _GITHUB_TOOLS + _IMAGE_TOOLS + _RUNTIME_TOOLS + _JOBS_TOOLS
                  + _SETTINGS_TOOLS + _PROFILE_TOOLS + _TELEGRAM_TOOLS + _MEMORY_TOOLS
                  + _GDRIVE_TOOLS + _GCALENDAR_TOOLS + _GDOCS_TOOLS + _GSHEETS_TOOLS
                  + _EGRESS_ADMIN_TOOLS + _AGENT_TOOLS + _PACKS_TOOLS
                  + _PROVIDERS_TOOLS + _INTEGRATIONS_TOOLS + _MCP_TOOLS)
    if instance_profile.rag_enabled():
        native += list(_EU_CORPUS_TOOLS + _RAG_TOOLS)
    return native


def all_native_verb_names() -> list[str]:
    """Nomi di TUTTI i verbi nativi del gateway, indipendentemente dal chiamante.

    Distinto da `list_tools()`, che filtra per l'agente corrente: qui serve il
    CATALOGO, per poter espandere un wildcard nella scheda di un seed. Espone i
    soli nomi — nessuna descrizione, nessuno schema: chi chiede l'elenco vuole
    sapere cosa esiste, non come si invoca.
    """
    return [t.name for t in _all_native_tools()]


@app.list_tools()
async def list_tools() -> list[Tool]:
    """Return only the tools allowed for the calling agent (native + proxied)."""
    try:
        me = agent_name()
    except PermissionError:
        return []
    # Stessa sorgente della scheda del seed: due concatenazioni identiche
    # divergerebbero al primo namespace aggiunto — ed è già successo, vedi la nota
    # in `_native_tool_namespaces`.
    native = list(_all_native_tools())
    # C1: tool dei backend MCP montati (namespaced), aggregati dal proxy.
    try:
        proxied = await proxy.list_proxied_tools()
    except Exception:
        proxied = []
    if is_on_behalf():
        # Umano: vede i tool consentiti dal suo RUOLO (admin = tutti; user = solo
        # non super-only). Stesso PDP del dispatch.
        return [t for t in (native + proxied)
                if _human_tool_allowed(t.name) and _scoped_ceiling_ok(t.name)]
    if _is_super(me):
        return [t for t in (native + proxied) if not agent_denies(t.name, me)]
    # Risolvi l'albero una volta per discovery: l'archseed viene dal filesystem
    # e rifarlo per ciascun verbo renderebbe la fonte unica inutilmente costosa.
    resolved = _declared_tools(me) | set(current_scoped_tools())
    return [t for t in (native + proxied)
            if _agent_tool_reachable(t.name, me, resolved)
            and not agent_denies(t.name, me)]


def _gate_notify_principal(agent: str, gate_key: str, principal: str | None) -> bool:
    """Notifica best-effort un gate in attesa al PRINCIPAL sui suoi canali di contatto
    (telegram/email dalla scheda agent). Se il principal non è umano/assente, ripiega
    sul superadmin umano dell'istanza. Non solleva mai (best-effort)."""
    try:
        import httpx as _hx
        from .tools.runtime import AGENT_SERVER_URL
        with _hx.Client(timeout=6.0) as c:
            data = c.get(f"{AGENT_SERVER_URL}/api/agents").json()
        rows = data.get("agents", data) if isinstance(data, dict) else data
        by_name = {a.get("name"): a for a in rows}
        target = by_name.get(principal) if principal else None
        if not target or target.get("type") != "human":
            target = next((a for a in rows if a.get("type") == "human"
                           and a.get("role") == "superadmin"),
                          next((a for a in rows if a.get("type") == "human"), None))
        if not target:
            return False
        ch = target.get("contact_channels") or {}
        tg_id = ch.get("telegram") or target.get("telegram")
        msg = (f"🔐 Gate in attesa: '{agent}' vuole usare `{gate_key}` (azione non "
               f"presidiata). Approva/nega dalla sezione gate della webui.")
        if tg_id:
            from .tools import telegram as _tg
            _tg.send(str(tg_id), msg)
            return True
    except Exception:  # noqa: BLE001 — notifica best-effort, mai bloccante
        pass
    return False


def _reply_recipient(arguments: dict) -> str:
    """Destinatario reale di un `email.reply`, letto dal messaggio originale.

    Stringa vuota se non risolvibile → `egress.check` la tratta come
    destinazione ignota e nega. È la direzione d'errore giusta: «l'attaccante
    scrive, l'agente risponde con i dati» è il percorso dell'injection, e non
    sapere a chi si sta rispondendo non è una buona ragione per procedere.
    """
    try:
        msg = email.read_message(
            str(arguments.get("email_id") or ""),
            account=_email_account(arguments),
            folder=arguments.get("folder") or "INBOX")
    except Exception as e:  # noqa: BLE001 — non risolvibile ≠ consentito
        LOG.warning("egress: destinatario di email.reply non risolvibile (%s)", e)
        return ""
    from . import egress as _eg
    src = msg.get("from") or (msg.get("message") or {}).get("from") or ""
    return _eg.address_of(str(src))


def _push_destination(arguments: dict) -> dict:
    """`github.push` con il repository esplicitato, per chi deve giudicarne la
    destinazione.

    Il verbo riceve una DIRECTORY: il repository sta nel remote di quel working
    tree. Senza risolverlo il PDP vede «nessuna destinazione» e nega — che è la
    direzione giusta per un dubbio, ma qui il dubbio è nostro, non del chiamante.

    Stessa forma di `_reply_recipient` e per la stessa ragione: una lettura in
    più che avviene SOLO per il verbo che ne ha bisogno, al call-site, invece di
    far indovinare il PDP.
    """
    if arguments.get("repo"):
        return arguments
    from .tools import github_repo as _gh
    url = _gh.remote_url(str(arguments.get("dir") or ""))
    return {**arguments, "repo": url} if url else arguments


async def _require_gate_consent(
    agent: str, gate_key: str, *, consume: bool, reason: str = "",
    allow_delegation: bool = True,
) -> dict | None:
    """Block-and-wait sul consenso di gate per (agent, gate_key). Se assente crea
    la richiesta (popup) e ATTENDE la decisione umana (~180s), poi procede; solleva
    su diniego o timeout. `consume`=True → one-shot (verbi); False → time-boxed
    (cross-topic: l'intera operazione sul topic vale finché dura il consenso)."""
    from . import gate as _gate
    from . import observe as _obs
    if _obs.skipping():
        # Osservazione: si registra il gate che sarebbe scattato e si procede.
        # Prima della delega, così anche i gate coperti da delega restano contati
        # — servono a capire quali controlli servono, non a documentare gli unlock.
        _obs.note("gate", gate_key, agent, detail=reason[:120])
        return None
    inst = "-"
    # DELEGA PERMANENTE (async·A): se esiste una delega firmata dall'utente il cui
    # scope copre questa azione (verb == gate_key), è già autorizzata → unlock senza
    # richiesta né blocco. Modello: delega → verifica firma (CA) → covers → unlock.
    if allow_delegation:
        try:
            from . import delegation as _deleg
            _d = _deleg.find_covering(agent, gate_key)
            if _d:
                import logging as _lg
                _lg.getLogger("clodia-tools").info(
                    "gate '%s' autorizzato da delega permanente di '%s'",
                    gate_key, _d.get("principal"))
                return {"delegated": True, "principal": _d.get("principal")}
        except Exception:  # noqa: BLE001 — la delega è additiva: su errore, gate normale
            pass
    if not _gate.active(agent, inst, gate_key):
        req = _gate.request(agent, inst, gate_key, context=current_chat(),
                            human=current_principal(), chat=current_chat(),
                            reason=reason)
        # UX inline: se l'azione parte da un CANALE (chat=chan:tier:name:...),
        # posta un marker nel canale → il webui rende la card Approva/Nega
        # NELLA conversazione (come job-proposal), non nel popup staccato. I gate
        # senza contesto-canale restano gestiti dal popup di fallback.
        _ch = current_chat() or ""
        if _ch.startswith("chan:"):
            try:
                _parts = _ch.split(":")
                _tier, _name = _parts[1], _parts[2]
                # Il messaggio PORTA la richiesta, non solo un puntatore a essa.
                #
                # Conteneva il solo marcatore, e tutto ciò che si leggeva — chi
                # chiede, cosa, perché — veniva reso in diretta dalla coda dei
                # pending. Alla decisione la richiesta esce dalla coda e in chat
                # restava un riquadro vuoto: il motivo per cui qualcuno aveva
                # chiesto quella cosa spariva nel momento esatto in cui diventava
                # una decisione da ricordare. Ed era peggio dopo un ricarico:
                # senza la coda, un gate già deciso tornava a somigliare a uno
                # aperto.
                #
                # Il testo è la traccia durevole; il marcatore resta per i
                # bottoni finché la richiesta è viva.
                _perche = f" — {reason}" if reason else ""
                _cosa = (f"di accedere al topic {gate_key.split(':', 1)[1]}"
                         if gate_key.startswith("topic-access:")
                         else f"di usare `{gate_key}`")
                # kind=ai (non system): i system sono filtrati dal render del webui.
                _topics().post_message(
                    _tier, _name, author="gate",
                    text=(f"🛡️ **{agent}** chiede {_cosa}{_perche}\n"
                          f"<!-- gate={req.get('id')} -->"), kind="ai")
            except Exception:  # noqa: BLE001 — il gate resta valido anche senza marker
                pass
        # Un gate in canale ha una card nella stanza, ma questo non prova che
        # l'umano stia guardando quella stanza. Un gate fuori canale non ha
        # nemmeno la card. In entrambi i casi si notifica best-effort il
        # principal sui suoi canali di contatto; l'approvazione resta sempre la
        # stessa capability di gate.
        import os as _os
        is_channel_gate = _ch.startswith("chan:")
        loops = int(_os.environ.get("GATE_WAIT_LOOPS", "30"))  # 30 = ~60s
        _gate_notify_principal(agent, gate_key, current_principal())
        if not is_channel_gate:
            loops = int(_os.environ.get("GATE_WAIT_LOOPS_ASYNC", "3600"))  # ~2 ore (loop 2s)
        import asyncio as _aio
        approved = False
        for _ in range(loops):
            await _aio.sleep(2)
            if _gate.active(agent, inst, gate_key):
                approved = True
                break
            if not _gate.request_pending(agent, inst, gate_key):
                raise PermissionError(f"gate: '{gate_key}' negato dall'operatore")
        if not approved:
            _gate.resolve_request(agent, inst, gate_key)
            raise PermissionError(f"gate: '{gate_key}' non approvato entro il tempo limite")
    approval = _gate.details(agent, inst, gate_key)
    if not approval:
        raise PermissionError(f"gate: capability per '{gate_key}' non disponibile")
    if consume:
        _gate.consume(agent, inst, gate_key)
    return approval


def _spawn_compartment_mode() -> str:
    """`off` | `report` | `on`. Default `report`: si osserva prima di rifiutare.

    Stessa forma della catena origin, e per la stessa ragione — questa regola
    stringe un permesso che oggi è larghissimo, e stringerlo alla cieca
    romperebbe l'orchestrazione senza che nessuno sappia dove.
    """
    m = (_os.environ.get("CLODIA_SPAWN_COMPARTMENT") or "report").strip().lower()
    return m if m in ("off", "report", "on") else "report"


def _is_portable(meta: dict, agent: str | None) -> bool:
    """Il TOPIC si dichiara portabile, e chi lo porta ne è partecipante.

    Rovesciato l'8 ago 2026. Prima la portabilità era `carries` sul SEED — «i
    topic che questo agente si porta dietro» — ed era il lato sbagliato: un
    agente che si aggiunge un topic alla propria lista **si dà da solo un
    canale** fra le stanze. Dichiarata dal topic, la portabilità è una decisione
    di chi possiede i contenuti (specification §2.4).

    Restano due condizioni, non una: il topic è portabile **e** chi chiede ne è
    partecipante. La prima da sola aprirebbe un topic portabile a chiunque; la
    seconda da sola è la membership normale, che di per sé non attraversa i muri
    (voce 29).

    Nessuno usava `carries` — misurato su venere prima di toglierlo — quindi lo
    spostamento non migra nulla e non rompe niente.
    """
    if not bool(meta.get("portable")):
        return False
    return _topic_is_member(meta, agent or "")


def _require_room_carries(meta: dict, tier: str, tname: str, qui: str | None) -> None:
    """La portabilità avviene solo se la STANZA regge il tier del topic portato.

    Regola di Davide, 7 ago 2026: «se il topic portabile TP ha SEAL-3, allora di
    sicuro un participant Alice sarà SEAL-3 o superiore. Se Alice viene convocata
    in un topic T SEAL-1 semplicemente non avviene la portabilità dei dati».

    È l'anello più debole applicato al trasporto. Il vincolo non sta
    sull'appartenenza — chi partecipa a TP ha già la clearance — ma sulla stanza:
    dati SEAL-3 non entrano in una stanza SEAL-1, perché lì li leggerebbero i
    partecipanti di quella.

    **Si rifiuta, non si gata.** Un gate lascerebbe a qualcuno la facoltà di
    approvare proprio il travaso che questa regola esiste per impedire, e il
    consenso di un owner non alza il tier di una stanza.

    **In un job la regola vale come in una stanza**, da quando il tier del job
    viaggia nel claim firmato (8 ago 2026). Fino a quel giorno qui si consentiva
    e si loggava: era l'unico posto in cui questa regola era scritta e non
    applicata, e la ragione era che il gateway non sapeva il tier — non che il
    caso fosse innocuo.

    Se il tier del job non è dichiarato non si rifiuta: assente significa
    «nessun requisito», che è lo stato di ogni job esistente, e trasformare
    un'assenza in un divieto spegnerebbe lavoro che gira.
    """
    t_topic = _rank(meta.get("tier", tier))
    if not qui:
        from .whitelist import current_scope_tier
        t_job = current_scope_tier()
        if not t_job:
            if t_topic > 0:
                import logging as _lg
                _lg.getLogger("clodia-tools").warning(
                    "portabilità · %s/%s è SEAL-%s e si porta in un'esecuzione "
                    "che non dichiara un tier: consentito", tier, tname, t_topic)
            return
        if _rank(t_job) >= t_topic:
            return
        raise PermissionError(
            f"il topic portabile {tier}/{tname} è SEAL-{t_topic}, e questo job "
            f"dichiara {t_job}: qui la portabilità non avviene, e i suoi dati non "
            f"sono disponibili. Non è un permesso che manca — è il livello "
            f"dichiarato dal job. Alza il tier del job, o leggi quei dati in "
            f"{tier}/{tname}.")
        return
    q_tier, _, _q = qui.partition("/")
    if _rank(q_tier) >= t_topic:
        return
    raise PermissionError(
        f"il topic portabile {tier}/{tname} è SEAL-{t_topic}, questa stanza è "
        f"{q_tier}: qui la portabilità non avviene, e i suoi dati non sono "
        f"disponibili. Non è un permesso che manca — è il livello della stanza. "
        f"Se ti servono quei dati, il posto in cui leggerli è {tier}/{tname}.")


def _cross_topic_gate_key(name: str, arguments: dict, agent: str) -> str | None:
    """Chiave di gate per l'accesso CROSS-TOPIC.

    La regola, dal 7 ago 2026: **la membership del seed non basta più**. Conta la
    stanza in cui lo spawn STA, presa dal claim `chat` FIRMATO — mai da un
    argomento, che sarebbe la parola dell'agente su dove si trova.

        T == qui               → consentito   (agisci nel tuo scope)
        T portabile e sono suo → consentito   (dichiarato dal TOPIC)
        agent ∈ participants(T) → GATE        ← il cambiamento
        altrimenti             → GATE         (invariato)

    Perché. `_topic_is_member` confrontava il nome del SEED con i partecipanti, e
    nessuno guardava da dove partiva la chiamata. Su marte clodia è participant
    di 135 topic su 157: uno spawn di clodia, stando in una stanza qualunque,
    poteva leggere gli altri 134 senza gate e riversarli lì dentro. Il modello
    dichiara due assi — clearance E compartimento — ma il secondo compartimenta
    solo se valutato per SPAWN. Per seed era un permesso globale vestito da
    compartimento.

    La membership resta rilevante, ma cambia ruolo: non toglie il gate, decide a
    CHI è rivolto — l'owner della stanza bersaglio può approvare la propria.
    """
    if NS_SEP_DOT not in name:
        return None
    ns, verb = name.split(NS_SEP_DOT, 1)
    if ns != "topic" or verb not in _TOPIC_SCOPED_VERBS:
        return None
    tier, tname = arguments.get("tier"), arguments.get("name")
    if not (tier and tname):
        return None
    try:
        meta = _topics().open(tier, tname).get("meta", {})
    except Exception:  # noqa: BLE001 — topic inesistente → lascia decidere al dispatch
        return None
    target = f"{meta.get('tier', tier)}/{tname}"
    if _spawn_compartment_mode() == "off":
        return None if _topic_is_member(meta, agent) else f"topic-access:{target}"

    from .whitelist import current_channel
    qui = current_channel()
    if qui and _norm_scope(qui) == _norm_scope(target):
        return None                      # la propria stanza
    if _is_portable(meta, agent):
        _require_room_carries(meta, tier, tname, qui)
        return None                      # portabile, e la stanza lo regge
    if _spawn_compartment_mode() == "report":
        if _topic_is_member(meta, agent):
            import logging as _lg
            _lg.getLogger("clodia-tools").warning(
                "compartimento spawn · %s leggerebbe %s stando in %s: consentito "
                "solo perche' participant del seed (con enforcement: GATE)",
                agent, target, qui or "nessuna stanza")
            return None
        return f"topic-access:{target}"
    return f"topic-access:{target}"


def _norm_scope(x: str) -> str:
    """`SEAL-1/acme` in forma confrontabile. I due lati arrivano da sorgenti
    diverse — il claim firmato e il meta del topic — e un confronto per stringa
    grezza fallirebbe sul primo alias di tier."""
    t, _, n = (x or "").partition("/")
    return f"{_norm_tier_str(t)}/{n.strip().lower()}"


def _norm_tier_str(t: str) -> str:
    t = (t or "").strip().upper()
    legacy = {"P0": "SEAL-0", "P1": "SEAL-1", "P2": "SEAL-2", "P3": "SEAL-3"}
    return legacy.get(t, t)


#: L'UNICO verbo `topic.*` ammesso a una sessione non presidiata: spedire
#: informazione verso un topic. Non legge, non elenca, non scarica.
_UNATTENDED_TOPIC_ALLOW = frozenset({"topic.invoke_hook"})


#: Verbi che leggono FILE di un topic: la fonte è il topic (la sua cartella Drive,
#: o l'etichetta del file locale).
_TOPIC_READ_VERBS = frozenset({
    "topic.read_file", "topic.read_document", "topic.files", "topic.fetch",
    # `remote_pull` scarica il contenuto del remote: è una lettura da quella
    # fonte come le altre, e trattarla diversamente contaminerebbe anche il pull
    # da una cartella vagliata.
    "topic.remote_pull",
})

#: Verbi che leggono una risorsa Drive/Workspace PER ID: la fonte è quella
#: risorsa, e l'argomento che la identifica è diverso per ognuno. Senza questa
#: mappa un `gsheets.read` su un foglio dichiarato fidato contaminerebbe comunque
#: — la regola «verbo + fonte» applicata a metà.
_RESOURCE_READ_VERBS = {
    "gsheets.read": ("spreadsheet_id", "gsheets:{}"),
    "gsheets.list_tabs": ("spreadsheet_id", "gsheets:{}"),
    "gdocs.read": ("document_id", "gdrive:doc/{}"),
    "gdrive.download": ("file_id", "gdrive:file/{}"),
}


def _source_vetted(verb: str, a: dict, result: object = None) -> bool | None:
    """La sorgente di questa lettura è dichiarata fidata?

    `True` non contamina, `False`/`None` sì. `None` è «non determinabile», e si
    tratta come non fidata: una lettura di cui non sappiamo la provenienza non è
    una lettura fidata, e sbagliare in questa direzione è **silenzioso** — un
    taint che non si accende non lo si vede.

    Due sorgenti, due criteri:

    - **file di un topic**: se il topic è collegato a una cartella Drive, la fonte
      è la cartella e vale `source_allow`; altrimenti vale l'ETICHETTA del file —
      `trusted` (dichiarato dall'owner all'upload) e `agent` (output nostro) non
      contaminano, `untrusted` e `unknown` sì;
    - **web**: l'URL contro `source_allow`, per prefisso;
    - **una mail letta**: il mittente, preso dal RISULTATO della chiamata — è già
      lì, e una seconda fetch per rileggerlo sarebbe una richiesta in più per
      un'informazione che abbiamo in mano.

    Regola generale: si valuta solo una chiamata con UNA fonte identificabile.
    `email.list`, `telegram.inbox`, `gcalendar.list_events`
    mescolano più mittenti/autori in una risposta: non c'è una fonte da vagliare,
    e dichiararne una sarebbe peggio che ammettere di non poterlo fare.
    """
    from . import egress as _eg
    try:
        # Un backend MCP montato: la fonte è il SERVER, uno per namespace, e si
        # vaglia come qualunque altra — `mcp:normattiva.` in `source_allow`.
        # Il discriminante è `is_proxied`, non il nome: senza, un
        # `mcp:email.` in lista spegnerebbe il taint sulla posta in arrivo,
        # che ha una fonte diversa a ogni messaggio e va valutata sul risultato.
        if proxy.is_proxied(verb):
            return _eg.is_vetted_source(f"mcp:{verb}")
        if verb.startswith("web."):
            url = str((a or {}).get("url") or "").strip()
            return _eg.is_vetted_source(url) if url else None
        if verb == "email.read":
            src = ""
            if isinstance(result, dict):
                src = str(result.get("from")
                          or (result.get("message") or {}).get("from") or "")
            addr = _eg.address_of(src)
            return _eg.is_vetted_source(f"mailfrom:{addr}") if addr else None
        spec = _RESOURCE_READ_VERBS.get(verb)
        if spec:
            field, tmpl = spec
            rid = str((a or {}).get(field) or "").strip()
            return _eg.is_vetted_source(tmpl.format(rid)) if rid else None
        if verb not in _TOPIC_READ_VERBS:
            return None
        tier, name = (a or {}).get("tier"), (a or {}).get("name")
        if not (tier and name):
            return None
        svc = _topics()
        meta = svc.open(tier, name).get("meta", {})
        rem = (meta.get("remote") or {})
        if str(rem.get("type") or "") == "drive":
            folder = str((rem.get("config") or {}).get("folder") or "").strip()
            # Cartella Drive: la fonte è la cartella, non il singolo file. Le
            # etichette di provenienza non esistono su Drive — chi ci scrive lo fa
            # da fuori — quindi l'unica domanda sensata è se di quella cartella
            # l'owner risponda.
            return _eg.is_vetted_source(f"gdrive:folder/{folder}") if folder else None
        path = str((a or {}).get("path") or "").strip()
        if not path:
            return None                     # `topic.files`: elenco, nessun file
        rel = path[len("files/"):] if path.startswith("files/") else path
        prov = (svc.provenance_map(tier, name).get(rel) or {}).get("provenance")
        if prov in ("trusted", "agent"):
            return True
        return False                        # untrusted, o etichetta assente
    except Exception as e:  # noqa: BLE001 — non determinabile ≠ fidata
        LOG.warning("taint: provenienza di %s non determinabile (%s)", verb, e)
        return None


def agent_name_safe() -> str:
    """`agent_name()` solleva se l'identità non è impostata, e nel ramo di
    rifiuto quello è proprio uno dei casi possibili: un errore nel registro
    nasconderebbe il rifiuto che stiamo registrando."""
    try:
        return agent_name() or "?"
    except Exception:  # noqa: BLE001
        return "?"


#: Quando si toglie un verbo, l'errore deve dire COSA USARE al posto suo.
#: Senza, l'agente conclude che manca un grant o che il server è rotto — è
#: successo con `topic.put` a un postino, che ha provato tre strade e poi ha
#: chiesto aiuto in chat. Un rifiuto senza alternativa è un vicolo cieco, e il
#: modello ci sbatte dentro con l'insistenza di chi non ha altre mosse.
_DENY_HINT = {
    "topic.put": ("Per archiviare un allegato di posta usa "
                  "email.save_attachment(email_id, filename) — senza tier/name "
                  "va nel topic del canale, e i byte non passano da te. Per un "
                  "file che hai prodotto, chiedilo a un agente che ha topic.put."),
    "topic.write_file": ("Per archiviare un allegato di posta usa "
                         "email.save_attachment(email_id, filename)."),
    "topic.read_file": ("Non ti serve leggerlo per spedirlo: "
                        "email.send(topic_files=['files/x.pdf']) e "
                        "telegram.send_file(path='files/x.pdf') li allegano "
                        "senza che il contenuto passi da te."),
    "topic.read_document": ("Per spedire un documento usa "
                            "email.send(topic_files=[...]) senza leggerlo."),
    "topic.fetch": ("Per spedire un file del topic non serve portarselo nello "
                    "scratch: email.send(topic_files=[...])."),
    "topic.files": ("Non hai l'elenco dei file di proposito. Il path di ciò che "
                    "devi spedire arriva nella conversazione."),
    "mcp.add": "Si installa dalla pagina Packs, non da un turno di chat.",
    "packs.install_pip": "Si installa dalla pagina Packs.",
    "packs.install_npm": "Si installa dalla pagina Packs.",
    "settings.backup_run": "Il backup si esegue da un job, non da un turno di chat.",
}


def _current_topic() -> tuple[str | None, str | None]:
    """(tier, name) del canale in cui l'agente sta operando, dal claim `chat`.

    Il gateway lo SA — la chiave di sessione è `chan:<tier>:<nome>:<seed>#<n>` e
    arriva firmata nel token. Chiederlo all'agente significava chiedergli di
    indovinare dove si trova: `topic.open` richiede a sua volta tier+name, e i
    verbi di elenco a un postino sono negati dal §8, quindi l'unica mossa che gli
    restava era domandarlo all'utente. Ed è quello che ha fatto.
    """
    from . import taint as _t
    ch = _t.channel_of(current_chat())
    if not ch or "/" not in ch:
        return None, None
    tier, name = ch.split("/", 1)
    return tier, name


def _topic_of(a: dict) -> tuple[str | None, str | None]:
    """Topic indicato negli argomenti, altrimenti quello del canale corrente.

    L'argomento esplicito resta e vince: serve per operare su un ALTRO topic — e
    in quel caso scatta il gate cross-topic, che è esattamente la differenza fra
    «archivia qui» e «archivia là».
    """
    tier, name = a.get("tier"), a.get("name")
    if tier and name:
        return tier, name
    return _current_topic()


def _topic_attachments(a: dict, agent: str) -> tuple[list, str | None]:
    """Materializza `topic_files` in una dir temporanea e ne ritorna i path.

    Il compartimento vale come per ogni altro accesso: si passa da
    `_require_topic_member`, quindi un agente non allega file di topic di cui non
    è partecipante. La dir temporanea viene cancellata dal chiamante nel `finally`
    — i byte non restano né nello spawn né nel gateway.
    """
    rels = [r for r in (a.get("topic_files") or []) if str(r).strip()]
    if not rels:
        return [], None
    tier, tname = _topic_of(a)
    if not (tier and tname):
        raise ValueError(
            "topic_files: non sono riuscito a ricavare il topic dal canale in cui "
            "stai lavorando — passa tier + name")
    svc = _topics()
    _require_topic_member(svc, tier, tname)
    import os as _os
    import tempfile as _tf
    d = _tf.mkdtemp(prefix="attach-")
    out = []
    for rel in rels:
        data = svc.read_file(tier, tname, str(rel))
        dest = _os.path.join(d, _os.path.basename(str(rel)))
        with open(dest, "wb") as f:
            f.write(data)
        out.append(dest)
    LOG.info("email: %d allegati per riferimento da %s/%s (agent %s)",
             len(out), tier, tname, agent)
    return out, d


def _unattended_denial(verb: str) -> str | None:
    """Blocco per le sessioni di job (clodia-platform#104, decisione 2 ago 2026).

    «Per i job asincroni blocco totale, nessun accesso ai dati dei topic, unica
    possibilità invocare hook dei topic per spedire informazioni.»

    Il motivo per cui un job non si difende con i gate come una chat: **non c'è
    nessuno che possa rispondere**. Un gate in una sessione non presidiata non è
    una protezione, è uno stallo fino al timeout — è la lezione di #116, dove la
    riconciliazione al boot tentava verbi gated senza canale e produceva decine di
    popup fuori contesto. Qui si nega prima, invece di chiedere a nessuno.

    Vale sul CLAIM FIRMATO nel token, non su un parametro: l'agente non può
    dichiararsi presidiato.
    """
    if not is_unattended():
        return None
    if verb.startswith("topic.") and verb not in _UNATTENDED_TOPIC_ALLOW:
        return (f"'{verb}' non è disponibile in un job schedulato: una sessione non "
                f"presidiata non accede ai dati dei topic. Per spedire informazioni "
                f"a un topic usa topic.invoke_hook.")
    return None


def _channel_participants(channel: str | None) -> list:
    """Partecipanti del canale corrente, per la firma di composizione."""
    ch = (channel or "").strip()
    if not ch or "/" not in ch:
        return []
    tier, tname = ch.split("/", 1)
    try:
        return list(_topics().open(tier, tname).get("meta", {}).get("participants") or [])
    except Exception:  # noqa: BLE001 — canale non leggibile → composizione vuota
        return []


def _context_gate_needed(verb: str, agent: str, egress_verdict: dict) -> tuple[str | None, str]:
    """Gate di CONTESTO (#104 §6, #77 «danger score di contesto + gate di uscita»).

    Presidia la TRANSIZIONE, non il verbo: l'agente vive a due lati e i verbi che
    accendono il terzo restano dichiarati e inerti, ma la loro invocazione **in un
    contesto contaminato** passa da un umano.

    Condizione 1 di #77 — serve ANCHE la contaminazione. Col solo punteggio di
    capacità scatterebbe quasi sempre (150 canali su 156 sono a 3/3) e un gate
    approvato per riflesso è peggio di nessun gate.

    Sul primo lato (dati privati) non si interroga il punteggio, che vive in
    clodia-logic: un canale di topic **è** dato privato — i suoi file, il summary,
    la conversazione. Chiamare l'agent-server nel percorso caldo per riscoprirlo
    aggiungerebbe una dipendenza di rete a ogni verbo di uscita.

    UNA DESTINAZIONE CENSITA È PERIMETRO (regola dell'owner, 17 ago 2026):

        «se la destinazione è censita in whitelist allora va considerata come
         parte del perimetro e non deve essere un segnale che fa scattare il gate
         o incrementare il trifecta»

    Questo ROVESCIA la lettura precedente della condizione 2 di #77, che diceva:
    se la destinazione è già in whitelist nessuno guarda, quindi questo gate DEVE
    scattare. Conseguenza misurata su `fullstack-dev`: censire `github.com/
    r-clodia/*` non spegneva niente — lo rendeva obbligatorio. Il ciclo di lavoro
    di un agente di sviluppo lo ri-armava da sé, perché `github.issue_read` e
    `topic.read_file` contaminano e `github.push` è l'uscita: leggi la issue →
    contaminato → push → gate → approvi → declassificato → leggi il file dopo →
    di nuovo. Un'approvazione per ciclo, per progetto. Un gate che si ripete a
    ogni giro è la definizione di consent fatigue, e si approva per riflesso.

    Restano presidiati i due casi in cui la destinazione NON è censita e nessuno
    guarda comunque: `report` (would_deny — fuori lista, non bloccata) e un tipo
    non controllato (`checked: False`, incluso il modo `off`). Lì non c'è nessuna
    dichiarazione di perimetro da rispettare: è l'assenza di confinamento.
    """
    from . import egress as _eg
    if not _eg.spec_for(verb):
        return None, ""                      # non è un verbo di uscita
    if egress_verdict.get("action") == "gate":
        # La destinazione è nuova: l'umano vede già questa chiamata. Un secondo
        # dialog sullo stesso invio è consent fatigue, non controllo in più.
        return None, ""
    if (egress_verdict.get("checked") and egress_verdict.get("allowed")
            and not egress_verdict.get("would_deny")):
        # Destinazione dichiarata: l'uscita è dentro il perimetro. Non è una
        # deduplicazione — è che non c'è nulla da chiedere. Chi ha censito quella
        # destinazione ha già deciso che scriverci è ammesso, e il canale
        # contaminato non cambia dove stanno andando i dati.
        return None, ""
    from . import taint as _t
    from .whitelist import current_chat
    chat = current_chat()
    st = _t.status(chat)
    if not st.get("tainted"):
        return None, ""
    ch = st["channel"]
    key = _t.context_gate_key(chat, _channel_participants(ch))
    srcs = ", ".join(f"{x.get('kind')}:{x.get('detail')}"
                     for x in (st.get("sources") or [])[-3:]) or "sorgente non registrata"
    reason = (f"@{agent} sta per usare {verb} da un canale CONTAMINATO ({ch}): "
              f"è entrato contenuto non fidato ({srcs}). Un'istruzione nascosta in "
              f"quel contenuto potrebbe essere ciò che sta chiedendo questa uscita. "
              f"Approvando, il canale viene declassificato: le uscite successive non "
              f"chiederanno più, finché non entra contenuto nuovo.")
    return key, reason


@app.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    try:
        # Enforcement whitelist per-richiesta: in HTTP multi-agente non basta
        # il filtro di list_tools (un client può invocare un tool non elencato).
        # I super-agent (clodia/ophelia) bypassano: accesso a tutti i tool. I
        # tool dei connettori (email.*, telegram.*) sono concessi anche a chi ha il
        # relativo grant nel vault (delega per-agent, persistente).
        _ag = agent_name()
        if is_on_behalf():
            # Richiesta ON-BEHALF di un umano: autorizza sul RUOLO umano (PDP
            # unico), NON sul carrier-agent. Un umano non-admin non può invocare
            # i tool super-only anche se il carrier è clodia (super).
            if not _human_tool_allowed(name):
                raise PermissionError(
                    f"tool '{name}' riservato agli admin (umano '{current_principal()}' "
                    f"ruolo '{current_human_role() or 'user'}')")
            if not _scoped_ceiling_ok(name):
                raise PermissionError(
                    f"tool '{name}' fuori dal token di '{current_principal()}': "
                    f"questo token concede {sorted(current_scoped_tools())}. "
                    f"Il ruolo non c'entra — è il token a essere ristretto.")
        elif not _is_super(_ag) and not _agent_tool_reachable(name, _ag):
            raise PermissionError(
                f"tool '{name}' non in whitelist per agent '{_ag}'")
        # DENY PER-AGENTE (#104 §8). Prima di tutto il resto: è una sottrazione
        # da `*`, e se un allow potesse sovrascriverla non toglierebbe nulla.
        # Vale anche per i super-agent, che è il punto — la lista esiste proprio
        # per ritagliare eccezioni al wildcard di clodia.
        if not is_on_behalf() and agent_denies(name, _ag):
            from . import observe as _obs
            if _obs.skipping():
                _obs.note("deny", name, _ag or "", detail="denied_tools")
            else:
                _hint = _DENY_HINT.get(name)
                raise PermissionError(
                    f"tool '{name}' non disponibile per l'agent '{_ag}'. "
                    + (_hint or "È stato escluso deliberatamente dalle sue "
                                "capacità: se ti serve, chiedi a Davide."))
        # BLOCCO DELLE SESSIONI NON PRESIDIATE (#104). Prima dei gate, di
        # proposito: negare non richiede il consenso di nessuno, e chiedere un
        # consenso che nessuno può dare è esattamente il difetto di #116.
        _un = _unattended_denial(name)
        if _un:
            from . import observe as _obs
            if _obs.skipping():
                _obs.note("deny", name, _ag or "", detail="unattended")
            else:
                raise PermissionError(_un)
        # INTERSEZIONE DELLA CATENA D'ORIGINE (docs/specification.md §3.3).
        # Prima dei gate, e non per efficienza: se la catena non regge, chiedere
        # l'approvazione del VERBO sarebbe la domanda sbagliata. «@messaggero
        # vuole spedire, approvi?» nasconde il fatto rilevante, che è «Giovanni
        # non ha questo permesso e sta usando messaggero per averlo».
        #
        # Intersezione, non sostituzione: far girare la chiamata sull'autorità di
        # chi ha iniziato rovescerebbe il difetto — Davide che chiede shell.exec a
        # un postino riuscirebbe. Entrambi devono permettere, e ogni anello.
        _org = origin.evaluate(_origin_chain(name), name)
        if _org.get("action") == "deny":
            _omode = origin.mode()
            if _omode == "on":
                raise PermissionError(origin.denial_message(_org))
            if _omode == "report":
                # Osservazione: decide e registra, non blocca. L'enforcement segue
                # la misura — è così che sono state trovate le lacune delle altre
                # due whitelist, e accenderlo alla cieca produrrebbe rifiuti su
                # lavoro legittimo e la perdita di fiducia nel controllo.
                from . import observe as _obs_o
                _obs_o.note("would_deny", name, _ag or "",
                            detail=f"origin:{_org.get('refused_by')}")
                LOG.info("origin (osservazione): %s rifiuterebbe %s — catena %s",
                         _org.get("refused_by"), name, " → ".join(_org["chain"]))

        # M-gate: conferma umana su azioni sotto supervisione — UN SOLO meccanismo.
        # (a) VERBI gated (packs/mcp/agents/…): consenso one-shot per-verbo.
        # (b) CROSS-TOPIC: un agente che tocca un topic di cui NON è participant
        #     richiede un consenso per-topic (time-boxed, così l'intera operazione
        #     cross-topic procede). Sostituisce il vecchio sudo cross-topic.
        # Il gate non concede nuovi tool: per il cross-topic apre soltanto il
        # compartimento target, sempre entro clearance e whitelist già verificate.
        gate_approval = None
        if not is_on_behalf():
            from . import gate as _gate
            # Gated GLOBALE (pericoloso per chiunque) oppure PER-AGENTE (§8: le
            # scritture di impiegato-tomato, quelle github di fullstack-dev —
            # stessi verbi che per altri restano liberi, quindi la granularità
            # non può essere globale).
            # Tre ragioni per un gate, e vanno distinte nel messaggio perché
            # chiedono all'umano di valutare cose diverse:
            #  - GLOBALE: il verbo è pericoloso per chiunque
            #  - PER-AGENTE: è pericoloso per QUESTO agente (le scritture di un
            #    factotum, le mutazioni github di uno sviluppatore)
            #  - FUORI PROFILO: l'agente può raggiungerlo ma non lo dichiara come
            #    proprio mestiere. Non è «pericoloso», è «inatteso da lui».
            _off_profile = outside_profile(name, _ag)
            # `gated_in_channel` era il QUARTO motivo ed è stato RITIRATO il
            # 7 ago 2026, dopo B2. Era il surrogato della domanda «chi sta
            # chiedendo?», posta per approssimazione — *qualcuno è in un canale* —
            # e l'approssimazione era grossolana due volte: una DM È un canale,
            # quindi chiedeva l'approvazione anche all'owner per la propria
            # richiesta nella propria DM; e non guardava affatto CHI avesse
            # chiesto, che è precisamente ciò che diceva di proteggere.
            #
            # Ora quella domanda ha una risposta esatta: la catena `origin` nomina
            # `human:giovanni` o `human:davide` e ne interseca il ruolo nella
            # stanza. Un surrogato che sopravvive alla cosa che surrogava diventa
            # un secondo controllo che dice altro — e due controlli sulla stessa
            # domanda divergono, come oggi hanno fatto tre volte.
            # UNA destinazione già ammessa non si fa approvare due volte.
            # `github.push` verso un repository che sta nella whitelist egress è
            # la stessa decisione presa due volte: il perimetro l'ha già detto sì
            # — quella è la porta — e il gate qui chiede di nuovo la stessa cosa
            # a ogni singola pubblicazione. Misurato su questo topic: 17 gate per
            # fullstack-dev, di cui otto per `create_branch`. Un agente che
            # pubblica di mestiere accumula conferme finché non si approva senza
            # leggere, e allora il gate non protegge più niente: resta l'attrito.
            #
            # Vale SOLO per i verbi che hanno una destinazione dichiarata e
            # confrontabile (`egress.spec_for`), e solo se TUTTE le destinazioni
            # combaciano. Fuori dal perimetro il gate resta — ed è lì che serve,
            # perché è lì che il confine si sposta.
            #
            # Vale per TUTTE E TRE le ragioni (globale, per-agente, fuori
            # profilo). Fino al 19 ago 2026 lo sconto si calcolava solo per il
            # gate globale, e le altre due lo rimettevano con un `or`: la
            # whitelist restava inefficace proprio dove l'attrito si misura —
            # `clodia` non dichiara `github.push` nel profilo, e un dev che abbia
            # ancora `github.*` nei `gated_tools` della copia del gateway (il seed
            # li ha rimossi il 17 ago) ricadeva nell'altro ramo. Card a ogni
            # pubblicazione verso un repository già approvato: clodia-platform#254.
            #
            # Quello che NON cambia: le tre ragioni restano distinte nel testo
            # della card, e il perimetro tace su chi non gli fa una domanda di
            # destinazione (`gate.perimeter_answers`) — `topic.remote_*`,
            # `agents.grant_*`, `egress.allow`, e `web.post`, la cui whitelist
            # censisce un host senza il path.
            _perimetro_ok = False
            _gated_globale = _gate.is_gated(name)
            _gated_agente = agent_gates(name, _ag)
            if ((_gated_globale or _gated_agente or _off_profile)
                    and _gate.perimeter_answers(name)):
                try:
                    from . import egress as _eg
                    # Stessa risoluzione che userà il PDP: due copie della
                    # stessa domanda divergono, e questa settimana è successo
                    # tre volte.
                    _args_dest = (_push_destination(arguments)
                                  if name == "github.push" else arguments)
                    _perimetro_ok = _eg.destinations_already_allowed(name, _args_dest)
                except Exception:  # noqa: BLE001 — in dubbio si chiede
                    _perimetro_ok = False
                if _perimetro_ok:
                    LOG.info("gate saltato per '%s' (@%s): destinazione già nel "
                             "perimetro autorizzato", name, _ag)
            if _gate.needs_consent(name, globally_gated=_gated_globale,
                                   agent_gated=_gated_agente,
                                   off_profile=_off_profile,
                                   perimeter_ok=_perimetro_ok):
                if name == "web.post":
                    reason = web_post.gate_summary(arguments)
                elif name in ("egress.allow", "ingress.allow"):
                    # Il verbo che ALLARGA un permesso è quello in cui il dialog
                    # deve dire cosa costa, non solo cosa concede.
                    from . import egress as _eg
                    _u = _eg.canonical(str(arguments.get("uri") or ""))
                    _dir = name.split(".", 1)[0]
                    _what = ("fonte fidata" if _dir == "ingress"
                             else "destinazione ammessa")
                    reason = (f"@{_ag} chiede di aggiungere {_u} come {_what}, "
                              f"per TUTTI gli agenti. {_eg.admin_note(_dir, _u)}")
                elif _off_profile:
                    reason = (f"@{_ag} chiede di usare `{name}`, che PUÒ raggiungere ma "
                              f"non dichiara nel proprio profilo. Non è un verbo "
                              f"pericoloso di per sé: è un verbo fuori dal suo mestiere. "
                              f"Se lo usa spesso, dichiararlo nel profilo del suo pack "
                              f"toglie questa domanda; approvarlo qui vale una volta.")
                else:
                    reason = ""
                gate_approval = await _require_gate_consent(
                    _ag, name, consume=True, reason=reason,
                    # Niente delega prefirmata sui verbi che aprono un'uscita o
                    # allargano un'autorità: una delega li renderebbe silenziosi
                    # per tutta la sua finestra, cioè l'opposto del motivo per cui
                    # il gate esiste.
                    allow_delegation=name not in {
                        "web.post", "agents.grant_scoped", "agents.revoke_scoped"},
                )
            _ck = _cross_topic_gate_key(name, arguments, _ag)
            if _ck:
                await _require_gate_consent(_ag, _ck, consume=False)
        # WHITELIST DI DESTINAZIONE (clodia-platform#104 §7, passo 5). Uscita da
        # capacità binaria a capacità circoscritta: «può inviare mail» diventa
        # «può inviare mail a queste destinazioni». Dopo il gate, non prima: se
        # il verbo è gated l'umano ha già visto la richiesta, e il verdetto sulla
        # destinazione è l'ultima cosa che deve poter fermare la chiamata.
        # Default `report`: decide e logga senza bloccare, perché la lista vera
        # delle destinazioni si impara dal traffico reale — è così che sono state
        # trovate le quattro lacune della whitelist di rete.
        # NON si applica alle richieste on-behalf: l'utente autenticato dall'UI
        # è trusted (§2), e una sua mail non è un'uscita dell'agente — è la sua.
        if not is_on_behalf():
            from . import egress as _egress
            try:
                _acfg = agent_config(_ag) if _ag else {}
            except KeyError:
                # Agent non in config (clone, connettore): nessuna regola
                # dichiarata → decide `egress.check`. Non si inventa un default
                # permissivo.
                _acfg = {}
            # `email.reply` non ha il destinatario negli argomenti: viene dal
            # messaggio a cui risponde, cioè da contenuto non fidato. Senza
            # risolverlo il verdetto sarebbe «destinazione ignota → nego», che
            # romperebbe il caso d'uso legittimo (un messaggero risponde alla
            # posta in arrivo: è il suo lavoro). Si risolve qui, al call-site,
            # con una lettura in più che avviene SOLO per questo verbo.
            _eargs = arguments
            if name == "email.reply" and not (arguments.get("to") or ""):
                _eargs = {**arguments, "to": _reply_recipient(arguments)}
            elif name == "github.push":
                # Senza questa riga il PDP vede un verbo con destinazione ignota
                # e nega: è ciò che è successo il 17 ago 2026 appena `push` è
                # entrato nel perimetro — reso governato e insieme impossibile.
                _eargs = _push_destination(arguments)
            # In una sessione non presidiata il modo `gate` non ha senso: la
            # richiesta resterebbe appesa fino al timeout. Si nega.
            from . import observe as _obs2
            _ev = _egress.check(_ag or "", _acfg, name, _eargs,
                                unattended=is_unattended())
            if _ev.get("action") == "deny":
                from . import observe as _obs
                if _obs.skipping():
                    _obs.note("deny", name, _ag or "",
                              detail=f"egress:{_ev.get('type') or '?'}")
                else:
                    raise _egress.denied_error(_ag or "", _ev)
            if _ev.get("action") == "gate":
                # Destinazione non vagliata → si CHIEDE, non si rifiuta. La
                # whitelist nasce vuota e si popola con l'uso (decisione del
                # 3 ago 2026): approvando, l'invio procede e la destinazione
                # resta ammessa. Il gate è sulla DESTINAZIONE, non sul verbo:
                # «scrivi a mario@x.it» è la domanda che l'umano sa valutare,
                # «puoi mandare mail» no.
                await _require_gate_consent(
                    _ag, _ev["gate_key"], consume=True,
                    reason=_ev.get("gate_reason", ""),
                    # Nessuna delega prefirmata su una destinazione nuova: è
                    # l'atto che la rende silenziosa per sempre (§7 proprietà 2),
                    # e una delega la renderebbe silenziosa senza che nessuno la
                    # veda mai.
                    allow_delegation=False)
                # In osservazione NON si ricorda: nessuno ha approvato, e una
                # whitelist che si popola da sé mentre i gate sono spenti sarebbe
                # una whitelist scritta dagli agenti. Al ritorno all'enforcement
                # ci si troverebbe tutto già consentito — cioè il contrario dello
                # scopo di questa modalità.
                if not _obs2.skipping():
                    _egress.remember(_ag or "", _ev["type"], _ev.get("remember") or [])
            # GATE DI CONTESTO (#104 §6): destinazione già ammessa, ma il canale
            # è contaminato. È il rischio che la whitelist non copre — l'uscita
            # verso una destinazione legittima di dati raccolti sotto injection.
            _ctx_key, _ctx_reason = _context_gate_needed(name, _ag or "", _ev)
            if _ctx_key:
                await _require_gate_consent(
                    _ag, _ctx_key, consume=False, reason=_ctx_reason,
                    # Nessuna delega: la condizione 2 di #77 esiste proprio per
                    # il caso in cui una delega copre il gate del verbo e
                    # nessuno guarda. Ammetterla qui riaprirebbe quel buco.
                    allow_delegation=False)
                # L'approvazione È l'«ultimo unlock» della definizione: il flag si
                # azzera, e si ri-arma da sé se entra contenuto nuovo.
                _taint.clear(current_chat(), by=current_principal() or "human")
        if name == "fs.list_dir":
            result = fs.list_dir(arguments["path"])
        elif name == "web.fetch":
            result = await asyncio.to_thread(web_fetch.fetch, arguments, agent=_ag or "")
        elif name == "web.post":
            result = await asyncio.to_thread(web_post.post, arguments, agent=_ag or "")
        elif name == "logs.tail":
            result = logs.tail(arguments.get("lines", 100), arguments.get("level", ""))
        elif name == "email.send":
            # ALLEGATI PER RIFERIMENTO (#104 §8, requisito derivato della riga
            # `messaggero`). Il gateway legge i file dal topic e li allega: il
            # contenuto NON entra nel contesto dell'agente, esattamente come per
            # i segreti. È ciò che permette di togliere a un postino i verbi di
            # lettura dei file del topic senza togliergli il mestiere — e non
            # brucia token su un PDF.
            _extra, _tmpdir = _topic_attachments(arguments, _ag or "")
            try:
                result = email.send(
                    arguments["to"],
                    arguments["subject"],
                    arguments["body"],
                    account=_email_account(arguments),
                    cc=arguments.get("cc"),
                    attachments=(arguments.get("attachments") or []) + _extra,
                )
            finally:
                if _tmpdir:
                    import shutil as _sh
                    _sh.rmtree(_tmpdir, ignore_errors=True)
        elif name == "email.folders":
            result = email.folders(account=_email_account(arguments))
        elif name == "email.list":
            result = email.list_messages(
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
                limit=arguments.get("limit", 10),
            )
        elif name == "email.read":
            result = email.read_message(
                arguments["email_id"],
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
            )
        elif name == "email.get_attachment":
            result = email.get_attachment(
                arguments["email_id"],
                arguments["filename"],
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
            )
        elif name == "email.save_attachment":
            # I byte dell'allegato NON passano dal modello: decodifica server-side
            # e scrittura o nello scratch validato, o DIRETTAMENTE nei file del
            # topic. La seconda è la RICEZIONE per riferimento, simmetrica
            # all'invio (`email.send(topic_files=…)`): esisteva il modo di
            # spedire un file senza vederlo, non quello di archiviarne uno.
            #
            # Serve perché il flusso documentato era `save_attachment` →
            # `topic.put`, e a un postino `topic.put` è negato (§8 di #104): un
            # allegato in arrivo non era archiviabile da chi la posta la riceve.
            _dest_arg = (arguments.get("dest") or "").strip()
            # Senza `dest` si archivia nel topic: quello indicato, o QUELLO IN CUI
            # SI STA LAVORANDO. Il default è ciò che rende il verbo usabile da un
            # agente che non ha modo di scoprire dove si trova.
            _tier, _tname = (None, None) if _dest_arg else _topic_of(arguments)
            if not (_dest_arg or (_tier and _tname)):
                raise ValueError(
                    "non sono riuscito a ricavare il topic dal canale: passa "
                    "`tier`+`name`, oppure `dest` per scrivere nel tuo scratch")
            raw, meta = email.get_attachment_bytes(
                arguments["email_id"],
                arguments["filename"],
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
            )
            if _tier and _tname:
                svc = _topics()
                _require_topic_member(svc, _tier, _tname)
                fn = arguments["filename"]
                # Provenienza `untrusted` d'ufficio: un file introdotto da un verbo
                # non ha nessuno da interrogare (#104 §3), e la posta in arrivo è
                # la definizione di sorgente non controllata.
                r = svc.put_file(_tier, _tname, fn, raw, "untrusted", by=_ag or "")
                # E CONTAMINA il canale: è contenuto di terzi che entra adesso.
                # Senza questo l'archiviazione sarebbe un modo di far entrare un
                # PDF ostile senza lasciare traccia nel flag.
                _taint.mark(f"{_tier}/{_tname}", "file", fn, _ag or "")
                LOG.info("email: allegato '%s' archiviato in %s/%s (untrusted)",
                         fn, _tier, _tname)
                result = {"topic_path": r["path"], "provenance": r["provenance"],
                          "size": len(raw), "tainted": True, **meta}
            else:
                dest = _safe_scratch_path(_dest_arg)
                _os.makedirs(_os.path.dirname(dest), exist_ok=True)
                with open(dest, "wb") as f:
                    f.write(raw)
                result = {"local_path": dest, "size": len(raw), **meta}
        elif name == "email.search":
            result = email.search(
                arguments["query"],
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
                limit=arguments.get("limit", 20),
            )
        elif name == "email.reply":
            result = email.reply(
                arguments["email_id"],
                arguments["body"],
                account=_email_account(arguments),
                folder=arguments.get("folder", "INBOX"),
                cc=arguments.get("cc"),
                attachments=arguments.get("attachments"),
            )
        elif name.startswith("topic."):
            # offload su thread: _dispatch_topic tocca lo storage e (suggest_team/
            # participants) fa httpx SINCRONO all'agent-server. Se girasse
            # nell'event loop lo bloccherebbe → deadlock bilaterale con
            # topics_client (agent-server → gateway, anch'esso sync). to_thread
            # propaga i contextvars → agent_name() resta valido.
            result = await asyncio.to_thread(_dispatch_topic, name, arguments)
        elif name.startswith("image."):
            result = await _dispatch_image(arguments)
        elif name.startswith("artifact."):
            result = _dispatch_artifact(arguments)
        elif name.startswith("profile."):
            result = _dispatch_profile(name, arguments, _ag)
        elif name.startswith("settings."):
            result = _dispatch_settings(name, arguments, _ag)
        elif name.startswith("telegram."):
            result = _dispatch_telegram(name, arguments)
        elif name.startswith("memory."):
            result = _dispatch_memory(name, arguments)
        elif name in _GITHUB_NATIVE_NAMES:
            # NON `startswith("github.")`: il namespace è condiviso col backend
            # MCP montato con lo stesso nome (vedi `_GITHUB_NATIVE_NAMES`). Solo i
            # verbi che il gateway implementa passano di qui; gli altri `github.*`
            # scendono fino al ramo `proxy.is_proxied`, che è il loro.
            result = await asyncio.to_thread(_dispatch_github, name, arguments)
        elif name.startswith("gdrive."):
            result = _dispatch_gdrive(name, arguments)
        elif name.startswith("gcalendar."):
            result = _dispatch_gcalendar(name, arguments)
        elif name.startswith("gdocs."):
            result = _dispatch_gdocs(name, arguments)
        elif name.startswith("gsheets."):
            result = _dispatch_gsheets(name, arguments)
        elif name.startswith(("egress.", "ingress.")):
            result = _dispatch_egress_admin(name, arguments)
        elif name.startswith("runtime."):
            # proxy httpx SINCRONO all'agent-server → offload su thread (no blocco loop)
            result = await asyncio.to_thread(_dispatch_runtime, name, arguments, _ag)
        elif name.startswith("jobs."):
            result = await asyncio.to_thread(_dispatch_jobs, name, arguments, _ag)
        elif name.startswith("packs."):
            result = await asyncio.to_thread(_dispatch_packs, name, arguments)
        elif name.startswith("providers."):
            result = await asyncio.to_thread(_dispatch_providers, name, arguments)
        elif name.startswith("integrations."):
            result = await asyncio.to_thread(_dispatch_integrations, name, arguments)
        elif name.startswith("mcp."):
            result = await asyncio.to_thread(_dispatch_mcp, name, arguments)
        elif name.startswith("agents."):
            result = await asyncio.to_thread(
                _dispatch_agents, name, arguments, _ag, gate_approval)
        elif name == "eu_corpus.search":
            # alias morbido: eu_corpus.* == rag.* sulla collection eu-normativa.
            _rag_authorize("eu-normativa", write=False)
            result = eu_corpus.search(
                arguments["query"],
                k=arguments.get("k", 5),
                doc=arguments.get("doc"),
            )
        elif name == "eu_corpus.ingest":
            # Legge il PDF dal topic server-side (i byte NON passano dal modello),
            # con controllo participant+clearance, poi lo invia al micro-servizio.
            _rag_authorize("eu-normativa", write=True)
            svc = _topics()
            tier, tname, path = arguments["tier"], arguments["name"], arguments["path"]
            _require_topic_member(svc, tier, tname)
            data = svc.read_file(tier, tname, path)
            filename = path.rsplit("/", 1)[-1]
            result = eu_corpus.ingest_bytes(
                data, filename,
                arguments["doc_name"], arguments["version"],
                url=arguments.get("url"),
                supersede=bool(arguments.get("supersede", False)),
            )
        elif name == "eu_corpus.list":
            _rag_authorize("eu-normativa", write=False)
            result = eu_corpus.list_documents()
        elif name == "eu_corpus.remove":
            _rag_authorize("eu-normativa", write=True)
            result = eu_corpus.remove(arguments["doc_name"], arguments.get("version"))
        elif name.startswith("rag."):
            result = _dispatch_rag(name, arguments)
        elif proxy.is_proxied(name):
            # C1: instrada al backend MCP montato (già passato il check whitelist).
            text = await proxy.call_proxied(name, arguments)
            # I verbi GitHub e gli MCP esterni passano DA QUI: sono la sorgente di
            # contenuto di terzi più ovvia, e marcare solo il ritorno nativo
            # avrebbe lasciato scoperto proprio il vettore del caso Invariant Labs.
            _taint.note_verb(name, _ag or "",
                              vetted=_source_vetted(name, arguments))
            _tlm.record(name, _ag or "", "ok", channel=current_chat(),
                        unattended=is_unattended())
            return [TextContent(type="text", text=text)]
        else:
            raise ValueError(f"unknown tool: {name}")
        # CONTAMINAZIONE (#104 §4, passo 7). Dopo l'esecuzione e solo in caso di
        # successo: il taint nasce quando il contenuto di terzi è EFFETTIVAMENTE
        # entrato nel contesto, non quando è stato chiesto. La §4 riformula la
        # colonna `untrusted_input` del catalogo da «questo agente è esposto» a
        # «questo verbo produce taint», ed è questo il punto in cui accade.
        _taint.note_verb(name, _ag or "",
                          vetted=_source_vetted(name, arguments, result))
        _tlm.record(name, _ag or "", "ok", channel=current_chat(),
                    unattended=is_unattended())
        return [TextContent(type="text", text=json.dumps(result, indent=2, ensure_ascii=False))]
    except PermissionError as e:
        # Il motivo è una CLASSE, non il messaggio: i messaggi contengono nomi di
        # file e indirizzi, e questo registro non deve diventare una rubrica.
        _why = ("egress" if "uscita non consentita" in str(e)
                else "unattended" if "job schedulato" in str(e)
                else "denied_tools" if "denied_tools" in str(e)
                else "whitelist" if "non in whitelist" in str(e)
                else "clearance" if "clearance" in str(e).lower()
                else "other")
        _tlm.record(name, agent_name_safe(), "denied", channel=current_chat(),
                    unattended=is_unattended(), detail=_why)
        return [TextContent(type="text", text=f"DENIED: {e}")]
    except VersionConflict as e:
        return [TextContent(type="text", text=(
            "CONFLICT: il summary è cambiato durante il lavoro — rileggi con "
            f"topic.open e riapplica le tue modifiche, non sovrascrivere. {e}"))]
    except TopicError as e:
        return [TextContent(type="text", text=f"ERROR: {e}")]
    except Exception as e:
        return [TextContent(type="text", text=f"ERROR: {type(e).__name__}: {e}")]


# Estensioni che indicano contenuto BINARIO: un file con questo suffisso non può
# essere scritto come testo (es. un .xlsx scritto come testo = file corrotto che
# Excel non apre). Per questi forziamo la decodifica base64.
_BINARY_EXTS = {
    "xlsx", "xls", "xlsm", "docx", "doc", "pptx", "ppt", "pdf", "odt", "ods", "odp",
    "zip", "tar", "gz", "tgz", "7z", "rar",
    "png", "jpg", "jpeg", "gif", "webp", "bmp", "tiff", "tif", "ico", "heic",
    "mp3", "wav", "ogg", "mp4", "mov", "avi", "mkv", "webm",
    "bin", "exe", "woff", "woff2", "ttf", "otf",
}

# Soglia oltre la quale i byte NON devono viaggiare come base64 nei parametri di
# una tool-call (ARG_MAX, troncamento, token bruciati): sopra questo, read_file/
# write_file rifiutano e indirizzano a topic.fetch/topic.put (transfer via scratch,
# mediato dal gateway). ~128KB grezzi ≈ ~170KB di base64.
_B64_INLINE_CAP = 128 * 1024


def _decode_b64_strict(content: str, filename: str) -> bytes:
    """Decodifica base64 in modo robusto: tollera whitespace/newline e padding
    mancante (errori comuni quando un LLM passa un blob lungo), ma su input non
    valido solleva un errore CHIARO — così l'agente rigenera il base64 invece di
    far scrivere spazzatura senza accorgersene."""
    import base64 as _b64
    import binascii
    t = "".join((content or "").split())          # togli spazi/newline
    t += "=" * ((-len(t)) % 4)                      # ripristina padding mancante
    try:
        return _b64.b64decode(t, validate=True)
    except (binascii.Error, ValueError) as e:
        raise ValueError(
            f"Il content per '{filename}' non è base64 valido ({e}). I file binari "
            f"(xlsx/pdf/docx/zip/immagini) vanno passati come base64 con "
            f"encoding='base64'; rigenera il base64 COMPLETO del file e riprova."
        ) from e


_SPAWNS_ROOT = _os.environ.get("CLODIA_SPAWNS_ROOT", "/datadir/spawns")


def _safe_scratch_path(p: str) -> str:
    """Valida che `p` stia DENTRO lo scratch di uno spawn, non nel cortile.

    Il gateway scrive/legge solo qui, mai nel topic store o nei secrets, anche se
    l'agent passa un path arbitrario. Difesa contro path-traversal/abuso.

    La radice `/datadir/spawns` NON è una destinazione. Accettarla — la versione
    precedente lo faceva con `rp == root` e con uno `startswith` che non chiedeva
    un livello — ha prodotto su marte **226 documenti sciolti in cima al
    cortile**: PDF, DOCX, lettere su carta intestata, root:root e modo 644.

    Perché contava, misurato: la cartella è `drwx--x--x`, quindi un agente non
    può ELENCARLI (`ls` → Permission denied), ma può traversarla, e 644 significa
    che chi ne conosce il nome li legge. Un file finito lì esce dal perimetro del
    suo scope senza che nulla lo dica, ed è il modo silenzioso di sbagliare:
    l'operazione riesce.

    **Il confinamento per spawn** (specification §1.1) è la seconda metà, e dal
    7 ago 2026 c'è: il token porta `execution_id`, quindi il gateway sa non solo
    QUALE SEED chiama ma QUALE SPAWN. Prima conosceva solo il seed, e la voce 2
    — «uno spawn possiede il proprio scratch e non raggiunge quello di un
    altro» — era promessa dalla specifica e non mantenuta dal codice.

    Se il claim manca — un chiamante vecchio, un percorso interno — resta il solo
    controllo del livello. Rifiutare lì trasformerebbe l'assenza di un campo
    nuovo in un guasto, e la direzione della retrocompatibilità va verso «come
    prima».
    """
    rp = _os.path.realpath(p or "")
    root = _os.path.realpath(_SPAWNS_ROOT)
    if not rp.startswith(root + "/"):
        raise ValueError(
            f"path non consentito (deve stare sotto {_SPAWNS_ROOT}): {p}. "
            f"Se stai archiviando un allegato di posta non serve un file locale: "
            f"email.save_attachment(email_id, filename) lo scrive nel topic del "
            f"canale, e i byte non passano da te.")
    resto = rp[len(root) + 1:]
    if "/" not in resto:
        raise ValueError(
            f"path non consentito: '{p}' finirebbe nella RADICE del cortile "
            f"degli spawn, non nello scratch di uno spawn. Usa un path dentro la "
            f"tua directory di lavoro (es. `{_SPAWNS_ROOT}/<spawn>/{resto}`), "
            f"oppure — se stai archiviando un allegato — `email.save_attachment` "
            f"senza `dest`, che lo scrive nel topic del canale.")
    from .whitelist import current_spawn
    mio = current_spawn()
    suo = resto.split("/", 1)[0]
    if mio and suo != mio:
        raise ValueError(
            f"path non consentito: '{p}' sta nello scratch di '{suo}', e tu sei "
            f"'{mio}'. Lo scratch di uno spawn è suo — scrivi sotto "
            f"{_SPAWNS_ROOT}/{mio}/. Se devi passare un file a un altro spawn, "
            f"la strada è il topic del canale, non il suo filesystem.")
    return rp


# Verbi topic.* che accedono ai dati di UN topic specifico → richiedono che il
# caller sia participant/owner (compartimento, need-to-know). `new`/`list`/`search`
# sono gestiti a parte (creazione / risultati filtrati per membership).
_TOPIC_SCOPED_VERBS = {
    "open", "save_summary", "save_agents_md", "add_minute", "archive", "set_portable",
    "telegram_bind", "telegram_unbind",
    "files", "read_file",
    "read_document", "write_file", "fetch", "put", "delete_file", "migrate_storage",
    "post_message", "messages", "my_mentions", "mark_seen",
    "remote_enable", "remote_disable", "remote_add", "remote_commit",
    "remote_push", "remote_pull", "remote_status",
}


_SEAL_RANK = {"SEAL-0": 0, "SEAL-1": 1, "SEAL-2": 2, "SEAL-3": 3, "SEAL-4": 4}


def _rank(tier: str | None) -> int:
    return _SEAL_RANK.get(str(tier or "SEAL-0").strip().upper(), 0)


def _topic_is_member(meta: dict, caller: str) -> bool:
    return caller == meta.get("owner") or caller in (meta.get("participants") or [])


#: Verbi `topic.*` che MUTANO lo stato dello scope. Il resto legge.
#:
#: `post_message` non c'è, di proposito: parlare non è mutare. Un reader — umano
#: o agente — resta nella stanza per seguirne il lavoro e dire la sua, e
#: azzittirlo sarebbe una cosa diversa da quella che il ruolo descrive.
#:
#: Non ci sono nemmeno `add_participant`/`remove_participant`: non passano da
#: `_TOPIC_SCOPED_VERBS`, quindi classificarli qui sarebbe una regola che non si
#: applica mai e che sembra applicarsi. Sono GATED, che è un controllo più
#: stretto del ruolo. L'ha colto un test scritto apposta — «una classificazione
#: su un verbo che non esiste è una regola che non si applica mai».
_TOPIC_MUTATING_VERBS = frozenset({
    "save_summary", "save_agents_md", "add_minute", "archive", "set_portable",
    "telegram_bind", "telegram_unbind",
    "write_file", "put", "delete_file", "migrate_storage",
    "remote_enable", "remote_disable", "remote_add", "remote_commit",
    "remote_push", "remote_pull",
})


def _chat_binds_this_topic(tier_t: str, name: str) -> bool:
    """True se il claim `chat` del token lega la sessione a QUESTO topic.

    Il claim vale `chan:<tier>:<topic>:<chi>` ed è firmato: chi lo porta non può
    riscriverlo per affacciarsi su un'altra stanza. È questa proprietà — non il
    nome della persona — a rendere sicuro il ramo umano di `_require_topic_member`.
    """
    c = current_chat() or ""
    if not c.startswith("chan:"):
        return False
    p = c[len("chan:"):].split(":")
    return len(p) >= 2 and p[0] == str(tier_t) and p[1] == str(name)


def _token_is_bound_to_a_room() -> bool:
    """True se questo token dichiara UNA stanza — cioè è il token di un client
    umano, non una sessione della webui."""
    return is_on_behalf() and (current_chat() or "").startswith("chan:")


def _require_person_of_this_room(meta, tier, name, tier_t, mutating: bool) -> None:
    """Gli stessi tre assi di un agente, letti sulla PERSONA: compartimento
    (partecipa?), ruolo (reader non muta), livello (clearance ≥ tier)."""
    chi = current_principal() or ""
    if not _topic_is_member(meta, chi):
        raise PermissionError(
            f"accesso negato al topic {tier}/{name}: '{chi or 'ignoto'}' non è "
            "partecipante. Il token resta valido: è la stanza che non è sua. "
            "Chi ti ha invitato può aggiungerti dai partecipanti del topic.")
    if mutating:
        from .topics.service import TopicService as _T
        if _T.participant_role(meta, chi) == _T.ROLE_READER:
            raise PermissionError(
                f"'{chi}' è reader in {tier_t}/{name}: può leggere e parlare, "
                "non modificare.")
    if _rank(current_clearance()) < _rank(tier_t):
        raise PermissionError(
            f"'{chi}': clearance insufficiente per il tier {tier_t} di "
            f"{tier}/{name} (accesso negato: livello)")


def _require_topic_member(svc, tier, name, mutating: bool = False) -> None:
    """ACL compartimento (need-to-know).

    Consentito solo se l'AGENTE è participant/owner del target oppure esiste un
    consenso M-gate cross-topic attivo. La membership del principal umano non è
    un bypass: altrimenti un owner di molti topic annullerebbe il compartimento
    per qualunque agente che opera per suo conto.
    """
    caller = agent_name()
    try:
        meta = svc.open(tier, name).get("meta", {})
    except Exception:  # noqa: BLE001 — topic inesistente/illeggibile → nega
        raise PermissionError(f"topic {tier}/{name}: accesso negato")
    tier_t = meta.get("tier", tier)
    # CLIENT MCP DI UNA PERSONA. Un token legato a UNA stanza (`chat` firmato) e
    # coniato on-behalf non porta un agente che lavora per conto di qualcuno:
    # porta la persona. Chiedere la membership del carrier chiuderebbe l'accesso
    # a Giovanni perché l'agente che gli firma il token non partecipa — un
    # rifiuto che non riguarda nessuno dei due.
    #
    # Non contraddice la regola sopra, che resta: il principal NON è un bypass
    # generale. Qui il perimetro non lo dà il principal, lo dà il `chat` — vale
    # per QUESTO topic e per nessun altro, e chi lo porta non può cambiarlo
    # perché è firmato. Un token del genere non apre altre stanze: le apre di
    # meno, non di più.
    if _token_is_bound_to_a_room():
        # E se la stanza NON è quella, si finisce qui: **rifiuto**, non ripiego
        # sul ramo dell'agente. Il primo disegno cadeva sul ramo del carrier, e
        # in esercizio il carrier è `clodia`, che partecipa a tutto: il token
        # «legato a una stanza» apriva ogni altra stanza, e lo faceva
        # rispondendo `200`. Il confinamento sembrava esserci perché il caso
        # felice funzionava — trovato usando il token per davvero, non nei test,
        # che chiedevano al ramo giusto se faceva la cosa giusta e mai all'altro
        # se stava zitto.
        if not _chat_binds_this_topic(tier_t, name):
            raise PermissionError(
                f"questo token vale per {current_chat() or '—'} e non per "
                f"{tier}/{name}. Un client MCP è collegato a UNA stanza: per "
                "un altro topic si conia un altro collegamento.")
        _require_person_of_this_room(meta, tier, name, tier_t, mutating)
        return
    agent_ok = _topic_is_member(meta, caller)
    # cross-topic: consentito con un CONSENSO GATE attivo per questo topic
    # (topic-access:<tier>/<name>), concesso via popup (M-gate). Sostituisce sudo.
    from . import gate as _gate
    cross_ok = _gate.active(caller, "-", f"topic-access:{tier_t}/{name}")
    if not (agent_ok or cross_ok):
        raise PermissionError(
            f"accesso negato al topic {tier}/{name}: l'agente '{caller}' non è "
            "partecipante (compartimento need-to-know; "
            f"il cross-topic richiede un consenso gate)")
    # asse RUOLO: un reader non muta. Fino al 7 ago 2026 il ruolo era applicato
    # solo sul percorso UMANO (gli endpoint della webui); un agente passa da qui,
    # quindi metterlo a `reader` non aveva alcun effetto — poteva comunque
    # chiamare `topic.put` o `save_summary`. Il ruolo esisteva e nessuno lo
    # guardava dove serviva di più: un agente reader è precisamente il caso
    # dell'osservatore che deve poter guardare e commentare senza toccare.
    if mutating and agent_ok:
        from .topics.service import TopicService as _T
        if _T.participant_role(meta, caller) == _T.ROLE_READER:
            raise PermissionError(
                f"'{caller}' è reader in {tier_t}/{name}: può leggere e parlare, "
                "non modificare. Chiedi all'owner del topic di cambiargli ruolo. "
                "(Il consenso puntuale per una singola modifica arriverà col "
                "modello a gate: oggi è un rifiuto.)")
    # asse livello: clearance ≥ tier (difesa in profondità oltre al compartimento).
    if _rank(current_clearance()) < _rank(tier_t):
        raise PermissionError(
            f"agent '{caller}': clearance insufficiente per il tier {tier_t} del "
            f"topic {tier}/{name} (accesso negato: livello)")


def _require_local_hook_caller(svc, tier, name) -> str:
    """ACL stretta dell'hook locale: niente principal umano né gate cross-topic."""
    caller = agent_name()
    if not caller:
        raise PermissionError("invocazione hook locale riservata agli agenti")
    if caller == "messaggero":
        return caller
    try:
        meta = svc.open(tier, name).get("meta", {})
    except Exception:  # noqa: BLE001
        raise PermissionError(f"topic {tier}/{name}: accesso negato")
    if not _topic_is_member(meta, caller):
        raise PermissionError(
            f"agent '{caller}' non è participant di {tier}/{name}")
    if _rank(current_clearance()) < _rank(meta.get("tier", tier)):
        raise PermissionError(
            f"agent '{caller}': clearance insufficiente per {tier}/{name}")
    return caller


def _filter_member_rows(rows: list, caller: str) -> list:
    """Filtra allo scope need-to-know dell'AGENTE, su ENTRAMBI gli assi.

    La membership umana non amplia l'elenco: l'accesso aggiuntivo è per-topic e
    passa dal gate, non diventa un lasciapassare globale.

    FAIL CLOSED. Una riga che non si può valutare viene ESCLUSA. Prima passava
    invariata, e quel ramo difensivo era l'unico percorso vivo: `search` non
    restituiva `participants` né `owner`, quindi il filtro non filtrava niente —
    un agente riceveva titolo e tldr di ogni topic corrispondente, compartimento
    e tier ignorati. Misurato: 97 righe a `segretario`, 27 delle quali SEAL-2.
    Il tldr è la prima riga del summary, cioè la riga più informativa di un
    dossier: era la peggior cosa da esporre.

    Un default che ammette ciò che non sa valutare non è difensivo: è una porta
    aperta con un commento rassicurante sopra.

    Si controlla anche il LIVELLO, non solo il compartimento: le due condizioni
    sono la stessa regola applicata su `open` e `read_file`, e un elenco che le
    tratta più larghe è un modo di leggere ciò che non si potrebbe aprire.
    """
    out = []
    my_rank = _rank(current_clearance())
    for r in rows:
        if not isinstance(r, dict):
            continue
        if "participants" not in r and "owner" not in r:
            # Nessun rumore per riga: un log per chiamata basta a scoprire un
            # cambio di forma, e un log per riga renderebbe illeggibile il resto.
            import logging as _lg
            _lg.getLogger("clodia-tools").warning(
                "filtro need-to-know: riga senza participants/owner, esclusa "
                "(forma inattesa: %s)", sorted(r)[:6])
            continue
        if not _topic_is_member(r, caller):
            continue
        tier_r = r.get("tier")
        if tier_r and my_rank < _rank(tier_r):
            continue
        out.append(r)
    return out


def _rag_grants(agent: str) -> dict[str, set[str]]:
    """Grant live dal core; qualunque errore nega l'accesso (fail closed)."""
    try:
        return runtime.rag_grants(agent)
    except PermissionError:
        raise
    except Exception as e:  # noqa: BLE001 — backend irraggiungibile/malformato
        raise PermissionError(
            f"impossibile verificare i grant RAG dell'agent '{agent}'") from e


def _rag_readable(grants: dict[str, set[str]]) -> set[str]:
    """Collection su cui l'agent ha lettura (read grant OR write grant)."""
    return set(grants.get("rag_read") or []) | set(grants.get("rag_write") or [])


def _rag_provisioners() -> set[str]:
    return {
        x.strip() for x in _os.environ.get("CLODIA_RAG_PROVISIONERS", "sysadmin").split(",")
        if x.strip()
    }


def _rag_authorize(collection: str, write: bool) -> None:
    """Reference monitor per-collection: grant read/write (arg-aware, dal
    AgentSpec autorevole nel core) + tiering (clearance ≥ tier della collection).
    Super-agent → bypass dei grant, MA il vincolo del profilo (rag off/single)
    è strutturale e vale per tutti. Solleva PermissionError su violazione."""
    instance_profile.rag_check_collection(collection)
    ag = agent_name()
    if _is_super(ag):
        return
    if write and ag in _rag_provisioners():
        tier = eu_corpus.collection_tier(collection)
        if _rank(current_clearance()) < _rank(tier):
            raise PermissionError(
                f"agent '{ag}': clearance insufficiente per la collection '{collection}' "
                f"(tier {tier})")
        return
    grants = _rag_grants(ag)
    if write:
        if collection not in grants["rag_write"]:
            raise PermissionError(
                f"agent '{ag}' senza grant di SCRITTURA sulla collection '{collection}'")
    else:
        if collection not in _rag_readable(grants):
            raise PermissionError(
                f"agent '{ag}' senza grant di LETTURA sulla collection '{collection}'")
    # asse livello: clearance(agent) ≥ tier(collection). Difesa in profondità.
    tier = eu_corpus.collection_tier(collection)
    if _rank(current_clearance()) < _rank(tier):
        raise PermissionError(
            f"agent '{ag}': clearance insufficiente per la collection '{collection}' "
            f"(tier {tier})")


def _dispatch_rag(name: str, a: dict):
    verb = name.split(".", 1)[1]
    if not instance_profile.rag_enabled():
        raise PermissionError("feature 'rag' disabilitata dal profilo dell'istanza")
    if verb == "collections":
        res = eu_corpus.collections()
        # Profilo rag:single → la lista mostra solo la collection dell'edizione.
        if instance_profile.rag_mode() == "single":
            only = instance_profile.load()["rag"].get("collection") or ""
            res = {"collections": [c for c in res.get("collections", [])
                                   if c.get("collection") == only]}
        if not _is_super(agent_name()):
            if agent_name() in _rag_provisioners():
                res = {"collections": [
                    c for c in res.get("collections", [])
                    if _rank(current_clearance()) >= _rank(c.get("tier", "SEAL-0"))
                ]}
            else:
                allowed = _rag_readable(_rag_grants(agent_name()))
                res = {"collections": [c for c in res.get("collections", [])
                                       if c.get("collection") in allowed]}
        return res
    if verb == "create_collection":
        collection = a["collection"]
        _rag_authorize(collection, write=True)
        tier = a.get("tier", "SEAL-1")
        if _rank(current_clearance()) < _rank(tier):
            raise PermissionError(
                f"agent '{agent_name()}': clearance insufficiente per creare "
                f"collection tier {tier}")
        return eu_corpus.create_collection(
            collection,
            tier=tier,
            description=a.get("description"),
        )
    collection = a["collection"]
    if verb == "search":
        _rag_authorize(collection, write=False)
        return eu_corpus.search(a["query"], k=a.get("k", 5), doc=a.get("doc"),
                                collection=collection)
    if verb == "list":
        _rag_authorize(collection, write=False)
        return eu_corpus.list_documents(collection)
    if verb == "ingest":
        _rag_authorize(collection, write=True)
        svc = _topics()
        tier, tname, path = a["tier"], a["name"], a["path"]
        _require_topic_member(svc, tier, tname)
        data = svc.read_file(tier, tname, path)
        filename = path.rsplit("/", 1)[-1]
        return eu_corpus.ingest_bytes(
            data, filename, a["doc_name"], a["version"],
            url=a.get("url"), supersede=bool(a.get("supersede", False)),
            collection=collection)
    if verb == "remove":
        _rag_authorize(collection, write=True)
        return eu_corpus.remove(a["doc_name"], a.get("version"), collection)
    raise ValueError(f"unknown rag verb: {name}")


async def _dispatch_image(a: dict):
    """image.generate → genera un PNG e lo salva nei files/ del topic.
    La API key OpenAI è letta server-side dal vault, mai esposta all'agente."""
    from .tools import image as image_tool
    if not image_tool.has_key():
        return {"ok": False, "error": "nessuna API key OpenAI nel vault "
                "(Tools → Image generation)."}
    svc = _topics()
    tier, name = a.get("tier"), a.get("name")
    _require_topic_member(svc, tier, name)
    prompt = (a.get("prompt") or "").strip()
    if not prompt:
        return {"ok": False, "error": "serve un prompt"}
    filename = (a.get("filename") or "image.png").strip().lstrip("/")
    if not filename.lower().endswith(".png"):
        filename += ".png"
    png = await asyncio.to_thread(
        image_tool.generate, prompt,
        size=a.get("size") or "1024x1024",
        quality=a.get("quality") or "auto",
        background=a.get("background") or "auto")
    svc.put_file(tier, name, filename, png)
    return {"ok": True, "path": f"files/{filename}", "bytes": len(png)}


def _dispatch_artifact(a: dict):
    """artifact.render → snapshot del canvas live in files/artifact.html del topic
    (persistente; la finestra di anteprima lo mostra col suo polling)."""
    svc = _topics()
    tier, name = a.get("tier"), a.get("name")
    _require_topic_member(svc, tier, name)
    data = (a.get("html") or "").encode("utf-8")
    svc.put_file(tier, name, "artifact.html", data)
    return {"ok": True, "path": "files/artifact.html", "bytes": len(data)}


def _extract_document_text(filename: str, data: bytes) -> tuple[str, int | None]:
    """Estrae testo da PDF/DOCX/XLSX (server-side). Ritorna (testo, n_pagine|None).
    Fallback: prova a decodificare come testo UTF-8."""
    import io
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = [(p.extract_text() or "") for p in reader.pages]
        return "\n\n".join(pages), len(pages)
    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts), None
    if ext in ("xlsx", "xlsm"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"# Foglio: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                out.append("\t".join("" if v is None else str(v) for v in row))
        return "\n".join(out), None
    # fallback testo
    return data.decode("utf-8", errors="replace"), None


def _dispatch_topic(name: str, a: dict):
    svc = _topics()
    verb = name.split(".", 1)[1]
    if verb in _TOPIC_SCOPED_VERBS:
        _require_topic_member(svc, a.get("tier"), a.get("name"),
                              mutating=verb in _TOPIC_MUTATING_VERBS)
    if verb == "suggest_team":
        # proposta di squadra: proxy read-only all'agent-server (registry+rilevanza)
        return runtime.suggest_team(a.get("tier") or "SEAL-0", a.get("description") or "")
    if verb in ("add_participant", "remove_participant"):
        # topic.add_participant/remove_participant sono verbi GATED → il
        # gate di call_tool ha già richiesto la conferma umana (block-and-wait)
        # prima di arrivare qui. L'owner umano usa invece l'endpoint webui
        # dedicato, già autorizzato sul suo ruolo.
        return runtime.set_participant(a["tier"], a["name"], (a.get("agent") or "").strip(),
                                       by=agent_name() or "", add=(verb == "add_participant"))
    if verb == "new":
        # Profilo topics:single → solo il workspace unico (DM sempre permessi).
        instance_profile.topic_creation_check(a["name"])
        hook_enabled = bool(a.get("hook_enabled", True))
        meta = svc.new(
            a.get("tier"), a["name"],
            {**(a.get("meta") or {}), "hook_enabled": hook_enabled})
        if hook_enabled:
            runtime.ensure_topic_hook(
                meta["tier"], a["name"], by=agent_name() or "platform")
        return meta
    if verb == "invoke_hook":
        caller = _require_local_hook_caller(svc, a["tier"], a["name"])
        return runtime.invoke_topic_hook(
            a["tier"], a["name"], a.get("payload") or "", caller=caller)
    if verb == "open":
        return svc.open(a["tier"], a["name"])
    if verb == "post_message":
        # Prerogativa di messaggero e dei super: posta una bolla nella chat del
        # topic (es. mail in arrivo / handoff). Se il testo ha una @menzione,
        # innesca il risponditore → l'agente taggato prende in carico il messaggio.
        ag = agent_name()
        # Nessun controllo sul nome qui: il dispatch ha GIÀ verificato che
        # `topic.post_message` sia fra i verbi di questo agente. Prima c'era
        # `if not (_is_super(ag) or ag == "messaggero"): raise` — un elenco di
        # nomi che produceva due errori opposti, misurati su entrambe le istanze:
        #
        #   sysadmin   dichiara=True   ammesso=False   → dichiarava e restava muto
        #   clodia     dichiara=False  ammesso=True    → parlava senza dichiararlo
        #
        # Cioè la dichiarazione non contava in nessuna delle due direzioni. E un
        # elenco di nomi va aggiornato a mano ogni volta che nasce un agente che
        # deve rispondere in chat: `sysadmin` è entrato nei canali con la modalità
        # debug e il suo diritto di parlare è rimasto indietro.
        #
        # Rimosso invece che riscritto: rifarlo qui duplicherebbe la whitelist in
        # un secondo punto — il «doppio gate incoerente» di cui avverte
        # `whitelist.tool_allowed` — e due copie della stessa regola divergono.
        text = a.get("text") or ""
        # CHI HA PARLATO. Firmare sempre col carrier-agent era corretto finché a
        # chiamare c'erano solo agenti. Da un client MCP di una persona la stessa
        # riga farebbe dire alla chat una cosa falsa: il messaggio di Giovanni
        # comparirebbe a nome dell'agente che porta il suo token.
        #
        # `kind` non è un dettaglio cosmetico che segue l'autore: è il campo su
        # cui poggiano due regole già scritte — «una menzione a una persona non
        # instrada un'AI» e «solo i messaggi umani accodano una notifica
        # Telegram». Sbagliarlo qui le disattiverebbe entrambe in silenzio, che è
        # il modo in cui una regola smette di valere senza che nessuno la tocchi.
        #
        # `on_behalf` NON basta a dire «una persona»: un proxy — un sistema terzo
        # ammesso nella stanza — parla per conto di un principal ammesso, quindi
        # è on-behalf come Giovanni, e finché `kind` si derivava da qui il
        # messaggio di un terzo si persisteva `human` (clodia-platform#248).
        # `clodia-logic` poteva solo coercirlo in lettura, che è un ponte: ogni
        # lettore futuro deve ricordarsene. L'etichetta si scrive dove nasce.
        umano = current_principal() if is_on_behalf() else None
        res = svc.post_message(a["tier"], a["name"],
                               author=umano or ag or "agente",
                               text=text, kind=message_kind())
        import re as _re
        # CHI FA PARTIRE UN TURNO. La condizione «solo se c'è una @menzione» era
        # giusta finché a postare c'erano solo agenti: un agente che deposita una
        # bolla (una mail in arrivo, un handoff) non deve svegliare nessuno, e
        # senza quel filtro due agenti che si parlano si rispondono all'infinito.
        #
        # Per una PERSONA è sbagliata, e in un modo che avevo dichiarato di voler
        # evitare. Dalla webui un messaggio umano senza menzione instrada
        # comunque, per rilevanza; da un client MCP restava senza risposta e
        # bisognava riscriverlo a mano nella webui. Cioè il client era diventato
        # **una seconda porta sulla stessa stanza con regole diverse** — la cosa
        # per cui avevo rifiutato di aggiungere un `topic.ask`, ricomparsa dove
        # non la stavo guardando.
        #
        # Chi decide resta uno solo: `channel_trigger` porta il messaggio al
        # router di clodia-logic, che applica le regole di ieri — menzione a un
        # umano → nessuna AI, tag → quell'agente, altrimenti rilevanza con il
        # tetto di una risposta sola. Qui non si sceglie il destinatario: si
        # smette di decidere al posto suo.
        if (umano and text.strip()) or _re.search(r"@[a-z0-9][a-z0-9_-]{0,30}", text):
            try:
                runtime.channel_trigger(a["tier"], a["name"], text,
                                        by=umano or ag or "")
                res["triggered"] = True
            except Exception as e:  # noqa: BLE001 — il post resta valido anche se il trigger fallisce
                res["trigger_error"] = str(e)[:120]
        return res
    if verb == "messages":
        return {"messages": svc.list_messages(a["tier"], a["name"],
                                              limit=int(a.get("limit") or 200))}
    if verb in ("my_mentions", "mark_seen"):
        # «Chi mi ha chiamato?» — la domanda è sempre su CHI CHIEDE, mai su un
        # nome passato come argomento: un parametro renderebbe il verbo un modo
        # per leggere la casella di un altro. Per un agente l'identità è il suo
        # nome, per una persona il principal firmato.
        chi = (current_principal() if is_on_behalf() else agent_name()) or ""
        if not chi:
            raise PermissionError("menzioni: chiamante non identificato")
        if verb == "my_mentions":
            return svc.my_mentions(a["tier"], a["name"], chi,
                                   limit=int(a.get("limit") or 50),
                                   only_unseen=bool(a.get("only_unseen", True)))
        return svc.mark_seen(a["tier"], a["name"], chi, a.get("seen_through") or "")
    if verb == "save_summary":
        return svc.save_summary(a["tier"], a["name"], a["text"], a.get("base_version"))
    if verb == "save_agents_md":
        return svc.save_agents_md(a["tier"], a["name"], a["text"], a.get("base_version"))
    if verb == "add_minute":
        return svc.add_minute(a["tier"], a["name"], a["text"])
    if verb == "archive":
        return svc.archive(a["tier"], a["name"])
    if verb == "telegram_bind":
        return svc.telegram_bind(a["tier"], a["name"], a["chat_id"],
                                 mode=a.get("mode") or "excerpt",
                                 people=a.get("people"),
                                 mount_name=a.get("mount"))
    if verb == "telegram_unbind":
        return svc.telegram_unbind(a["tier"], a["name"], a.get("mount"))
    if verb == "set_portable":
        return svc.set_portable(a["tier"], a["name"], bool(a["portable"]))
    if verb in ("list", "search"):
        # `list` e `search` non passano da `_require_topic_member`: filtrano da
        # sé, per membership del chiamante. E il chiamante lo leggevano da
        # `agent_name()` — il CARRIER — anche quando il token era di una persona
        # legata a una stanza. In esercizio il carrier è `clodia`, che partecipa
        # a tutto: dal client di Giovanni una ricerca rispondeva con i titoli dei
        # topic di clodia. Non un accesso ai contenuti, ma una mappa di stanze
        # che non lo riguardano — e una perdita che non somiglia a un errore,
        # perché una lista di titoli sembra sempre plausibile.
        if _token_is_bound_to_a_room():
            solo = current_chat()[len("chan:"):].split(":")[:2]
            righe = (svc.list(a.get("tier"), a.get("include_archived", False))
                     if verb == "list" else svc.search(a["query"], a.get("mode", "lexical")))
            if not isinstance(righe, list):
                return righe
            return [r for r in righe
                    if str(r.get("tier")) == solo[0] and str(r.get("name")) == solo[1]]
        if verb == "list":
            return _filter_member_rows(
                svc.list(a.get("tier"), a.get("include_archived", False)), agent_name())
        res = svc.search(a["query"], a.get("mode", "lexical"))
        return _filter_member_rows(res, agent_name()) if isinstance(res, list) else res
    if verb == "files":
        return svc.list_files(a["tier"], a["name"], a.get("subpath", ""))
    if verb == "read_file":
        data = svc.read_file(a["tier"], a["name"], a["path"])
        try:
            return {"path": a["path"], "encoding": "utf-8", "content": data.decode("utf-8")}
        except UnicodeDecodeError:
            # File binario: NON riversare base64 grossi nel contesto (si tronca, brucia
            # token, spesso fallisce). Sopra soglia → indirizza a topic.fetch (copia nello
            # scratch, byte fuori dal modello). Vedi anche topic.read_document per il testo.
            if len(data) > _B64_INLINE_CAP:
                return {"ok": False, "path": a["path"], "size": len(data),
                        "error": (f"file binario di {len(data)} byte: troppo grande per "
                                  "read_file (base64 nel contesto). USA topic.fetch(tier, name, "
                                  f"path='{a['path']}', dest=<path nel tuo scratch>) e lavora sul "
                                  "file locale; per il solo testo usa topic.read_document.")}
            import base64 as _b64
            return {"path": a["path"], "encoding": "base64",
                    "content": _b64.b64encode(data).decode("ascii"),
                    "note": "file binario (PDF/immagine/...): decodifica da base64"}
    if verb == "read_document":
        data = svc.read_file(a["tier"], a["name"], a["path"])
        cap = int(a.get("max_chars") or 60000)
        try:
            text, pages = _extract_document_text(a["path"].rsplit("/", 1)[-1], data)
        except Exception as e:  # noqa: BLE001
            return {"ok": False, "error": f"estrazione fallita: {str(e)[:160]}"}
        trunc = len(text) > cap
        return {"path": a["path"], "text": text[:cap], "chars": len(text),
                "pages": pages, "truncated": trunc}
    if verb == "write_file":
        fn = a["filename"]
        enc = (a.get("encoding") or "text").lower()
        ext = fn.rsplit(".", 1)[-1].lower() if "." in fn else ""
        # Un file con estensione binaria NON può essere testo: lo decodifichiamo
        # sempre come base64 (anche se l'agente ha dimenticato encoding='base64'),
        # con errore chiaro se il base64 è malformato. Era questo il bug del
        # Travel_reimbursement.xlsx: base64 scritto come testo → file corrotto.
        if enc == "base64" or ext in _BINARY_EXTS:
            data = _decode_b64_strict(a["content"], fn)
            # Base64 grosso nei parametri = anti-pattern (ARG_MAX/troncamento): se hai
            # già il file nello scratch, caricalo con topic.put (il gateway legge i byte
            # dal path, niente base64 nel modello).
            if len(data) > _B64_INLINE_CAP:
                return {"ok": False, "filename": fn, "size": len(data),
                        "error": (f"payload di {len(data)} byte troppo grande per write_file. "
                                  "Scrivi il file nel tuo scratch e usa topic.put(tier, name, "
                                  f"filename='{fn}', src=<path nel tuo scratch>).")}
        else:
            data = (a["content"] or "").encode("utf-8")
        # Output dell'agente: provenienza `agent`, non `untrusted`. Non è
        # contenuto di terzi per sé — se il canale è già contaminato la
        # contaminazione è del canale, e non serve ri-etichettarla sul file.
        return svc.put_file(a["tier"], a["name"], fn, data, "agent", agent_name())
    if verb == "fetch":
        # I byte attraversano il solo volume /shared come envelope cifrato per
        # lo spawn destinatario; agent-server decifra e materializza `dest`.
        data = svc.read_file(a["tier"], a["name"], a["path"])
        chat_id = current_chat()
        if not chat_id:
            raise ValueError("topic.fetch richiede una sessione agent con chat_id")
        # `dest` opzionale: senza, il file prende il proprio nome dentro lo
        # scratch. Pretenderlo obbligatorio e assoluto costringeva l'agente a
        # indovinare un path d'infrastruttura che non conosce.
        _dest = str(a.get("dest") or "").strip() or _os.path.basename(a["path"])
        return transfer_channel.fetch_to_agent(
            data, chat_id=chat_id, dest=_dest, sender=agent_name())
    if verb == "put":
        # agent-server legge lo scratch della sola sessione, cifra per il gateway
        # e deposita un envelope effimero su /shared; qui viene decifrato e consumato.
        chat_id = current_chat()
        if not chat_id:
            raise ValueError("topic.put richiede una sessione agent con chat_id")
        data = transfer_channel.put_from_agent(chat_id=chat_id, src=a["src"])
        return svc.put_file(a["tier"], a["name"], a["filename"], data,
                            "agent", agent_name())
    if verb == "delete_file":
        return svc.delete_file(a["tier"], a["name"], a["path"])
    if verb == "migrate_storage":
        return svc.migrate_storage(a["tier"], a["name"], a["target"])
    # Remote pluggable (git/drive): storage sempre local, sync opzionale/manuale.
    if verb == "remote_status":
        return svc.remote_status(a["tier"], a["name"], a.get("mount"))
    if verb == "remote_enable":
        return svc.remote_enable(a["tier"], a["name"], a["type"], a.get("config"),
                                 mount_name=a.get("mount"))
    if verb == "remote_disable":
        return svc.remote_disable(a["tier"], a["name"], a.get("mount"))
    if verb == "remote_add":
        return svc.remote_add(a["tier"], a["name"], a["path"], a.get("mount"))
    if verb == "remote_commit":
        return svc.remote_commit(a["tier"], a["name"], a.get("message", ""), a.get("mount"))
    if verb == "remote_push":
        return svc.remote_push(a["tier"], a["name"], a.get("mount"))
    if verb == "remote_pull":
        return svc.remote_pull(a["tier"], a["name"], a.get("mount"))
    raise ValueError(f"unknown topic verb: {name}")


async def main():
    try:
        agent = agent_name()
        print(f"[mcp-tools-server v{__version__}] serving agent={agent}", file=sys.stderr)
    except PermissionError as e:
        print(f"[mcp-tools-server v{__version__}] {e}", file=sys.stderr)
        sys.exit(2)
    async with stdio_server() as (read, write):
        await app.run(read, write, app.create_initialization_options())


if __name__ == "__main__":
    asyncio.run(main())
